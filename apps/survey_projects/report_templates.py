"""
Ministry-prescribed DGDE survey-report template generator.

Produces the standard "Report on Survey of Defence Land" .docx for a SurveyArea,
with a clickable (bookmark-linked) index and the 17 prescribed sections. Dynamic
sections — the Statement of Survey Numbers and Area Computation — are auto-filled
from the area's GIS features; narrative sections carry the standard DGDE text with
fill-in blanks the surveyor completes in OnlyOffice.
"""
from __future__ import annotations

import io
import json
import re
from datetime import date


# Attribute keys we try (case/format-insensitive) when reading a survey number or
# the Raksha Bhoomi ID off a feature's attribute table.
_SN_KEYS = ('survey_number', 'survey_no', 'surveyno', 'sy_no', 'syno',
            'khasra', 'khasra_no', 'plot_no', 'plot')
_RB_KEYS = ('raksha_bhoomi_id', 'raksha_bhoomi', 'raksha_id', 'rb_id', 'rbid',
            'rb_no', 'rbno', 'rb', 'rakshabhoomi_id', 'rb_code', 'office_code')
# Raksha Bhoomi IDs look like "DCHEN10001" — a short uppercase office code + digits.
_RB_VALUE_RE = re.compile(r'^[A-Z]{2,6}\d{3,}$')


def _attr_lookup(attrs: dict, keys) -> str:
    if not isinstance(attrs, dict):
        return ''
    low = {str(k).lower().replace(' ', '_'): v for k, v in attrs.items()}
    for k in keys:
        v = low.get(k)
        if v not in (None, ''):
            return str(v)
    return ''


def _detect_rb_id(attrs: dict) -> str:
    """Raksha Bhoomi ID by known key, else by value pattern (e.g. DCHEN10001)."""
    v = _attr_lookup(attrs, _RB_KEYS)
    if v:
        return v
    if isinstance(attrs, dict):
        for val in attrs.values():
            s = str(val).strip()
            if _RB_VALUE_RE.match(s):
                return s
    return ''


def _is_lonlat(x, y) -> bool:
    return -180.0 <= x <= 180.0 and -85.06 <= y <= 85.06


def _area_acres(geom) -> float:
    """Area of a (WGS-84) geometry in acres, via the UTM zone of its centroid."""
    try:
        c = geom.centroid
        # Guard against features whose coordinates are not valid lon/lat (e.g. stored
        # in a projected CRS) — otherwise the UTM zone is out of range (invalid EPSG).
        if not _is_lonlat(c.x, c.y):
            return 0.0
        zone = min(max(int((c.x + 180) / 6) + 1, 1), 60)
        g = geom.transform(32600 + zone, clone=True)
        return g.area / 4046.8564224  # m² → acres
    except Exception:
        return 0.0


def _gather_features(area, user):
    """Return (rows[(survey_no, acres)], total_acres, raksha_bhoomi_id, geoms)."""
    from .analysis import _get_folder_ids_for_survey_areas, _scope_survey_qs
    from .models import GISFeature
    folder_ids = _get_folder_ids_for_survey_areas([area.id])
    qs = _scope_survey_qs(
        GISFeature.objects.filter(is_deleted=False, folder_id__in=folder_ids), user
    ).only('geometry', 'attributes')
    rows, total, rb_id, geoms = [], 0.0, '', []
    for f in qs[:5000]:
        attrs = f.attributes if isinstance(f.attributes, dict) else {}
        if not rb_id:
            rb_id = _detect_rb_id(attrs)
        sn = _attr_lookup(attrs, _SN_KEYS)
        ac = _area_acres(f.geometry)
        total += ac
        if f.geometry is not None:
            geoms.append(f.geometry)
        if sn or ac:
            rows.append((sn or '—', ac))
    rows.sort(key=lambda r: (r[0] == '—', r[0]))
    return rows, total, rb_id, geoms


# ── Logo + basemap-with-features image rendering (Pillow + OSM tiles) ──────────

def _fetch_logo_bytes(url: str):
    """Return raw bytes of the configured branding logo (URL or local path)."""
    if not url:
        return None
    try:
        if url.startswith('http://') or url.startswith('https://'):
            import httpx
            r = httpx.get(url, timeout=6.0, headers={'User-Agent': 'RakshaGIS-Report/1.0'})
            if r.status_code == 200:
                return r.content
        else:
            import os
            from django.conf import settings
            rel = url.split('?')[0].lstrip('/')
            rel = rel[len('media/'):] if rel.startswith('media/') else rel
            for base in (settings.MEDIA_ROOT, getattr(settings, 'STATIC_ROOT', None),
                         str(settings.BASE_DIR)):
                if not base:
                    continue
                p = os.path.join(str(base), rel)
                if os.path.exists(p):
                    with open(p, 'rb') as fh:
                        return fh.read()
    except Exception:
        pass
    return None


def _geom_rings(geom):
    """Exterior rings [[(lon,lat),...]] for Polygon/MultiPolygon; [] otherwise."""
    try:
        gj = json.loads(geom.geojson)
    except Exception:
        return []
    t, c = gj.get('type'), gj.get('coordinates')
    rings = []
    if t == 'Polygon' and c:
        rings.append(c[0])
    elif t == 'MultiPolygon' and c:
        for poly in c:
            if poly:
                rings.append(poly[0])
    return rings


def _render_area_map_png(geoms, target_w=1000, max_tiles=6, tile_url=None):
    """
    Render the area's feature outlines over an OSM basemap and return PNG bytes.
    Stitches web-mercator tiles for the feature bbox and draws the polygons on top.
    Falls back to a plain light background (outline only) if tiles are unavailable.
    """
    import math
    from PIL import Image, ImageDraw

    rings, minlon, minlat, maxlon, maxlat = [], 1e9, 1e9, -1e9, -1e9
    for g in geoms or []:
        for ring in _geom_rings(g):
            # Skip rings whose coordinates aren't valid lon/lat — they would push the
            # bbox/tile math out of range (negative tiles, invalid zoom).
            if not ring or any(not _is_lonlat(pt[0], pt[1]) for pt in ring):
                continue
            rings.append(ring)
            for lon, lat in ring:
                minlon, maxlon = min(minlon, lon), max(maxlon, lon)
                minlat, maxlat = min(minlat, lat), max(maxlat, lat)
    if not rings or minlon > maxlon:
        return None

    dlon = (maxlon - minlon) or 0.002
    dlat = (maxlat - minlat) or 0.002
    minlon -= dlon * 0.15; maxlon += dlon * 0.15
    minlat -= dlat * 0.15; maxlat += dlat * 0.15

    def deg2px(lon, lat, z):
        n = 2 ** z
        x = (lon + 180.0) / 360.0 * n * 256.0
        latr = math.radians(lat)
        y = (1.0 - math.asinh(math.tan(latr)) / math.pi) / 2.0 * n * 256.0
        return x, y

    zoom = 1
    for z in range(18, 0, -1):
        x0, _ = deg2px(minlon, maxlat, z)
        x1, _ = deg2px(maxlon, minlat, z)
        _, y0 = deg2px(minlon, maxlat, z)
        _, y1 = deg2px(maxlon, minlat, z)
        if (int(x1 // 256) - int(x0 // 256) + 1) <= max_tiles and \
           (int(y1 // 256) - int(y0 // 256) + 1) <= max_tiles:
            zoom = z
            break

    px0, py0 = deg2px(minlon, maxlat, zoom)
    px1, py1 = deg2px(maxlon, minlat, zoom)
    tx0, ty0 = int(px0 // 256), int(py0 // 256)
    tx1, ty1 = int(px1 // 256), int(py1 // 256)
    ox, oy = tx0 * 256, ty0 * 256

    canvas = Image.new('RGB', ((tx1 - tx0 + 1) * 256, (ty1 - ty0 + 1) * 256), (236, 236, 232))
    url_tpl = tile_url or 'https://tile.openstreetmap.org/{z}/{x}/{y}.png'
    try:
        import httpx
        fails = 0
        with httpx.Client(timeout=httpx.Timeout(5.0, connect=3.0),
                          headers={'User-Agent': 'RakshaGIS-Report/1.0'}) as client:
            for tx in range(tx0, tx1 + 1):
                if fails >= 2:
                    break  # tiles unreachable (offline) → abort fast, fall back to plain bg
                for ty in range(ty0, ty1 + 1):
                    if fails >= 2:
                        break
                    try:
                        r = client.get(url_tpl.format(z=zoom, x=tx, y=ty))
                        if r.status_code == 200:
                            tile = Image.open(io.BytesIO(r.content)).convert('RGB')
                            canvas.paste(tile, ((tx - tx0) * 256, (ty - ty0) * 256))
                        else:
                            fails += 1
                    except Exception:
                        fails += 1
    except Exception:
        pass

    overlay = Image.new('RGBA', canvas.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    for ring in rings:
        pts = [(deg2px(lon, lat, zoom)[0] - ox, deg2px(lon, lat, zoom)[1] - oy)
               for lon, lat in ring]
        if len(pts) >= 3:
            draw.polygon(pts, fill=(255, 102, 0, 60))
            draw.line(pts + [pts[0]], fill=(200, 30, 0, 255), width=3)
    canvas = Image.alpha_composite(canvas.convert('RGBA'), overlay).convert('RGB')

    cx0, cy0 = max(0, int(px0 - ox)), max(0, int(py0 - oy))
    cx1, cy1 = min(canvas.width, int(px1 - ox)), min(canvas.height, int(py1 - oy))
    if cx1 > cx0 and cy1 > cy0:
        canvas = canvas.crop((cx0, cy0, cx1, cy1))
    if canvas.width:
        canvas = canvas.resize((target_w, max(1, int(canvas.height * target_w / canvas.width))),
                               Image.LANCZOS)
    out = io.BytesIO()
    canvas.save(out, 'PNG')
    out.seek(0)
    return out.read()


# ── docx low-level helpers (bookmarks + internal links for a clickable index) ──

def _add_bookmark(paragraph, name, bm_id):
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bm_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bm_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text, anchor):
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    h = OxmlElement('w:hyperlink')
    h.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)


# Section titles (the 17 prescribed sections + their indicative page ranges).
_SECTIONS = [
    ('Introduction', '1-2'),
    ('Objectives', '2'),
    ('Scope of Survey Work and Demarcation', '2'),
    ('Details of Defence Land Pocket', '3-6'),
    ('Area Outcome / Findings of Phase-I and Phase-II', '6-9'),
    ('Action Taken on the Findings / Recommendations of Phase-I & II', '9-11'),
    ('Methodology — CoRS Real-Time Service', '11-12'),
    ('Advantages of CoRS Network (RTK) Survey', '12'),
    ('Preparation of Base Maps (Methodology)', '12-13'),
    ('Base Maps', '13-14'),
    ('Statement of Survey Numbers (MLR/Revenue) before/after Land Settlement/Consolidation',
     '14-23'),
    ('Statement of Peripheral Survey Numbers before/after Settlement/Consolidation', '23-27'),
    ('Constitution of Survey Team (Revenue/Settlement Department)', '27-28'),
    ('Actual Ground Survey Team (Reference Points & Demarcation)', '28'),
    ('List of Coordinates & Preparation of Map', '29-38'),
    ('Area Computation', '39'),
    ('Comparison Statement of Area Outcome — Phase-I, Phase-II with Phase-III', '39-40'),
]


def build_ministry_survey_report(area, user) -> bytes:
    """Build the ministry-format survey report .docx and return its bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    project = area.project
    org = getattr(project, 'organisation', None)
    office_name = org.name if org else '<Office Name>'
    year = date.today().year
    rows, total_acres, rb_id, geoms = _gather_features(area, user)
    rb_display = rb_id or '<Fetch from Attribute Table>'

    # Branding logo + basemap-with-features image (best-effort; None on failure).
    logo_bytes = None
    try:
        from apps.core.models import BrandingConfig
        from django.conf import settings as _s
        logo_bytes = _fetch_logo_bytes(getattr(BrandingConfig.get_solo(), 'logo_url', ''))
        tile_url = getattr(_s, 'REPORT_BASEMAP_TILE_URL', None)
    except Exception:
        tile_url = None
    try:
        map_png = _render_area_map_png(geoms, tile_url=tile_url)
    except Exception:
        map_png = None

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    def centered(text, size=14, bold=True, color=None, after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return p

    def centered_image(data, width_inches):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(data), width=Inches(width_inches))
        return p

    # ── Title page ────────────────────────────────────────────────────────────
    if logo_bytes:
        centered_image(logo_bytes, 1.6)
    else:
        centered('[ National Emblem of India ]            [ Defence Estates Organisation Logo ]',
                 size=10, bold=False, after=24)
    centered('Report on Survey of Defence Land at', size=16)
    centered(area.name or '<Survey Area Name>', size=16)
    centered(project.name or '<Project Name>', size=14)
    centered(f'YEAR OF SURVEY: {year}', size=13)
    doc.add_paragraph()
    centered(f'Raksha Bhoomi ID: {rb_display}', size=13)
    doc.add_paragraph()
    # Outline shapefile over basemap.
    if map_png:
        centered_image(map_png, 5.6)
    else:
        centered('< Outline shape file with basemap — insert map image here >',
                 size=11, bold=False, after=24)
    doc.add_paragraph()
    centered('GOVERNMENT OF INDIA', size=13)
    centered('MINISTRY OF DEFENCE', size=13)
    centered('Office of the Defence Estates Officer', size=12)
    centered(office_name, size=12)
    doc.add_page_break()

    # ── Index (clickable) ───────────────────────────────────────────────────────
    h = doc.add_heading('INDEX OF THE CONTENT', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run('SL. NO.').bold = True
    hdr[1].paragraphs[0].add_run('INDEX OF THE CONTENT').bold = True
    hdr[2].paragraphs[0].add_run('PAGE NO.').bold = True
    for i, (title, page) in enumerate(_SECTIONS, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        _add_internal_link(cells[1].paragraphs[0], title, f'sec{i}')
        cells[2].text = page
    doc.add_page_break()

    # ── Section helpers ────────────────────────────────────────────────────────
    def section(idx, title):
        p = doc.add_heading(f'{idx}. {title}', level=1)
        _add_bookmark(p, f'sec{idx}', 100 + idx)
        return p

    def body(text):
        return doc.add_paragraph(text)

    def bullets(items):
        for it in items:
            doc.add_paragraph(it, style='List Bullet')

    def placeholder(note):
        p = doc.add_paragraph()
        run = p.add_run(f'[ {note} ]')
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def grid(headers, data_rows, note=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        for c, head in zip(t.rows[0].cells, headers):
            c.paragraphs[0].add_run(head).bold = True
        for row in data_rows:
            cells = t.add_row().cells
            for c, val in zip(cells, row):
                c.text = '' if val is None else str(val)
        if note:
            placeholder(note)
        return t

    # ── 1. Introduction (from the prescribed template) ──────────────────────────
    section(1, 'Introduction')
    body('MoD has tasked DGDE to complete the survey of ________ lakh acres of Defence lands '
         'spread over ________ pockets. Survey of the entire Defence land in the country was '
         'completed twice as Phase-I and Phase-II survey using latest survey technology like '
         'Electronic Total Station (ETS) and Differential Global Positioning System (DGPS) and in '
         'a large number of pockets in association with various State Governments.')
    body('For the first time, Drone imagery-based survey technology was used for survey of lakhs of '
         'acres of Defence land. Besides this, Satellite Imagery based survey was done for certain '
         'pockets. As a result of these technological interventions, the survey progressed at a '
         'faster pace.')
    body('Ministry of Defence has developed a web portal (Land Management System) wherein the '
         'updated GIS Maps of Defence lands across the country can be accessed. The GIS maps, '
         'digitized shape files and updated land data have enabled MoD to access land data on a '
         'real-time basis, reducing delays in decisions on land optimization, transfer for '
         'infrastructure projects and new land acquisitions.')
    body('DGDE have issued directions for carrying out Phase-III survey of Defence land such that '
         '1/3rd part of Defence land held by DEOs is surveyed each calendar year w.e.f. 2024 till '
         '2026. Phase-III has been conducted using Continuously Operating Reference Stations (CoRS) '
         'enabled DGPS equipment with reference to the CoRS Network established by Survey of India '
         '(SoI) to fix the boundary of Defence land.')

    # ── 2. Objectives ───────────────────────────────────────────────────────────
    section(2, 'Objectives')
    bullets([
        'Preparation of Base Maps.',
        'Identification of peripheral survey numbers of the base maps on latest revenue maps.',
        'Demarcation of the peripheral survey numbers of Defence land on ground as per base map in '
        'association with State Revenue Authorities.',
        'Identification of the Defence land boundary as per the actual ground survey and generating '
        'CoRS-enabled digital maps.',
        'Identify the discrepancy in the area and alignment of Defence land boundary and report the '
        'same in the Survey Report.',
        'Create a Centralized GIS-based database of Defence land spread across the jurisdiction of DEOs.',
        'Identify the actual Defence land boundary as per the base maps.',
        'Preparation of survey report & resolution of discrepancies.',
        'Updation of Defence land records in MLR/GLR/Raksha Bhoomi software/RTRM etc.',
    ])

    # ── 3. Scope of Survey Work and Demarcation ─────────────────────────────────
    section(3, 'Scope of Survey Work and Demarcation')
    bullets([
        'Demarcation of actual Defence land boundary in association with Revenue Authorities.',
        'Start DGPS ground survey using latest technology e.g. CoRS Network for the entire Defence land.',
        'GIS-based database integration, validation, incorporation of field survey data on GIS '
        'platform for the Defence land.',
        'Integrate GIS Map for visualisation and references.',
        'Digitized in GIS environment; update the map, boundary limits of encroached area & update '
        'the status of land records.',
    ])

    # ── 4. Details of Defence Land Pocket ───────────────────────────────────────
    section(4, 'Details of Defence Land Pocket')
    body('Details of the Defence land pocket(s) covered under this survey are tabulated below.')
    grid(['S.No.', 'Pocket / Site Name', 'Village / Location', 'GLR/MLR Reference',
          'No. of Survey Numbers', 'Area (acres)'],
         [['1', area.name, '< village/location >', '< GLR/MLR ref >',
           str(len(rows)), f'{total_acres:.3f}']],
         note='Add a row per Defence land pocket and complete the location/GLR-MLR references.')

    # ── 5. Area Outcome / Findings of Phase-I and Phase-II ──────────────────────
    section(5, 'Area Outcome / Findings of Phase-I and Phase-II')
    body('Summary of area outcomes and findings recorded during Phase-I and Phase-II surveys:')
    grid(['Phase', 'Area as per Records (acres)', 'Area as Surveyed (acres)', 'Difference (acres)',
          'Key Findings'],
         [['Phase-I', '', '', '', ''], ['Phase-II', '', '', '', '']],
         note='Fill the recorded vs surveyed areas and summarise findings for each phase.')

    # ── 6. Action Taken on Findings / Recommendations of Phase-I & II ───────────
    section(6, 'Action Taken on the Findings / Recommendations of Phase-I & II')
    body('Action taken on the discrepancies, encroachments and recommendations identified during '
         'Phase-I and Phase-II is summarised below.')
    grid(['S.No.', 'Finding / Recommendation', 'Action Taken', 'Status'],
         [['1', '', '', '']],
         note='Add a row per finding/recommendation with the corresponding action and status.')

    # ── 7. Methodology — CoRS Real-Time Service ─────────────────────────────────
    section(7, 'Methodology — CoRS Real-Time Service')
    body('The survey was carried out using the Continuously Operating Reference Stations (CoRS) '
         'network established by Survey of India. CoRS provides real-time differential corrections '
         '(RTK) to roving DGPS receivers over GSM/NTRIP, enabling centimetre-level positional '
         'accuracy without the need for a local base station.')
    body('Field procedure: the rover is initialised against the nearest CoRS station, fixed '
         '(RTK-fixed) solutions are confirmed, and boundary pillars / turning points of the Defence '
         'land are observed and recorded. Observed points are exported and processed into the GIS '
         'environment for map preparation and area computation.')

    # ── 8. Advantages of CoRS Network (RTK) Survey ──────────────────────────────
    section(8, 'Advantages of CoRS Network (RTK) Survey')
    bullets([
        'Centimetre-level (RTK) positional accuracy referenced to the national CoRS datum.',
        'No requirement to set up and occupy a local base station — faster mobilisation.',
        'Real-time fixed solutions reduce post-processing time and field re-visits.',
        'Consistent, repeatable coordinates that integrate directly with the GIS database.',
        'Reliable demarcation of boundary pillars and detection of peripheral encroachment.',
    ])

    # ── 9. Preparation of Base Maps (Methodology) ───────────────────────────────
    section(9, 'Preparation of Base Maps (Methodology)')
    body('Base maps were prepared by overlaying the CoRS-surveyed boundary on the latest revenue '
         'maps and existing GLR/MLR records. Survey numbers were identified and georeferenced, the '
         'Defence land boundary was digitized in the GIS environment, and discrepancies in area and '
         'alignment were flagged for resolution with the State Revenue Authorities.')

    # ── 10. Base Maps ───────────────────────────────────────────────────────────
    section(10, 'Base Maps')
    if map_png:
        body('Base map of the survey area showing the surveyed Defence land boundary over the basemap.')
        centered_image(map_png, 6.2)
    else:
        placeholder('Insert the prepared base map(s) / outline shapefile over basemap for this survey area.')

    # ── 11. Statement of Survey Numbers (auto-filled) ───────────────────────────
    section(11, 'Statement of Survey Numbers (MLR/Revenue) before/after Land '
                'Settlement/Consolidation')
    body('Survey numbers comprising the Defence land in this area, with their corresponding area '
         '(auto-compiled from the GIS features; verify against MLR/Revenue records).')
    data = [[str(i), sn, f'{ac:.3f}', '', '']
            for i, (sn, ac) in enumerate(rows, start=1)] or [['1', '', '', '', '']]
    grid(['S.No.', 'Survey No. (MLR/Revenue)', 'Area as Surveyed (acres)',
          'Area before Settlement', 'Area after Settlement'], data)
    body(f'Total number of survey numbers: {len(rows)}.  '
         f'Total surveyed area: {total_acres:.3f} acres '
         f'({total_acres * 0.404686:.3f} hectares).')

    # ── 12. Statement of Peripheral Survey Numbers ──────────────────────────────
    section(12, 'Statement of Peripheral Survey Numbers before/after '
                'Settlement/Consolidation')
    grid(['S.No.', 'Peripheral Survey No.', 'Adjoining Owner / Use', 'Area (acres)',
          'Before Settlement', 'After Settlement'],
         [['1', '', '', '', '', '']],
         note='List the peripheral (boundary-adjoining) survey numbers and their areas.')

    # ── 13. Constitution of Survey Team (Revenue/Settlement) ────────────────────
    section(13, 'Constitution of Survey Team (Revenue/Settlement Department)')
    grid(['S.No.', 'Name', 'Designation', 'Department / Office'],
         [['1', '', '', '']],
         note='List the Revenue/Settlement Department officials constituting the survey team.')

    # ── 14. Actual Ground Survey Team ───────────────────────────────────────────
    section(14, 'Actual Ground Survey Team (Reference Points & Demarcation)')
    body('Detail of the ground survey team, reference points used, and demarcation of the actual/true '
         'boundary of the Defence land pocket by the State Revenue Authorities.')
    grid(['S.No.', 'Name', 'Role', 'Reference Point / Remarks'],
         [['1', '', '', '']],
         note='Record team members, roles and the reference points / CoRS stations used.')

    # ── 15. List of Coordinates & Preparation of Map ────────────────────────────
    section(15, 'List of Coordinates & Preparation of Map')
    body('List of surveyed boundary coordinates (turning points / pillars) used to prepare the map.')
    grid(['Point ID', 'Survey No.', 'Latitude', 'Longitude', 'Easting (m)', 'Northing (m)'],
         [['P1', '', '', '', '', '']],
         note='Populate the boundary/turning-point coordinates from the field survey.')

    # ── 16. Area Computation (auto-filled) ──────────────────────────────────────
    section(16, 'Area Computation')
    body('Area computed from the surveyed GIS geometry of this area.')
    grid(['Description', 'Area (acres)', 'Area (hectares)', 'Area (sq. m)'],
         [['Total surveyed Defence land', f'{total_acres:.3f}',
           f'{total_acres * 0.404686:.3f}', f'{total_acres * 4046.8564224:.1f}']],
         note='Computed from feature geometry; reconcile with MLR/Revenue figures.')

    # ── 17. Comparison Statement Phase-I/II/III ─────────────────────────────────
    section(17, 'Comparison Statement of Area Outcome — Phase-I, Phase-II with Phase-III')
    grid(['Phase', 'Area Outcome (acres)', 'Difference vs Phase-III (acres)', 'Remarks'],
         [['Phase-I', '', '', ''],
          ['Phase-II', '', '', ''],
          ['Phase-III (this survey)', f'{total_acres:.3f}', '—', 'Auto-computed from GIS']],
         note='Enter Phase-I and Phase-II area outcomes to complete the comparison.')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

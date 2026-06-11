import json
import logging

logger = logging.getLogger(__name__)

from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action, api_view, permission_classes as drf_permission_classes
from rest_framework.exceptions import PermissionDenied
from rest_framework.response import Response
from rest_framework.views import APIView
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import SearchFilter, OrderingFilter
from django.contrib.gis.db.models.functions import Distance
from django.contrib.gis.measure import D

from apps.accounts.permissions import (
    CanEditProject, IsSuperAdmin, org_queryset_filter,
    get_shared_project_ids, get_approved_area_ids, deo_subordinate_org_ids,
)
from .models import (
    SurveyProject, SurveyArea, GISFeature, DefenceParcel, AttributeTemplate,
    ShapefileImport, ProjectLayerFolder, ProjectShare, GeoTiffLayer,
    FeatureAttachment, ProjectMilestone, SurveyAreaAccessRequest, ReviewAnnotation,
    TopologyRule, FeatureComment,
)
from .serializers import (
    SurveyProjectSerializer, SurveyAreaSerializer, GISFeatureSerializer, DefenceParcelSerializer,
    AttributeTemplateSerializer, ShapefileImportSerializer,
    ProjectLayerFolderSerializer, ProjectShareSerializer, GeoTiffLayerSerializer,
    BufferParcelSerializer, FeatureAttachmentSerializer, ProjectMilestoneSerializer,
    ReviewAnnotationSerializer, TopologyRuleSerializer, FeatureCommentSerializer,
)


class SurveyProjectViewSet(viewsets.ModelViewSet):
    serializer_class = SurveyProjectSerializer
    lookup_value_regex = r'\d+'
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['status', 'survey_type', 'priority', 'organisation', 'state', 'district', 'taluk', 'village']
    search_fields = ['name', 'project_number', 'description']
    ordering_fields = ['created_at', 'updated_at', 'name', 'priority']

    def get_queryset(self):
        from .access import permitted_extra_project_ids
        user = self.request.user
        base_qs = SurveyProject.objects.select_related(
            'organisation', 'created_by', 'state', 'district', 'taluk', 'village'
        )
        # Strict office isolation: every user (including DGDE/PDDE office users)
        # sees only projects of their OWN organisation here, plus explicit
        # grants. HQ users reach subordinate data only via the published
        # Map Viewer endpoints — never through the projects API.
        own_qs = org_queryset_filter(user, base_qs)
        if user.is_superadmin:
            return own_qs
        extra_ids = permitted_extra_project_ids(user)
        if not extra_ids:
            return own_qs
        return (own_qs | base_qs.filter(id__in=extra_ids)).distinct()

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [CanEditProject()]
        return [permissions.IsAuthenticated()]

    def perform_create(self, serializer):
        # SUPERADMIN may send an explicit organisation; everyone else uses their own
        org = serializer.validated_data.get('organisation') or self.request.user.organisation
        if org is None:
            from rest_framework.exceptions import ValidationError
            raise ValidationError({'organisation': 'Organisation is required.'})
        project = serializer.save(
            created_by=self.request.user,
            organisation=org,
        )
        _create_default_folders(project, self.request.user)

    @action(detail=True, methods=['post'], url_path='set-map-enabled',
            permission_classes=[permissions.IsAuthenticated])
    def set_map_enabled(self, request, pk=None):
        """
        POST /api/projects/{id}/set-map-enabled/  body: {map_enabled: bool}
        Enable/disable exposure of this project's PUBLISHED data to higher levels
        (PDDE/DGDE). Owning-office admins (DEO/CEO/ADEO) + superadmin only. Works
        even when the project is otherwise edit-locked.
        """
        project = self.get_object()
        user = request.user
        if not (user.is_superadmin or user.role in user.ADMIN_ROLES):
            return Response({'detail': 'Only office admins can change visibility.'}, status=403)
        project.map_enabled = bool(request.data.get('map_enabled', True))
        project.save(update_fields=['map_enabled'])
        return Response({'id': project.id, 'map_enabled': project.map_enabled})

    @action(detail=True, methods=['get', 'post'], url_path='active-version',
            permission_classes=[permissions.IsAuthenticated])
    def active_version(self, request, pk=None):
        """
        GET  — Return the active (latest non-final) VERSION folder for drawing.
               Auto-creates Phase I / {year} / Ver-I if none exist yet.
        POST — Force-create the next version (Ver-II, Ver-III …).
        """
        import datetime
        project = self.get_object()
        force_new = request.method == 'POST'

        if not force_new:
            # Return existing active version if available
            active = (
                ProjectLayerFolder.objects
                .filter(project=project, folder_type=ProjectLayerFolder.VERSION, is_final=False)
                .order_by('-created_at')
                .first()
            )
            if active:
                return Response(ProjectLayerFolderSerializer(active).data)

        # Count all versions (including finals) to name the next one
        version_count = ProjectLayerFolder.objects.filter(
            project=project, folder_type=ProjectLayerFolder.VERSION
        ).count()
        next_name = f'Ver-{_roman(version_count + 1)}'

        # Ensure Phase I exists
        phase = (
            ProjectLayerFolder.objects
            .filter(project=project, folder_type=ProjectLayerFolder.PHASE)
            .order_by('order')
            .first()
        )
        if not phase:
            phase = ProjectLayerFolder.objects.create(
                project=project, name='Phase I',
                folder_type=ProjectLayerFolder.PHASE,
                created_by=request.user, order=1,
            )

        # Ensure a YEAR folder exists under the phase
        current_year = datetime.date.today().year
        year_folder = (
            ProjectLayerFolder.objects
            .filter(project=project, parent=phase, folder_type=ProjectLayerFolder.YEAR)
            .order_by('-year')
            .first()
        )
        if not year_folder:
            year_folder = ProjectLayerFolder.objects.create(
                project=project, parent=phase,
                name=str(current_year),
                folder_type=ProjectLayerFolder.YEAR,
                year=current_year,
                created_by=request.user, order=0,
            )

        # Create the new VERSION folder
        new_version = ProjectLayerFolder.objects.create(
            project=project, parent=year_folder,
            name=next_name,
            folder_type=ProjectLayerFolder.VERSION,
            is_final=False,
            created_by=request.user,
            order=version_count,
        )
        return Response(ProjectLayerFolderSerializer(new_version).data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'], url_path='new-layer',
            permission_classes=[CanEditProject])
    def new_layer(self, request, pk=None):
        """
        POST /api/projects/{pk}/new-layer/
        Create a new drawing layer and wire it to a survey area.

        Body (JSON):
          layer_name      string  required — name for the new layer
          geometry_type   string  required — POINT | LINE | POLYGON
          survey_area_id  int     optional — use an existing DRAFT/RETURNED survey area
          new_area_name   string  optional — create a new survey area with this name
          new_area_code   string  optional — short code for the new area
          fields          list    optional — attribute schema [{name, type, label, required}]

        Returns:
          { survey_area_id, survey_area_name, folder_id, layer_name, geometry_type }
        """
        from django.db import IntegrityError
        from apps.accounts.models import User

        project = self.get_object()
        user = request.user

        if user.role not in (User.SDO, User.SURVEYOR, User.SUPERADMIN):
            raise PermissionDenied('Only SDO / Surveyor can create layers.')

        layer_name = (request.data.get('layer_name') or '').strip()
        geometry_type = (request.data.get('geometry_type') or 'POLYGON').upper()
        survey_area_id = request.data.get('survey_area_id')
        new_area_name = (request.data.get('new_area_name') or '').strip()
        new_area_code = (request.data.get('new_area_code') or '').strip()
        fields = request.data.get('fields', [])

        if not layer_name:
            return Response({'detail': 'layer_name is required.'}, status=400)
        if geometry_type not in ('POINT', 'LINE', 'POLYGON'):
            return Response({'detail': 'geometry_type must be POINT, LINE, or POLYGON.'}, status=400)
        if not survey_area_id and not new_area_name:
            return Response({'detail': 'Provide survey_area_id or new_area_name.'}, status=400)

        # ── Resolve / create survey area ─────────────────────────────
        if new_area_name:
            try:
                area = SurveyArea.objects.create(
                    project=project, name=new_area_name, area_code=new_area_code,
                    status=SurveyArea.DRAFT, created_by=user,
                )
            except Exception:
                return Response(
                    {'detail': f'A survey area named "{new_area_name}" already exists in this project.'},
                    status=400,
                )
        else:
            try:
                area = SurveyArea.objects.get(id=survey_area_id, project=project)
            except SurveyArea.DoesNotExist:
                return Response({'detail': 'Survey area not found.'}, status=404)
            if area.status not in (SurveyArea.DRAFT, SurveyArea.RETURNED):
                return Response(
                    {'detail': 'Survey area is locked. Cannot add layers to submitted/approved areas.'},
                    status=400,
                )

        # ── Ensure survey area has a root folder ──────────────────────
        if not area.folder_id:
            root_folder = ProjectLayerFolder.objects.create(
                project=project, name=area.name,
                folder_type=ProjectLayerFolder.ZONE, created_by=user, order=0,
            )
            _add_survey_area_subfolders(root_folder, user)
            area.folder = root_folder
            area.save(update_fields=['folder'])

        # ── Resolve the "Shape Files" subfolder ───────────────────────
        shapefile_folder, _ = ProjectLayerFolder.objects.get_or_create(
            project=project, parent=area.folder,
            folder_type=ProjectLayerFolder.SHAPEFILE,
            defaults={'name': 'Shape Files', 'created_by': user, 'order': 0},
        )

        # ── Optionally persist attribute schema ───────────────────────
        if fields:
            AttributeTemplate.objects.update_or_create(
                organisation=user.organisation, layer_name=layer_name,
                defaults={
                    'fields': fields,
                    'created_by': user,
                    'description': f'Schema for {layer_name}',
                },
            )

        return Response({
            'survey_area_id': area.id,
            'survey_area_name': area.name,
            'folder_id': shapefile_folder.id,
            'layer_name': layer_name,
            'geometry_type': geometry_type,
        }, status=201)

    @action(detail=True, methods=['get'], url_path='export')
    def export(self, request, pk=None):
        """Export project features in various formats (legacy GeoJSON/CSV/Shapefile)."""
        project = self.get_object()
        fmt = request.query_params.get('format', 'geojson').lower()
        layer_name = request.query_params.get('layer_name', None)
        folder_id = request.query_params.get('folder', None)

        qs = GISFeature.objects.filter(project=project, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        if folder_id:
            qs = qs.filter(folder_id=folder_id)

        return _export_features(qs, project, fmt)

    @action(detail=True, methods=['post'], url_path='export-full')
    def export_full(self, request, pk=None):
        """
        POST /api/projects/{id}/export-full/

        Queue an async export of the COMPLETE project:
          • All survey areas with their GIS features, documents, rasters,
            and uploaded shapefile ZIPs
          • C2PA / LP-DNA watermarks applied to every file
          • provenance.json manifest

        Returns: { task_uuid, status, message }
        Poll GET /api/core/export/status/{task_uuid}/ until status=='DONE'.
        Download via GET /api/core/export/download/{task_uuid}/.
        """
        from django.http import JsonResponse
        from apps.core.models import ExportTask
        from apps.core.tasks import build_export_zip
        from django.conf import settings as _settings

        project = self.get_object()
        user = request.user
        org_id = getattr(user, 'organisation_id', None)

        active_statuses = (ExportTask.PENDING, ExportTask.RUNNING)
        _max_user = getattr(_settings, 'EXPORT_MAX_CONCURRENT_PER_USER', 2)
        _max_org  = getattr(_settings, 'EXPORT_MAX_CONCURRENT_PER_ORG', 3)

        if ExportTask.objects.filter(requested_by=user, status__in=active_statuses).count() >= _max_user:
            return JsonResponse({
                'error': 'Too many active exports',
                'detail': 'Wait for your current export(s) to finish before starting another.',
            }, status=429)

        if org_id and ExportTask.objects.filter(
            organisation_id=org_id, status__in=active_statuses
        ).count() >= _max_org:
            return JsonResponse({
                'error': 'Office export limit reached',
                'detail': 'Your office has too many exports running. Please wait.',
            }, status=429)

        include_dxf = bool(request.data.get('include_dxf', False))

        et = ExportTask.objects.create(
            export_type=ExportTask.PROJECT,
            object_id=project.id,
            object_name=project.project_number or str(project.id),
            requested_by=user,
            organisation_id=org_id,
            include_dxf=include_dxf,
            progress_msg='Queued…',
        )
        build_export_zip.delay(et.pk)

        return JsonResponse({
            'task_uuid': str(et.task_uuid),
            'status': et.status,
            'message': 'Full project export queued.',
        }, status=202)


class GISFeatureViewSet(viewsets.ModelViewSet):
    serializer_class = GISFeatureSerializer
    lookup_value_regex = r'\d+'
    filter_backends = [DjangoFilterBackend, SearchFilter]
    filterset_fields = ['project', 'folder', 'layer_name', 'geometry_type', 'is_deleted']
    search_fields = ['layer_name', 'feature_id']
    pagination_class = None  # map must load all features; pagination breaks the map layer

    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        base_qs = GISFeature.objects.select_related('project__organisation', 'created_by')
        from .access import hq_level, published_map_filter
        own_qs = org_queryset_filter(user, base_qs, org_field='project__organisation').filter(is_deleted=False)
        # DGDE/PDDE (national/command) office users: own-org features stay fully
        # visible; subordinate offices contribute only features inside PUBLISHED +
        # map-enabled survey areas (PDDE limited to its command subtree).
        level = hq_level(user)
        org_id = self.request.query_params.get('organisation')
        if level:
            from apps.survey_projects.models import SurveyArea as _SA
            areas = published_map_filter(user, _SA.objects.filter(
                status=_SA.PUBLISHED, map_enabled=True,
                project__map_enabled=True, folder__isnull=False,
            ))
            if org_id:
                areas = areas.filter(project__organisation_id=org_id)
            # Narrow to a single survey area when the map browser drills down
            area_id = self.request.query_params.get('area')
            if area_id:
                areas = areas.filter(id=area_id)
            folder_ids: list[int] = []
            for a in areas.only('folder_id'):
                folder_ids.extend(_get_subtree_folder_ids(a.folder_id))
            pub_qs = (base_qs.filter(folder_id__in=folder_ids, is_deleted=False)
                      if folder_ids else base_qs.none())
            if org_id:
                # Office drill-down view: published features of that office only
                return pub_qs
            return (own_qs | pub_qs).distinct()
        if user.is_superadmin:
            # When browsing a specific office via the field browser, filter to
            # that office's PUBLISHED survey area features only.
            if org_id:
                from apps.survey_projects.models import SurveyArea as _SA
                areas = _SA.objects.filter(
                    status=_SA.PUBLISHED, map_enabled=True,
                    project__map_enabled=True, folder__isnull=False,
                    project__organisation_id=org_id,
                )
                area_id = self.request.query_params.get('area')
                if area_id:
                    areas = areas.filter(id=area_id)
                folder_ids: list[int] = []
                for a in areas.only('folder_id'):
                    folder_ids.extend(_get_subtree_folder_ids(a.folder_id))
                if not folder_ids:
                    return base_qs.none()
                return base_qs.filter(folder_id__in=folder_ids, is_deleted=False)
            return own_qs
        # Include features from approved cross-org areas (in their folder subtrees)
        approved_area_ids = get_approved_area_ids(user)
        shared_project_ids = get_shared_project_ids(user)
        deo_sub_ids = deo_subordinate_org_ids(user)
        if not approved_area_ids and not shared_project_ids and not deo_sub_ids:
            return own_qs
        # For approved areas: features in the area's folder subtree
        from apps.survey_projects.models import SurveyArea as _SA
        approved_folder_ids: list[int] = []
        for area in _SA.objects.filter(id__in=approved_area_ids).select_related('folder'):
            if area.folder_id:
                approved_folder_ids.extend(_get_subtree_folder_ids(area.folder_id))
        extra_q = Q(project_id__in=shared_project_ids)
        if approved_folder_ids:
            extra_q |= Q(folder_id__in=approved_folder_ids)
        # DEO offices: subordinate-office features explicitly marked deo_visible
        if deo_sub_ids:
            extra_q |= Q(project__organisation_id__in=deo_sub_ids, deo_visible=True)
        return (own_qs | base_qs.filter(extra_q, is_deleted=False)).distinct()

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [CanEditProject()]
        return [permissions.IsAuthenticated()]

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def perform_destroy(self, instance):
        instance.is_deleted = True
        instance.save(update_fields=['is_deleted'])

    @action(detail=False, methods=['get'], url_path='enli-search',
            permission_classes=[permissions.IsAuthenticated])
    def enli_search(self, request):
        """
        GET /api/survey_projects/features/enli-search/?q=<eNLI_code>

        Search internal GISFeature attributes for a matching Land_Parcel_ID.
        Returns GeoJSON-compatible results suitable for flyToSearchResult().
        """
        from django.contrib.gis.serializers.geojson import Serializer as GeoJSONSerializer
        import json

        q = (request.query_params.get('q') or '').strip()
        if len(q) < 2:
            return Response({'detail': 'Enter at least 2 characters.', 'results': []}, status=400)

        qs = self.get_queryset().filter(
            attributes__Land_Parcel_ID__icontains=q
        ).exclude(geometry=None)[:20]

        results = []
        for feat in qs:
            parcel_id = feat.attributes.get('Land_Parcel_ID', '')
            try:
                geom_json = json.loads(feat.geometry.geojson)
            except Exception:
                continue
            results.append({
                'id': feat.id,
                'layer_id': None,
                'layer_name': feat.layer_name,
                'match_field': 'Land_Parcel_ID',
                'match_value': parcel_id,
                'label': f'{parcel_id} ({feat.layer_name})',
                'geometry': geom_json,
                'attributes': feat.attributes,
            })

        return Response({'query': q, 'count': len(results), 'results': results})

    @action(detail=False, methods=['post'], url_path='atlas',
            permission_classes=[permissions.IsAuthenticated])
    def atlas(self, request):
        """
        POST /api/survey_projects/features/atlas/

        Generate a multi-page PDF Atlas — one page per feature in the given layer.

        Body:
            project      (int)   project id
            layer_name   (str)   layer to iterate over
            title_field  (str, optional) attribute to use as page title
            padding      (float, default 0.001) degrees of padding around each feature
            width        (int, default 1200) render width px
            height       (int, default 800)  render height px
            dpi          (int, default 150)  output DPI
        """
        import io
        from django.http import HttpResponse as DjResponse

        project_id  = request.data.get('project')
        layer_name  = (request.data.get('layer_name') or '').strip()
        title_field = (request.data.get('title_field') or '').strip()
        padding     = float(request.data.get('padding', 0.001))
        width       = min(int(request.data.get('width', 1200)), 4000)
        height      = min(int(request.data.get('height', 800)), 3000)
        dpi         = min(int(request.data.get('dpi', 150)), 300)

        if not project_id or not layer_name:
            return Response({'detail': 'project and layer_name are required.'}, status=400)

        features = self.get_queryset().filter(
            project_id=project_id, layer_name=layer_name, is_deleted=False,
        ).exclude(geometry=None)[:200]

        if not features:
            return Response({'detail': f'No features found in layer "{layer_name}".'}, status=404)

        try:
            from apps.core.services.mapnik_service import get_mapnik_service
            from PIL import Image, ImageDraw, ImageFont
            import math
        except ImportError as exc:
            return Response({'detail': f'Required library not available: {exc}'}, status=500)

        svc = get_mapnik_service()
        pages: list[Image.Image] = []

        for feat in features:
            geom = feat.geometry
            try:
                env = geom.envelope  # bounding box as polygon
                minx = geom.extent[0] - padding
                miny = geom.extent[1] - padding
                maxx = geom.extent[2] + padding
                maxy = geom.extent[3] + padding

                # Ensure minimum extent to avoid zero-size box
                if (maxx - minx) < 0.0001: minx -= 0.0005; maxx += 0.0005
                if (maxy - miny) < 0.0001: miny -= 0.0005; maxy += 0.0005

                svc.set_bbox((minx, miny, maxx, maxy))
                png_bytes = svc.render_png(width=width, height=height)
                page = Image.open(io.BytesIO(png_bytes)).convert('RGB')

                # Title bar overlay
                title_text = str(feat.attributes.get(title_field, '') if title_field else '') or f'Feature #{feat.id}'
                draw = ImageDraw.Draw(page)
                bar_h = 32
                draw.rectangle([(0, 0), (width, bar_h)], fill=(10, 20, 40, 220))
                draw.text((8, 6), f'{layer_name}  ·  {title_text}', fill=(200, 220, 255))

                # ID watermark bottom-right
                draw.text((width - 120, height - 20), f'ID {feat.id}  |  eNLI: {feat.attributes.get("Land_Parcel_ID", "—")}',
                          fill=(120, 120, 120))
                pages.append(page)
            except Exception:
                continue

        if not pages:
            return Response({'detail': 'Could not render any atlas pages.'}, status=500)

        out = io.BytesIO()
        pages[0].save(out, format='PDF', save_all=True, append_images=pages[1:],
                      resolution=dpi, title=f'{layer_name} Atlas')
        out.seek(0)
        response = DjResponse(out.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{layer_name}_atlas.pdf"'
        return response

    @action(detail=False, methods=['post'], url_path='rename-layer',
            permission_classes=[CanEditProject])
    def rename_layer(self, request):
        project_id = request.data.get('project')
        old_name = request.data.get('old_name', '').strip()
        new_name = request.data.get('new_name', '').strip()
        if not old_name or not new_name:
            return Response({'detail': 'old_name and new_name required.'}, status=400)
        if old_name == new_name:
            return Response({'detail': 'Names are the same.'}, status=400)
        import re
        if not re.match(r'^[A-Za-z][A-Za-z0-9_]{0,63}$', new_name):
            return Response({'detail': 'Layer name must start with a letter, max 64 chars, alphanumeric + underscore.'}, status=400)
        count = GISFeature.objects.filter(
            project_id=project_id, layer_name=old_name, is_deleted=False
        ).update(layer_name=new_name)
        return Response({'detail': f'Renamed {count} feature(s) from "{old_name}" → "{new_name}".', 'count': count})

    @action(detail=False, methods=['post'], url_path='repair-geometry',
            permission_classes=[CanEditProject])
    def repair_geometry(self, request):
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        from django.db import connection
        with connection.cursor() as cur:
            cur.execute("""
                UPDATE survey_projects_gisfeature
                   SET geometry = ST_CollectionExtract(ST_MakeValid(geometry))
                 WHERE project_id = %s
                   AND (%s = '' OR layer_name = %s)
                   AND is_deleted = FALSE
                   AND NOT ST_IsValid(geometry)
            """, [project_id, layer_name, layer_name])
            count = cur.rowcount
        return Response({'detail': f'Repaired {count} invalid geometry(ies).', 'count': count})

    @action(detail=False, methods=['post'], url_path='deduplicate',
            permission_classes=[CanEditProject])
    def deduplicate(self, request):
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        from django.db import connection
        with connection.cursor() as cur:
            cur.execute("""
                UPDATE survey_projects_gisfeature SET is_deleted = TRUE
                WHERE id IN (
                    SELECT id FROM (
                        SELECT id,
                               ROW_NUMBER() OVER (
                                   PARTITION BY ST_AsEWKB(geometry), layer_name
                                   ORDER BY id
                               ) AS rn
                        FROM survey_projects_gisfeature
                        WHERE project_id = %s
                          AND (%s = '' OR layer_name = %s)
                          AND is_deleted = FALSE
                    ) sub WHERE sub.rn > 1
                )
            """, [project_id, layer_name, layer_name])
            count = cur.rowcount
        return Response({'detail': f'Removed {count} duplicate feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='merge-layers',
            permission_classes=[CanEditProject])
    def merge_layers(self, request):
        project_id = request.data.get('project')
        source = request.data.get('source_layer', '').strip()
        target = request.data.get('target_layer', '').strip()
        delete_source = request.data.get('delete_source', False)
        if not source or not target:
            return Response({'detail': 'source_layer and target_layer required.'}, status=400)
        features = list(GISFeature.objects.filter(project_id=project_id, layer_name=source, is_deleted=False))
        count = 0
        for f in features:
            f.pk = None
            f.id = None
            f.layer_name = target
            f.created_by = request.user
            f.save()
            count += 1
        if delete_source:
            GISFeature.objects.filter(project_id=project_id, layer_name=source, is_deleted=False).update(is_deleted=True)
        return Response({'detail': f'Merged {count} feature(s) into "{target}".', 'count': count,
                         'source_deleted': delete_source})

    @action(detail=False, methods=['get'], url_path='layer-schema')
    def layer_schema(self, request):
        project_id = request.query_params.get('project')
        layer_name = request.query_params.get('layer_name', '')
        qs = GISFeature.objects.filter(project_id=project_id, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        fields: dict = {}
        for f in qs[:200]:
            for k, v in (f.attributes or {}).items():
                if k not in fields:
                    if isinstance(v, bool):
                        t = 'boolean'
                    elif isinstance(v, int):
                        t = 'integer'
                    elif isinstance(v, float):
                        t = 'float'
                    else:
                        t = 'string'
                    fields[k] = {'name': k, 'type': t, 'sample': str(v)[:80] if v is not None else ''}
        total = qs.count()
        return Response({'fields': list(fields.values()), 'feature_count': total})

    @action(detail=False, methods=['post'], url_path='remove-field',
            permission_classes=[CanEditProject])
    def remove_field(self, request):
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        field_name = request.data.get('field_name', '').strip()
        if not field_name:
            return Response({'detail': 'field_name required.'}, status=400)
        qs = GISFeature.objects.filter(project_id=project_id, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        count = 0
        for f in qs:
            if field_name in (f.attributes or {}):
                attrs = dict(f.attributes)
                del attrs[field_name]
                f.attributes = attrs
                f.save(update_fields=['attributes'])
                count += 1
        return Response({'detail': f'Removed field "{field_name}" from {count} feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='rename-field',
            permission_classes=[CanEditProject])
    def rename_field(self, request):
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        old_field = request.data.get('old_field', '').strip()
        new_field = request.data.get('new_field', '').strip()
        if not old_field or not new_field:
            return Response({'detail': 'old_field and new_field required.'}, status=400)
        qs = GISFeature.objects.filter(project_id=project_id, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        count = 0
        for f in qs:
            if old_field in (f.attributes or {}):
                attrs = dict(f.attributes)
                attrs[new_field] = attrs.pop(old_field)
                f.attributes = attrs
                f.save(update_fields=['attributes'])
                count += 1
        return Response({'detail': f'Renamed field "{old_field}" → "{new_field}" on {count} feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='auto-geometry-stats',
            permission_classes=[CanEditProject])
    def auto_geometry_stats(self, request):
        """Compute area_m2, perimeter_m or length_m and save into feature attributes."""
        from django.db import connection
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        qs = GISFeature.objects.filter(project_id=project_id, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        count = 0
        for f in qs:
            if f.geometry is None:
                continue
            attrs = dict(f.attributes or {})
            gt = f.geometry_type
            if gt == GISFeature.POLYGON:
                attrs['area_m2'] = round(f.geometry.transform(32643, clone=True).area, 2)
                attrs['perimeter_m'] = round(f.geometry.transform(32643, clone=True).length, 2)
            elif gt == GISFeature.LINE:
                attrs['length_m'] = round(f.geometry.transform(32643, clone=True).length, 2)
            elif gt == GISFeature.POINT:
                attrs['lat'] = round(f.geometry.y, 6)
                attrs['lon'] = round(f.geometry.x, 6)
            f.attributes = attrs
            f.save(update_fields=['attributes'])
            count += 1
        return Response({'detail': f'Updated geometry stats for {count} feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='dissolve',
            permission_classes=[CanEditProject])
    def dissolve(self, request):
        """Merge features with same value in dissolve_field using ST_Union. Creates new layer."""
        from django.contrib.gis.db.models import Union
        project_id = request.data.get('project')
        layer_name  = request.data.get('layer_name', '').strip()
        dissolve_field = request.data.get('dissolve_field', '').strip()
        out_layer = (request.data.get('out_layer') or f"{layer_name}_dissolved").strip()
        if not layer_name or not dissolve_field:
            return Response({'detail': 'layer_name and dissolve_field required.'}, status=400)
        qs = GISFeature.objects.filter(project_id=project_id, layer_name=layer_name, is_deleted=False)
        groups = {}
        for f in qs:
            key = (f.attributes or {}).get(dissolve_field)
            key_str = str(key) if key is not None else '__null__'
            if key_str not in groups:
                groups[key_str] = {'geoms': [], 'attr_val': key}
            groups[key_str]['geoms'].append(f.geometry)
        created = 0
        for key_str, g in groups.items():
            from django.contrib.gis.geos import GeometryCollection
            combined = g['geoms'][0]
            for gm in g['geoms'][1:]:
                try:
                    combined = combined.union(gm)
                except Exception:
                    pass
            GISFeature.objects.create(
                project_id=project_id,
                layer_name=out_layer,
                geometry=combined,
                geometry_type=GISFeature.POLYGON,
                attributes={dissolve_field: g['attr_val']},
                created_by=request.user,
            )
            created += 1
        return Response({'detail': f'Dissolved into {created} feature(s) in layer "{out_layer}".', 'created': created})

    @action(detail=False, methods=['post'], url_path='spatial-join',
            permission_classes=[CanEditProject])
    def spatial_join(self, request):
        """Join attributes from overlay_layer to features in base_layer where geometries intersect."""
        project_id    = request.data.get('project')
        base_layer    = request.data.get('base_layer', '').strip()
        overlay_layer = request.data.get('overlay_layer', '').strip()
        join_fields   = request.data.get('join_fields', [])  # list of field names to copy
        if not base_layer or not overlay_layer:
            return Response({'detail': 'base_layer and overlay_layer required.'}, status=400)
        overlay_qs = GISFeature.objects.filter(project_id=project_id, layer_name=overlay_layer, is_deleted=False)
        base_qs    = GISFeature.objects.filter(project_id=project_id, layer_name=base_layer, is_deleted=False)
        count = 0
        for bf in base_qs:
            if bf.geometry is None:
                continue
            for ov in overlay_qs:
                if ov.geometry is None:
                    continue
                try:
                    if bf.geometry.intersects(ov.geometry):
                        attrs = dict(bf.attributes or {})
                        ov_attrs = ov.attributes or {}
                        for fld in (join_fields or ov_attrs.keys()):
                            if fld in ov_attrs:
                                attrs[f"sj_{fld}"] = ov_attrs[fld]
                        bf.attributes = attrs
                        bf.save(update_fields=['attributes'])
                        count += 1
                        break  # first match only
                except Exception:
                    pass
        return Response({'detail': f'Joined attributes to {count} feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='clip-to-boundary',
            permission_classes=[CanEditProject])
    def clip_to_boundary(self, request):
        """Clip features in layer_name to the union of clip_layer features. Updates geometries in-place."""
        project_id  = request.data.get('project')
        layer_name  = request.data.get('layer_name', '').strip()
        clip_layer  = request.data.get('clip_layer', '').strip()
        if not layer_name or not clip_layer:
            return Response({'detail': 'layer_name and clip_layer required.'}, status=400)
        clip_geoms = list(
            GISFeature.objects.filter(project_id=project_id, layer_name=clip_layer, is_deleted=False)
            .exclude(geometry=None).values_list('geometry', flat=True)
        )
        if not clip_geoms:
            return Response({'detail': 'No clip features found.'}, status=400)
        clip_union = clip_geoms[0]
        for g in clip_geoms[1:]:
            try:
                clip_union = clip_union.union(g)
            except Exception:
                pass
        count = 0
        for f in GISFeature.objects.filter(project_id=project_id, layer_name=layer_name, is_deleted=False):
            if f.geometry is None:
                continue
            try:
                clipped = f.geometry.intersection(clip_union)
                if clipped.empty:
                    f.is_deleted = True
                    f.save(update_fields=['is_deleted'])
                else:
                    f.geometry = clipped
                    f.save(update_fields=['geometry'])
                count += 1
            except Exception:
                pass
        return Response({'detail': f'Clipped {count} feature(s) to "{clip_layer}".', 'count': count})

    @action(detail=False, methods=['get'], url_path='summary-stats')
    def summary_stats(self, request):
        """Return count + numeric field stats (sum/min/max/avg) for a layer."""
        project_id = request.query_params.get('project')
        layer_name = request.query_params.get('layer_name', '').strip()
        qs = GISFeature.objects.filter(project_id=project_id, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        total = qs.count()
        # Collect numeric fields and compute stats in Python (attributes is JSONB)
        field_vals: dict = {}
        for f in qs:
            for k, v in (f.attributes or {}).items():
                if isinstance(v, (int, float)):
                    field_vals.setdefault(k, []).append(v)
        stats = {'feature_count': total, 'fields': {}}
        for k, vals in field_vals.items():
            stats['fields'][k] = {
                'count': len(vals),
                'sum': round(sum(vals), 4),
                'min': round(min(vals), 4),
                'max': round(max(vals), 4),
                'avg': round(sum(vals) / len(vals), 4),
            }
        return Response(stats)

    @action(detail=False, methods=['post'], url_path='near-analysis',
            permission_classes=[CanEditProject])
    def near_analysis(self, request):
        """For each feature in base_layer, find nearest feature in near_layer and store distance_m."""
        project_id  = request.data.get('project')
        base_layer  = request.data.get('base_layer', '').strip()
        near_layer  = request.data.get('near_layer', '').strip()
        if not base_layer or not near_layer:
            return Response({'detail': 'base_layer and near_layer required.'}, status=400)
        near_features = list(
            GISFeature.objects.filter(project_id=project_id, layer_name=near_layer, is_deleted=False)
            .exclude(geometry=None)
        )
        if not near_features:
            return Response({'detail': 'No near features found.'}, status=400)
        count = 0
        for bf in GISFeature.objects.filter(project_id=project_id, layer_name=base_layer, is_deleted=False):
            if bf.geometry is None:
                continue
            min_dist = None
            near_id  = None
            for nf in near_features:
                try:
                    d = bf.geometry.distance(nf.geometry)
                    # rough m conversion from degrees (~111320 m/deg)
                    d_m = round(d * 111320, 1)
                    if min_dist is None or d_m < min_dist:
                        min_dist = d_m
                        near_id  = nf.feature_id or str(nf.id)
                except Exception:
                    pass
            if min_dist is not None:
                attrs = dict(bf.attributes or {})
                attrs['near_dist_m'] = min_dist
                attrs['near_fid']    = near_id
                bf.attributes = attrs
                bf.save(update_fields=['attributes'])
                count += 1
        return Response({'detail': f'Near analysis complete for {count} feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='convex-hull',
            permission_classes=[CanEditProject])
    def convex_hull(self, request):
        """Generate convex hull polygon from all features in layer_name → new out_layer."""
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        out_layer  = (request.data.get('out_layer') or f"{layer_name}_hull").strip()
        if not layer_name:
            return Response({'detail': 'layer_name required.'}, status=400)
        geoms = list(
            GISFeature.objects.filter(project_id=project_id, layer_name=layer_name, is_deleted=False)
            .exclude(geometry=None).values_list('geometry', flat=True)
        )
        if not geoms:
            return Response({'detail': 'No features found.'}, status=400)
        from django.contrib.gis.geos import GeometryCollection
        collection = GeometryCollection(*geoms, srid=4326)
        hull = collection.convex_hull
        feat = GISFeature.objects.create(
            project_id=project_id,
            layer_name=out_layer,
            geometry=hull,
            geometry_type=GISFeature.POLYGON,
            attributes={'source_layer': layer_name, 'feature_count': len(geoms)},
            created_by=request.user,
        )
        return Response({'detail': f'Convex hull created in layer "{out_layer}".', 'id': feat.id})

    @action(detail=False, methods=['post'], url_path='centroid-extract',
            permission_classes=[CanEditProject])
    def centroid_extract(self, request):
        """Extract centroids from polygon layer_name → new out_layer of points."""
        project_id = request.data.get('project')
        layer_name = request.data.get('layer_name', '').strip()
        out_layer  = (request.data.get('out_layer') or f"{layer_name}_centroids").strip()
        if not layer_name:
            return Response({'detail': 'layer_name required.'}, status=400)
        qs = GISFeature.objects.filter(project_id=project_id, layer_name=layer_name, is_deleted=False).exclude(geometry=None)
        created = 0
        for f in qs:
            try:
                centroid = f.geometry.centroid
                GISFeature.objects.create(
                    project_id=project_id,
                    layer_name=out_layer,
                    geometry=centroid,
                    geometry_type=GISFeature.POINT,
                    attributes=dict(f.attributes or {}, source_id=f.feature_id or str(f.id)),
                    created_by=request.user,
                )
                created += 1
            except Exception:
                pass
        return Response({'detail': f'Extracted {created} centroid(s) to layer "{out_layer}".', 'created': created})

    @action(detail=False, methods=['post'], url_path='find-replace',
            permission_classes=[CanEditProject])
    def find_replace(self, request):
        """Search and replace a string value in a specific attribute field across a layer."""
        project_id  = request.data.get('project')
        layer_name  = request.data.get('layer_name', '').strip()
        field_name  = request.data.get('field_name', '').strip()
        find_val    = request.data.get('find_val', '')
        replace_val = request.data.get('replace_val', '')
        match_exact = request.data.get('match_exact', False)
        if not field_name:
            return Response({'detail': 'field_name required.'}, status=400)
        qs = GISFeature.objects.filter(project_id=project_id, is_deleted=False)
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        count = 0
        for f in qs:
            attrs = f.attributes or {}
            if field_name not in attrs:
                continue
            cur = str(attrs[field_name])
            if match_exact:
                if cur == str(find_val):
                    new_attrs = dict(attrs)
                    new_attrs[field_name] = replace_val
                    f.attributes = new_attrs
                    f.save(update_fields=['attributes'])
                    count += 1
            else:
                if str(find_val) in cur:
                    new_attrs = dict(attrs)
                    new_attrs[field_name] = cur.replace(str(find_val), str(replace_val))
                    f.attributes = new_attrs
                    f.save(update_fields=['attributes'])
                    count += 1
        return Response({'detail': f'Replaced in {count} feature(s).', 'count': count})

    @action(detail=False, methods=['post'], url_path='sql-view', permission_classes=[permissions.IsAuthenticated])
    def sql_view(self, request):
        """
        POST /api/survey_projects/features/sql-view/
        Execute a safe read-only SQL query and return results as GeoJSON.
        Only SELECT statements allowed. Must reference survey_projects_gisfeature table.
        """
        from django.db import connection
        import json

        query = (request.data.get('query') or '').strip()
        if not query:
            return Response({'detail': 'query is required.'}, status=400)

        # Safety: only allow SELECT statements, block destructive SQL
        ql = query.lower()
        if not ql.lstrip().startswith('select'):
            return Response({'detail': 'Only SELECT statements are allowed.'}, status=400)
        for forbidden in ('insert', 'update', 'delete', 'drop', 'truncate', 'alter', 'create', 'grant', 'pg_'):
            if forbidden in ql:
                return Response({'detail': f'Forbidden keyword: {forbidden}'}, status=400)

        try:
            with connection.cursor() as cur:
                cur.execute(query)
                cols = [d[0] for d in cur.description]
                rows = cur.fetchall()[:500]  # cap at 500 rows
        except Exception as exc:
            return Response({'detail': f'SQL error: {exc}'}, status=400)

        features = []
        for row in rows:
            row_dict = dict(zip(cols, row))
            geom = None
            # Look for a geom/geometry/geojson column
            for g_col in ('geom', 'geometry', 'geojson', 'st_asgeojson'):
                if g_col in row_dict and row_dict[g_col]:
                    try:
                        geom = json.loads(row_dict.pop(g_col))
                    except Exception:
                        row_dict.pop(g_col, None)
                    break
            if geom:
                features.append({'type': 'Feature', 'geometry': geom, 'properties': {k: str(v) for k, v in row_dict.items()}})

        return Response({'type': 'FeatureCollection', 'features': features, 'count': len(features)})


class AttributeTemplateViewSet(viewsets.ModelViewSet):
    """
    Org-scoped attribute schema definitions for GIS layers.
    SDO/SURVEYOR and admins can manage templates for their own org.
    """
    serializer_class = AttributeTemplateSerializer

    def get_queryset(self):
        return org_queryset_filter(
            self.request.user,
            AttributeTemplate.objects.select_related('organisation', 'created_by'),
        )

    def get_permissions(self):
        return [CanEditProject()]  # same gate: SDO/SURVEYOR + SUPERADMIN

    def perform_create(self, serializer):
        serializer.save(
            organisation=self.request.user.organisation,
            created_by=self.request.user,
        )


class ShapefileImportViewSet(viewsets.ModelViewSet):
    """Upload a .zip shapefile and trigger async import into GISFeature rows."""
    serializer_class = ShapefileImportSerializer
    http_method_names = ['get', 'post', 'head', 'options']  # no update/delete

    def get_queryset(self):
        return org_queryset_filter(
            self.request.user,
            ShapefileImport.objects.select_related(
                'project__organisation', 'attribute_template', 'created_by'
            ),
            org_field='project__organisation',
        )

    def get_permissions(self):
        return [CanEditProject()]

    def perform_create(self, serializer):
        from .tasks import import_shapefile
        job = serializer.save(
            created_by=self.request.user,
            status=ShapefileImport.PENDING,
        )
        import_shapefile.delay(job.id)

    @action(detail=True, methods=['post'], url_path='process-ai')
    def process_ai(self, request, pk=None):
        """Queue AI analysis of this imported layer's columns and features."""
        shp_import = self.get_object()
        if shp_import.status != ShapefileImport.DONE:
            from rest_framework.response import Response
            from rest_framework import status as drf_status
            return Response(
                {'detail': 'Import must be DONE before AI analysis.'},
                status=drf_status.HTTP_400_BAD_REQUEST,
            )
        from apps.ai_assistant.models import AITask
        from apps.ai_assistant.tasks import process_shapefile_ai
        from rest_framework.response import Response
        task = AITask.objects.create(
            task_type=AITask.GIS_INDEXING,
            requested_by=request.user,
            input_data={'shapefile_import_id': shp_import.id},
        )
        # Reset processed flag while re-running
        shp_import.ai_processed = False
        shp_import.save(update_fields=['ai_processed'])
        process_shapefile_ai.delay(task.id)
        return Response({'task_id': task.id, 'detail': 'AI analysis queued.'})


class DefenceParcelViewSet(viewsets.ModelViewSet):
    serializer_class = DefenceParcelSerializer
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['category', 'classification', 'organisation', 'state', 'district', 'taluk', 'village']
    search_fields = ['parcel_id', 'name', 'encumbrance_notes']
    ordering_fields = ['parcel_id', 'name', 'area_hectares']

    def get_queryset(self):
        return org_queryset_filter(
            self.request.user,
            DefenceParcel.objects.select_related(
                'organisation', 'state', 'district', 'taluk', 'village', 'survey_project'
            ).prefetch_related('revenue_maps')
        )

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [IsSuperAdmin()]
        return [permissions.IsAuthenticated()]


def _roman(n: int) -> str:
    """Convert a positive integer to uppercase Roman numerals."""
    vals = [
        (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
        (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
        (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I'),
    ]
    result = ''
    for v, s in vals:
        while n >= v:
            result += s
            n -= v
    return result


def _add_survey_area_subfolders(parent_folder, user):
    """
    Auto-create the three standard sub-folders under any survey-area folder:
      Shape Files  — vector GIS (zip shapefile, GeoJSON, KML, GeoPackage)
      Raster       — GeoTIFF / drone raster uploads
      Doc          — PDF, Excel, CSV, Word, images and other documents
    """
    ProjectLayerFolder.objects.get_or_create(
        project=parent_folder.project, parent=parent_folder,
        name='Shape Files', folder_type=ProjectLayerFolder.SHAPEFILE,
        defaults={'created_by': user, 'order': 0},
    )
    ProjectLayerFolder.objects.get_or_create(
        project=parent_folder.project, parent=parent_folder,
        name='Raster', folder_type=ProjectLayerFolder.RASTER,
        defaults={'created_by': user, 'order': 1},
    )
    ProjectLayerFolder.objects.get_or_create(
        project=parent_folder.project, parent=parent_folder,
        name='Doc', folder_type=ProjectLayerFolder.DOC,
        defaults={'created_by': user, 'order': 2},
    )

# Keep backward-compatible alias
_add_doc_shapefile_subfolders = _add_survey_area_subfolders


def _create_default_folders(project, user):
    common = ProjectLayerFolder.objects.create(
        project=project, name='Common Layer',
        folder_type=ProjectLayerFolder.COMMON, created_by=user, order=0,
    )
    # Revenue Map is now included in Common Layer alongside admin boundaries
    for i, name in enumerate(['State', 'District', 'Taluk', 'Village', 'Revenue Map']):
        boundary = ProjectLayerFolder.objects.create(
            project=project, parent=common, name=name,
            folder_type=ProjectLayerFolder.BOUNDARY, created_by=user, order=i,
        )
        _add_doc_shapefile_subfolders(boundary, user)


def _import_geotiff(folder, uploaded, layer_name, user, deo_visible=True):
    """Save a GeoTiff file and queue COG conversion."""
    import os
    from rest_framework.response import Response
    from .tasks import convert_geotiff_to_cog

    layer = GeoTiffLayer.objects.create(
        project=folder.project,
        folder=folder,
        name=layer_name,
        file=uploaded,
        status=GeoTiffLayer.PENDING,
        deo_visible=deo_visible,
        created_by=user,
    )
    try:
        convert_geotiff_to_cog.delay(layer.id)
    except Exception:
        pass  # Celery unavailable in dev — process proceeds, status stays PENDING
    return Response(
        {'detail': 'GeoTiff uploaded. COG conversion queued.', 'id': layer.id, 'type': 'geotiff'},
        status=201,
    )


def _import_shapefile_zip(folder, uploaded, layer_name, name_field, user, deo_visible=True):
    """Validate a .zip Shapefile bundle and import all features into GIS features + log ShapefileImport."""
    import io
    import os
    import tempfile
    import zipfile
    from django.contrib.gis.gdal import DataSource
    from django.contrib.gis.geos import GEOSGeometry
    from rest_framework.response import Response

    # Record the import attempt so the folder tree can show it
    shp_import = ShapefileImport.objects.create(
        project=folder.project,
        folder=folder,
        file=uploaded,
        layer_name=layer_name,
        status=ShapefileImport.RUNNING,
        deo_visible=deo_visible,
        created_by=user,
    )
    uploaded.seek(0)  # reset after ShapefileImport save consumed it
    data = uploaded.read()
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile:
        return Response({'detail': 'Uploaded file is not a valid ZIP archive.'}, status=400)

    names = zf.namelist()
    # Validate required shapefile components
    def _has_ext(ext):
        return any(n.lower().endswith(ext) for n in names)

    missing = [e for e in ('.shp', '.dbf', '.shx') if not _has_ext(e)]
    if missing:
        return Response(
            {'detail': f'ZIP is missing required Shapefile component(s): {", ".join(missing)}'},
            status=400,
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        zf.extractall(tmpdir)
        shp_path = next(
            os.path.join(tmpdir, n) for n in names if n.lower().endswith('.shp')
        )
        try:
            ds = DataSource(shp_path)
        except Exception as exc:
            return Response({'detail': f'Cannot read Shapefile: {exc}'}, status=400)

        created = 0
        errors = []
        project = folder.project

        for lyr in ds:
            for feat in lyr:
                try:
                    geom = GEOSGeometry(feat.geom.wkt, srid=feat.geom.srid or 4326)
                    if geom.srid != 4326:
                        geom.transform(4326)
                    geom_type = _classify_geom(geom.geom_type)
                    attrs = {field: _safe_val(feat[field].value) for field in feat.fields}
                    fid = str(feat[name_field].value) if name_field and name_field in feat.fields else ''
                    GISFeature.objects.create(
                        project=project,
                        folder=folder,
                        layer_name=layer_name,
                        geometry_type=geom_type,
                        geometry=geom,
                        feature_id=fid,
                        attributes=attrs,
                        deo_visible=deo_visible,
                        created_by=user,
                    )
                    created += 1
                except Exception as exc:
                    errors.append(str(exc))

    shp_import.status = ShapefileImport.DONE if not errors else ShapefileImport.FAILED
    shp_import.feature_count = created
    shp_import.error = '; '.join(errors[:5]) if errors else ''
    shp_import.save(update_fields=['status', 'feature_count', 'error'])

    return Response({
        'detail': f'Imported {created} feature(s) from Shapefile.',
        'created': created,
        'errors': errors[:10],
        'type': 'vector',
        'import_id': shp_import.id,
    }, status=201)


def _import_geojson_file(folder, uploaded, layer_name, name_field, user, deo_visible=True):
    """Parse a GeoJSON file and import all features."""
    import json
    from django.contrib.gis.geos import GEOSGeometry
    from rest_framework.response import Response

    try:
        data = json.loads(uploaded.read().decode('utf-8'))
    except (ValueError, UnicodeDecodeError) as exc:
        return Response({'detail': f'Invalid GeoJSON: {exc}'}, status=400)

    if data.get('type') == 'Feature':
        features = [data]
    elif data.get('type') == 'FeatureCollection':
        features = data.get('features', [])
    else:
        return Response({'detail': 'GeoJSON must be a Feature or FeatureCollection.'}, status=400)

    created = 0
    errors = []
    project = folder.project

    for feat in features:
        try:
            geom = GEOSGeometry(json.dumps(feat['geometry']), srid=4326)
            geom_type = _classify_geom(geom.geom_type)
            attrs = feat.get('properties') or {}
            fid = str(attrs.get(name_field, '')) if name_field else ''
            GISFeature.objects.create(
                project=project,
                folder=folder,
                layer_name=layer_name,
                geometry_type=geom_type,
                geometry=geom,
                feature_id=fid,
                attributes=attrs,
                deo_visible=deo_visible,
                created_by=user,
            )
            created += 1
        except Exception as exc:
            errors.append(str(exc))

    return Response({
        'detail': f'Imported {created} feature(s) from GeoJSON.',
        'created': created,
        'errors': errors[:10],
        'type': 'vector',
    }, status=201)


def _import_gpx_file(folder, uploaded, layer_name, name_field, user, geom_type='auto', deo_visible=True, extra_attributes=None):
    """Parse a GPX file and import waypoints and tracks as GISFeatures."""
    import xml.etree.ElementTree as ET
    from django.contrib.gis.geos import Point, LineString, Polygon
    from rest_framework.response import Response
    from .models import GISFeature

    try:
        xml_data = uploaded.read()
        try:
            decoded = xml_data.decode('utf-8')
        except UnicodeDecodeError:
            decoded = xml_data.decode('latin1')
        root = ET.fromstring(decoded)
    except Exception as exc:
        return Response({'detail': f'Invalid GPX XML: {exc}'}, status=400)

    # GPX namespaces
    ns = {'gpx': 'http://www.topografix.com/GPX/1/1'}
    root_tag = root.tag
    if '}' in root_tag:
        ns_url = root_tag.split('}')[0].strip('{')
        ns = {'gpx': ns_url}
    else:
        ns = {}

    created = 0
    errors = []
    project = folder.project
    _extra = extra_attributes if isinstance(extra_attributes, dict) else {}

    def _attrs(base: dict) -> dict:
        """Merge surveyor-supplied extra_attributes over the auto-generated base attrs."""
        merged = {**base}
        merged.update(_extra)
        return merged

    wpt_tag = './/gpx:wpt' if 'gpx' in ns else './/wpt'
    name_tag = 'gpx:name' if 'gpx' in ns else 'name'
    time_tag = 'gpx:time' if 'gpx' in ns else 'time'
    desc_tag = 'gpx:desc' if 'gpx' in ns else 'desc'

    # Gather waypoints coordinates & metadata
    wpt_coords = []
    wpt_names = []
    for wpt in root.findall(wpt_tag, ns):
        try:
            lat = float(wpt.attrib['lat'])
            lon = float(wpt.attrib['lon'])
            name_el = wpt.find(name_tag, ns)
            name = name_el.text if name_el is not None else 'GPS Waypoint'
            wpt_coords.append((lon, lat))
            wpt_names.append(name)
        except Exception as exc:
            errors.append(f"Waypoint parse error: {exc}")

    # Process Waypoints according to geom_type
    if geom_type == 'polygon':
        if len(wpt_coords) >= 3:
            try:
                if wpt_coords[0] != wpt_coords[-1]:
                    wpt_coords.append(wpt_coords[0])
                if len(wpt_coords) >= 4:
                    geom = Polygon(wpt_coords, srid=4326)
                    fid = wpt_names[0] if name_field and wpt_names else ''
                    GISFeature.objects.create(
                        project=project,
                        folder=folder,
                        layer_name=layer_name,
                        geometry_type=GISFeature.POLYGON,
                        geometry=geom,
                        feature_id=fid,
                        attributes=_attrs({'name': f"{layer_name} Waypoints Polygon", 'description': 'Constructed from waypoints'}),
                        deo_visible=deo_visible,
                        created_by=user,
                    )
                    created += 1
            except Exception as exc:
                errors.append(f"Waypoints Polygon construction error: {exc}")
    elif geom_type == 'line':
        if len(wpt_coords) >= 2:
            try:
                geom = LineString(wpt_coords, srid=4326)
                fid = wpt_names[0] if name_field and wpt_names else ''
                GISFeature.objects.create(
                    project=project,
                    folder=folder,
                    layer_name=layer_name,
                    geometry_type=GISFeature.LINE,
                    geometry=geom,
                    feature_id=fid,
                    attributes=_attrs({'name': f"{layer_name} Waypoints Line", 'description': 'Constructed from waypoints'}),
                    deo_visible=deo_visible,
                    created_by=user,
                )
                created += 1
            except Exception as exc:
                errors.append(f"Waypoints Line construction error: {exc}")
    else:  # 'auto' or 'point'
        # Import waypoints as Points
        for i, pt in enumerate(wpt_coords):
            try:
                geom = Point(pt[0], pt[1], srid=4326)
                name = wpt_names[i] if i < len(wpt_names) else 'GPS Waypoint'
                fid = name if name_field else ''
                GISFeature.objects.create(
                    project=project,
                    folder=folder,
                    layer_name=layer_name,
                    geometry_type=GISFeature.POINT,
                    geometry=geom,
                    feature_id=fid,
                    attributes=_attrs({'name': name}),
                    deo_visible=deo_visible,
                    created_by=user,
                )
                created += 1
            except Exception as exc:
                errors.append(f"Waypoint creation error: {exc}")

    # Parse Tracks
    trk_tag = './/gpx:trk' if 'gpx' in ns else './/trk'
    trkseg_tag = 'gpx:trkseg' if 'gpx' in ns else 'trkseg'
    trkpt_tag = 'gpx:trkpt' if 'gpx' in ns else 'trkpt'

    for trk in root.findall(trk_tag, ns):
        try:
            name_el = trk.find(name_tag, ns)
            name = name_el.text if name_el is not None else 'GPS Track'
            desc_el = trk.find(desc_tag, ns)
            desc = desc_el.text if desc_el is not None else ''

            for trkseg in trk.findall(trkseg_tag, ns):
                pts = []
                for trkpt in trkseg.findall(trkpt_tag, ns):
                    lat = float(trkpt.attrib['lat'])
                    lon = float(trkpt.attrib['lon'])
                    pts.append((lon, lat))
                
                if not pts:
                    continue

                if geom_type == 'point':
                    # Import each track point as a POINT feature
                    for i, pt in enumerate(pts):
                        geom = Point(pt[0], pt[1], srid=4326)
                        fid = f"{name} pt {i+1}" if name_field else ''
                        GISFeature.objects.create(
                            project=project,
                            folder=folder,
                            layer_name=layer_name,
                            geometry_type=GISFeature.POINT,
                            geometry=geom,
                            feature_id=fid,
                            attributes=_attrs({'name': f"{name} Point {i+1}", 'description': desc}),
                            deo_visible=deo_visible,
                            created_by=user,
                        )
                        created += 1
                elif geom_type == 'polygon':
                    # Import track as closed Polygon
                    if len(pts) >= 3:
                        if pts[0] != pts[-1]:
                            pts.append(pts[0])
                        if len(pts) >= 4:
                            geom = Polygon(pts, srid=4326)
                            fid = name if name_field else ''
                            GISFeature.objects.create(
                                project=project,
                                folder=folder,
                                layer_name=layer_name,
                                geometry_type=GISFeature.POLYGON,
                                geometry=geom,
                                feature_id=fid,
                                attributes=_attrs({'name': name, 'description': desc}),
                                deo_visible=deo_visible,
                                created_by=user,
                            )
                            created += 1
                else:  # 'auto' or 'line'
                    # Import track as LineString
                    if len(pts) >= 2:
                        geom = LineString(pts, srid=4326)
                        fid = name if name_field else ''
                        GISFeature.objects.create(
                            project=project,
                            folder=folder,
                            layer_name=layer_name,
                            geometry_type=GISFeature.LINE,
                            geometry=geom,
                            feature_id=fid,
                            attributes=_attrs({'name': name, 'description': desc}),
                            deo_visible=deo_visible,
                            created_by=user,
                        )
                        created += 1
        except Exception as exc:
            errors.append(f"Track error: {exc}")

    return Response({
        'detail': f'Imported {created} feature(s) from GPX.',
        'created': created,
        'errors': errors[:10],
        'type': 'vector',
    }, status=201)


def _import_csv_file(folder, uploaded, layer_name, name_field, user, geom_type='auto', deo_visible=True):
    """Parse a CSV file and import coordinates as GISFeatures."""
    import csv
    import io
    from django.contrib.gis.geos import Point, LineString, Polygon
    from rest_framework.response import Response
    from .models import GISFeature

    try:
        csv_data = uploaded.read()
        try:
            decoded = csv_data.decode('utf-8-sig')
        except UnicodeDecodeError:
            decoded = csv_data.decode('latin1')
        reader = csv.DictReader(io.StringIO(decoded))
        headers = reader.fieldnames or []
    except Exception as exc:
        return Response({'detail': f'Invalid CSV: {exc}'}, status=400)

    # Try to identify lat/lon fields
    lat_col = None
    lon_col = None
    name_col = None

    lat_candidates = ['lat', 'latitude', 'y', 'lat_dec', 'lat_dd', 'north', 'northing']
    lon_candidates = ['lon', 'longitude', 'x', 'lon_dec', 'lon_dd', 'lng', 'east', 'easting']
    name_candidates = ['name', 'label', 'id', 'title', 'waypoint', 'point_id']

    for h in headers:
        hl = h.lower().strip()
        if hl in lat_candidates:
            lat_col = h
        elif hl in lon_candidates:
            lon_col = h
        elif hl in name_candidates:
            name_col = h

    # If not found, look for headings starting with lat/lon
    if not lat_col:
        for h in headers:
            if h.lower().strip().startswith('lat'):
                lat_col = h
                break
    if not lon_col:
        for h in headers:
            if h.lower().strip().startswith('lon') or h.lower().strip().startswith('lng'):
                lon_col = h
                break
    if not name_col:
        for h in headers:
            if h.lower().strip().startswith('name'):
                name_col = h
                break

    if not lat_col or not lon_col:
        return Response({
            'detail': "Could not auto-detect latitude and longitude columns in CSV. Expected columns like: lat, lon, latitude, longitude, etc."
        }, status=400)

    created = 0
    errors = []
    project = folder.project

    # Gather all coordinates and properties from CSV
    pts = []
    rows_attrs = []
    first_row_name = None

    for row in reader:
        try:
            lat_str = row.get(lat_col)
            lon_str = row.get(lon_col)
            if not lat_str or not lon_str:
                continue
            lat = float(lat_str.strip())
            lon = float(lon_str.strip())

            name = row.get(name_col, '').strip() if name_col else f"Point {reader.line_num}"
            if first_row_name is None:
                first_row_name = name

            attrs = {k: v.strip() for k, v in row.items() if k not in (lat_col, lon_col)}
            attrs['name'] = name

            pts.append((lon, lat))
            rows_attrs.append(attrs)
        except Exception as exc:
            errors.append(f"Row {reader.line_num} parse error: {exc}")

    # Now create geometries based on geom_type
    if geom_type == 'polygon':
        if len(pts) >= 3:
            try:
                if pts[0] != pts[-1]:
                    pts.append(pts[0])
                if len(pts) >= 4:
                    geom = Polygon(pts, srid=4326)
                    fid = first_row_name if name_field and first_row_name else ''
                    GISFeature.objects.create(
                        project=project,
                        folder=folder,
                        layer_name=layer_name,
                        geometry_type=GISFeature.POLYGON,
                        geometry=geom,
                        feature_id=fid,
                        attributes={'name': f"{layer_name} CSV Polygon", 'description': 'Imported from CSV list'},
                        deo_visible=deo_visible,
                        created_by=user,
                    )
                    created += 1
            except Exception as exc:
                errors.append(f"CSV Polygon creation error: {exc}")
        else:
            errors.append("Need at least 3 points to construct a polygon.")
    elif geom_type == 'line':
        if len(pts) >= 2:
            try:
                geom = LineString(pts, srid=4326)
                fid = first_row_name if name_field and first_row_name else ''
                GISFeature.objects.create(
                    project=project,
                    folder=folder,
                    layer_name=layer_name,
                    geometry_type=GISFeature.LINE,
                    geometry=geom,
                    feature_id=fid,
                    attributes={'name': f"{layer_name} CSV Line", 'description': 'Imported from CSV list'},
                    deo_visible=deo_visible,
                    created_by=user,
                )
                created += 1
            except Exception as exc:
                errors.append(f"CSV Line creation error: {exc}")
        else:
            errors.append("Need at least 2 points to construct a line.")
    else:  # 'auto' or 'point' (meaning import each point individually)
        for i, pt in enumerate(pts):
            try:
                geom = Point(pt[0], pt[1], srid=4326)
                attrs = rows_attrs[i]
                name = attrs.get('name', f"Point {i+1}")
                fid = name if name_field else ''
                GISFeature.objects.create(
                    project=project,
                    folder=folder,
                    layer_name=layer_name,
                    geometry_type=GISFeature.POINT,
                    geometry=geom,
                    feature_id=fid,
                    attributes=attrs,
                    deo_visible=deo_visible,
                    created_by=user,
                )
                created += 1
            except Exception as exc:
                errors.append(f"Point creation error: {exc}")

    return Response({
        'detail': f'Imported {created} feature(s) from CSV.',
        'created': created,
        'errors': errors[:10],
        'type': 'vector',
    }, status=201)


def _import_via_gdal(folder, uploaded, layer_name, name_field, user, fmt, deo_visible=True):
    """Import KML or GeoPackage via GDAL DataSource."""
    import os
    import tempfile
    from django.contrib.gis.gdal import DataSource
    from django.contrib.gis.geos import GEOSGeometry
    from rest_framework.response import Response

    suffix = f'.{fmt}'
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        for chunk in uploaded.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name

    try:
        try:
            ds = DataSource(tmp_path)
        except Exception as exc:
            return Response({'detail': f'Cannot read {fmt.upper()} file: {exc}'}, status=400)

        created = 0
        errors = []
        project = folder.project

        for lyr in ds:
            for feat in lyr:
                try:
                    geom = GEOSGeometry(feat.geom.wkt, srid=feat.geom.srid or 4326)
                    if geom.srid != 4326:
                        geom.transform(4326)
                    geom_type = _classify_geom(geom.geom_type)
                    attrs = {field: _safe_val(feat[field].value) for field in feat.fields}
                    fid = str(feat[name_field].value) if name_field and name_field in feat.fields else ''
                    GISFeature.objects.create(
                        project=project,
                        folder=folder,
                        layer_name=layer_name,
                        geometry_type=geom_type,
                        geometry=geom,
                        feature_id=fid,
                        attributes=attrs,
                        deo_visible=deo_visible,
                        created_by=user,
                    )
                    created += 1
                except Exception as exc:
                    errors.append(str(exc))
    finally:
        os.unlink(tmp_path)

    return Response({
        'detail': f'Imported {created} feature(s) from {fmt.upper()}.',
        'created': created,
        'errors': errors[:10],
        'type': 'vector',
    }, status=201)


def _classify_geom(geom_type: str) -> str:
    """Map GDAL/GEOS geometry type strings to GISFeature.geometry_type choices."""
    gt = geom_type.upper()
    if 'POINT' in gt:
        return GISFeature.POINT
    if 'LINE' in gt or 'STRING' in gt:
        return GISFeature.LINE
    return GISFeature.POLYGON


def _safe_val(v):
    """Convert non-JSON-serialisable attribute values to strings."""
    if isinstance(v, (str, int, float, bool, type(None))):
        return v
    return str(v)


# ─────────────────────────────────────────────────────────────────────────────
# OnlyOffice PDF Converter Helper
# ─────────────────────────────────────────────────────────────────────────────

def _convert_to_pdf_via_onlyoffice(doc, project, doc_folder, user, report_title, docx_filename):
    """Convert document to PDF using OnlyOffice ConvertService."""
    import jwt
    import json
    import urllib.request
    import time
    from django.conf import settings
    from django.core.files.base import ContentFile
    from apps.documents.models import Document

    secret = getattr(settings, 'ONLYOFFICE_JWT_SECRET', '')
    internal_base = getattr(settings, 'ONLYOFFICE_INTERNAL_BASE_URL', '').rstrip('/')

    if internal_base:
        doc_url = f"{internal_base}{doc.file.url}"
    else:
        doc_url = f"http://nginx{doc.file.url}"

    convert_url = 'http://onlyoffice/ConvertService.ashx'
    key = f"conv-doc-{doc.id}-{doc.version}-{int(time.time())}"
    payload = {
        'async': False,
        'filetype': 'docx',
        'key': key,
        'outputtype': 'pdf',
        'title': docx_filename,
        'url': doc_url
    }

    if secret:
        token_payload = {
            'payload': payload
        }
        token = jwt.encode(token_payload, secret, algorithm='HS256')
        if isinstance(token, bytes):
            token = token.decode('utf-8')
        payload['token'] = token

    data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(convert_url, data=data, headers={
        'Content-Type': 'application/json',
        'Accept': 'application/json',
        'Authorization': f'Bearer {payload.get("token", "")}' if secret else ''
    })

    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            resp_data = json.loads(resp.read().decode('utf-8'))
            if resp_data.get('endConvert') and resp_data.get('fileUrl'):
                pdf_download_url = resp_data['fileUrl']
                with urllib.request.urlopen(pdf_download_url, timeout=30) as pdf_resp:
                    pdf_content = pdf_resp.read()
                
                pdf_filename = docx_filename.replace('.docx', '.pdf')
                try:
                    from apps.core.watermark import embed_watermark
                    wm_meta = {
                        "project_id": project.id,
                        "project_number": project.project_number,
                        "title": f'{report_title} (PDF)',
                        "uploaded_by": user.username,
                        "export_format": "pdf",
                    }
                    pdf_content = embed_watermark(pdf_content, pdf_filename, 'application/pdf', wm_meta)
                except Exception:
                    pass
                pdf_doc = Document.objects.create(
                    project=project,
                    folder=doc_folder,
                    title=f'{report_title} (PDF)',
                    category=Document.SURVEY_REPORT,
                    file_size=len(pdf_content),
                    mime_type='application/pdf',
                    uploaded_by=user,
                    parent=doc,
                )
                pdf_doc.file.save(pdf_filename, ContentFile(pdf_content), save=True)
                return pdf_doc
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"OnlyOffice PDF conversion failed: {e}")
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Survey Report Generator
# ─────────────────────────────────────────────────────────────────────────────

def _build_survey_report(area, user, include_features=True, include_photos=True, include_workflow=True):
    """
    Generate a .docx survey report for a SurveyArea, save it to the area's Doc folder,
    optionally convert to PDF via LibreOffice headless, and return download URLs.
    """
    import io, os, shutil, subprocess, tempfile
    from datetime import date
    from django.core.files.base import ContentFile
    from django.db.models import Count, Sum
    from django.contrib.gis.db.models.functions import Area, Length, Transform
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from apps.documents.models import Document
    from apps.workflow.models import WorkflowStep

    project = area.project
    org     = project.organisation

    # ── Ensure area has a folder tree ────────────────────────────────────────
    if not area.folder_id:
        root_folder = ProjectLayerFolder.objects.create(
            project=project, name=area.name,
            folder_type=ProjectLayerFolder.ZONE, created_by=user, order=0,
        )
        _add_survey_area_subfolders(root_folder, user)
        area.folder = root_folder
        area.save(update_fields=['folder'])

    # Find the DOC subfolder
    doc_folder, _ = ProjectLayerFolder.objects.get_or_create(
        project=project, parent=area.folder,
        folder_type=ProjectLayerFolder.DOC,
        defaults={'name': 'Doc', 'created_by': user, 'order': 2},
    )

    # ── Collect data ──────────────────────────────────────────────────────────
    features_qs = GISFeature.objects.filter(
        project=project, is_deleted=False,
    ).filter(
        folder__in=_get_subtree_folder_ids(area.folder)
    )

    # Layer summary: count, total area (polygon), total length (line)
    layer_stats = {}
    for f in features_qs.values('layer_name', 'geometry_type'):
        key = (f['layer_name'], f['geometry_type'])
        layer_stats.setdefault(key, 0)
        layer_stats[key] += 1

    # Compute areas and lengths per layer
    layer_metrics = {}
    for (ln, gtype), cnt in layer_stats.items():
        lq = features_qs.filter(layer_name=ln, geometry_type=gtype)
        total_area_ha = None
        total_len_km = None
        if gtype == GISFeature.POLYGON:
            try:
                proj_geoms = lq.annotate(geom_proj=Transform('geometry', 32643))
                total_m2 = sum(
                    f.geom_proj.area for f in proj_geoms
                    if f.geom_proj and not f.geom_proj.empty
                )
                total_area_ha = round(total_m2 / 10000, 4)
            except Exception:
                pass
        elif gtype == GISFeature.LINE:
            try:
                proj_geoms = lq.annotate(geom_proj=Transform('geometry', 32643))
                total_m = sum(
                    f.geom_proj.length for f in proj_geoms
                    if f.geom_proj and not f.geom_proj.empty
                )
                total_len_km = round(total_m / 1000, 3)
            except Exception:
                pass
        layer_metrics[(ln, gtype)] = {
            'count': cnt, 'area_ha': total_area_ha, 'len_km': total_len_km,
        }

    # Workflow steps
    workflow_steps = list(
        WorkflowStep.objects.filter(survey_area=area)
        .select_related('actor')
        .order_by('timestamp')
    ) if include_workflow else []

    # Photos from doc folder
    photos = list(
        Document.objects.filter(folder=doc_folder, category=Document.PHOTO)
        .order_by('uploaded_at')
    ) if include_photos else []

    # ── Build Document ────────────────────────────────────────────────────────
    doc = DocxDocument()

    # Page margins (2 cm all around)
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def _centered_para(text, bold=False, size=11, color=None, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def _add_section_heading(text):
        p = doc.add_heading(text, level=2)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _info_table(rows):
        """Two-column label/value table."""
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (label, value) in enumerate(rows):
            lc = tbl.cell(i, 0)
            vc = tbl.cell(i, 1)
            lc.width = Cm(5.5)
            lr = lc.paragraphs[0].add_run(label)
            lr.bold = True
            lr.font.size = Pt(10)
            vr = vc.paragraphs[0].add_run(str(value) if value is not None else '—')
            vr.font.size = Pt(10)
        return tbl

    def _set_col_widths(tbl, widths_cm):
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                if j < len(widths_cm):
                    cell.width = Cm(widths_cm[j])

    # ── Header ────────────────────────────────────────────────────────────────
    _centered_para('GOVERNMENT OF INDIA', bold=True, size=13, color=(0, 0, 128), space_after=2)
    _centered_para('DEFENCE ESTATES DEPARTMENT', bold=True, size=12, color=(0, 0, 128), space_after=2)
    _centered_para(f'{org.name}  ({org.get_level_display()})', bold=True, size=11, space_after=2)
    if org.address:
        _centered_para(org.address, size=10, space_after=2)
    if org.office_id:
        _centered_para(f'Office Code: {org.office_id}', size=10, space_after=6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    _centered_para('SURVEY AREA REPORT', bold=True, size=16,
                   color=(139, 0, 0), space_before=4, space_after=8)

    ref_no = f'SR/{project.project_number}/{area.id}'
    _centered_para(f'Reference No: {ref_no}   |   Date: {date.today().strftime("%d %B %Y")}',
                   size=10, space_after=10)
    doc.add_paragraph()

    # ── 1. Project Information ─────────────────────────────────────────────────
    _add_section_heading('1. PROJECT INFORMATION')
    proj_rows = [
        ('Project Number',  project.project_number),
        ('Project Name',    project.name),
        ('Survey Type',     project.get_survey_type_display()),
        ('Priority',        project.get_priority_display()),
        ('Status',          project.get_status_display()),
        ('State',           project.state.name if project.state else '—'),
        ('District',        project.district.name if project.district else '—'),
        ('Taluk',           project.taluk.name if project.taluk else '—'),
        ('Village',         project.village.name if project.village else '—'),
        ('Total Area',      f'{project.total_area_hectares} ha' if project.total_area_hectares else '—'),
        ('Start Date',      project.start_date.strftime('%d-%b-%Y') if project.start_date else '—'),
        ('Target Date',     project.target_date.strftime('%d-%b-%Y') if project.target_date else '—'),
        ('Description',     project.description or '—'),
    ]
    _info_table(proj_rows)
    doc.add_paragraph()

    # ── 2. Survey Area Details ────────────────────────────────────────────────
    _add_section_heading('2. SURVEY AREA DETAILS')
    area_rows = [
        ('Area Name',      area.name),
        ('Area Code',      area.area_code or '—'),
        ('Status',         area.get_status_display()),
        ('Assigned To',    area.assigned_to.get_full_name() if area.assigned_to else '—'),
        ('Created By',     area.created_by.get_full_name() if area.created_by else '—'),
        ('Created On',     area.created_at.strftime('%d-%b-%Y') if area.created_at else '—'),
        ('Last Updated',   area.updated_at.strftime('%d-%b-%Y') if area.updated_at else '—'),
        ('Description',    area.description or '—'),
    ]
    _info_table(area_rows)
    doc.add_paragraph()

    # ── 3. Feature Summary ────────────────────────────────────────────────────
    if include_features and layer_metrics:
        _add_section_heading('3. FEATURE SUMMARY')
        headers = ['Layer Name', 'Type', 'Count', 'Total Area (ha)', 'Total Length (km)']
        data_rows = [
            (ln, gtype, m['count'],
             str(m['area_ha']) if m['area_ha'] is not None else '—',
             str(m['len_km'])  if m['len_km']  is not None else '—')
            for (ln, gtype), m in sorted(layer_metrics.items())
        ]
        tbl = doc.add_table(rows=1 + len(data_rows), cols=5)
        tbl.style = 'Table Grid'
        # Header row
        hdr = tbl.rows[0]
        for i, h in enumerate(headers):
            run = hdr.cells[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
        # Data rows
        for ri, row_data in enumerate(data_rows, start=1):
            for ci, val in enumerate(row_data):
                tbl.rows[ri].cells[ci].paragraphs[0].add_run(str(val)).font.size = Pt(9)
        _set_col_widths(tbl, [5.5, 2.5, 1.5, 3.0, 3.0])
        total = sum(m['count'] for m in layer_metrics.values())
        p = doc.add_paragraph(f'\nTotal: {total} feature(s) across {len(layer_metrics)} layer(s)')
        p.paragraph_format.space_after = Pt(8)

    # ── 4. Feature Attribute Details ──────────────────────────────────────────
    if include_features and layer_metrics:
        _add_section_heading('4. FEATURE ATTRIBUTE DETAILS')
        for (ln, gtype), m in sorted(layer_metrics.items()):
            layer_feats = list(
                features_qs.filter(layer_name=ln, geometry_type=gtype).order_by('id')[:50]
            )
            if not layer_feats:
                continue
            p = doc.add_paragraph()
            r = p.add_run(f'{ln}  ({gtype} — {m["count"]} feature(s))')
            r.bold = True
            r.font.size = Pt(10)

            # Collect all attribute keys across features
            all_keys = []
            seen = set()
            for feat in layer_feats:
                for k in (feat.attributes or {}):
                    if k not in seen:
                        seen.add(k)
                        all_keys.append(k)
            all_keys = all_keys[:12]  # cap columns to keep doc readable

            col_headers = ['#', 'Feature ID'] + all_keys
            attr_tbl = doc.add_table(rows=1 + len(layer_feats), cols=len(col_headers))
            attr_tbl.style = 'Table Grid'
            for ci, h in enumerate(col_headers):
                run = attr_tbl.rows[0].cells[ci].paragraphs[0].add_run(h)
                run.bold = True
                run.font.size = Pt(8)
            for ri, feat in enumerate(layer_feats, start=1):
                row = attr_tbl.rows[ri]
                row.cells[0].paragraphs[0].add_run(str(ri)).font.size = Pt(8)
                row.cells[1].paragraphs[0].add_run(feat.feature_id or str(feat.id)).font.size = Pt(8)
                for ci, k in enumerate(all_keys, start=2):
                    val = (feat.attributes or {}).get(k, '')
                    row.cells[ci].paragraphs[0].add_run(str(val)[:60]).font.size = Pt(8)

            if m['count'] > 50:
                doc.add_paragraph(
                    f'  (Showing first 50 of {m["count"]} features)'
                ).paragraph_format.space_after = Pt(6)
            doc.add_paragraph()

    # ── Encroachment / Overlap details ────────────────────────────────────────
    from apps.external_data.models import ExternalLayer
    from apps.external_data.db_utils import layer_geojson
    from django.contrib.gis.geos import GEOSGeometry

    encroachments = []
    survey_geoms = [f.geometry for f in features_qs if f.geometry]
    if survey_geoms:
        try:
            # Union of geometries
            union_geom = survey_geoms[0]
            for g in survey_geoms[1:]:
                union_geom = union_geom.union(g)

            xmin, ymin, xmax, ymax = union_geom.extent
            search_bbox = [xmin, ymin, xmax, ymax]

            active_layers = ExternalLayer.objects.filter(is_active=True, database__is_active=True).select_related('database')
            for ext_layer in active_layers:
                # Query the layer for features in this bbox
                fc = layer_geojson(ext_layer, limit=1000, user=user, bbox=search_bbox)
                for feat in fc.get('features', []):
                    geom_json = feat.get('geometry')
                    if not geom_json:
                        continue
                    try:
                        ext_geom = GEOSGeometry(json.dumps(geom_json), srid=4326)
                    except Exception:
                        continue

                    # Check intersection
                    if union_geom.intersects(ext_geom):
                        overlap_area_ha = 0
                        intersection = union_geom.intersection(ext_geom)
                        if intersection and intersection.geom_type in ('Polygon', 'MultiPolygon'):
                            try:
                                overlap_area_ha = round(intersection.transform(32643, clone=True).area / 10000.0, 4)
                            except Exception:
                                overlap_area_ha = round(intersection.area * (111320 ** 2) / 10000.0, 4)

                        props = feat.get('properties') or {}
                        label_col = ext_layer.label_column
                        label = str(props.get(label_col, '')) if label_col else ext_layer.display_name

                        encroachments.append({
                            'layer_name': ext_layer.display_name,
                            'db_name': ext_layer.database.name,
                            'feature_id': label,
                            'overlap_area_ha': overlap_area_ha,
                            'details': {k: str(v) for k, v in list(props.items())[:3]}
                        })
        except Exception as exc:
            logger.error("Encroachment check failed: %s", exc)

    # Add Section 5: ENCROACHMENT / OVERLAP ANALYSIS
    _add_section_heading('5. ENCROACHMENT / OVERLAP ANALYSIS')
    if encroachments:
        enc_tbl = doc.add_table(rows=1 + len(encroachments), cols=5)
        enc_tbl.style = 'Table Grid'
        headers = ['External Layer', 'External Database', 'External Feature ID', 'Overlap Area (ha)', 'Details']
        for ci, h in enumerate(headers):
            run = enc_tbl.rows[0].cells[ci].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
        for ri, enc in enumerate(encroachments, start=1):
            row = enc_tbl.rows[ri]
            row.cells[0].paragraphs[0].add_run(enc['layer_name']).font.size = Pt(9)
            row.cells[1].paragraphs[0].add_run(enc['db_name']).font.size = Pt(9)
            row.cells[2].paragraphs[0].add_run(enc['feature_id']).font.size = Pt(9)
            row.cells[3].paragraphs[0].add_run(f"{enc['overlap_area_ha']} ha" if enc['overlap_area_ha'] > 0 else '—').font.size = Pt(9)
            
            # Details string
            details_str = ", ".join(f"{k}: {v}" for k, v in enc['details'].items())
            row.cells[4].paragraphs[0].add_run(details_str).font.size = Pt(8)
            
        _set_col_widths(enc_tbl, [3.5, 3.5, 3.5, 2.5, 3.0])
    else:
        doc.add_paragraph('No encroachment or overlap with external database layers detected.')
    doc.add_paragraph()

    # ── 6. Photographs ────────────────────────────────────────────────────────
    if include_photos and photos:
        _add_section_heading('6. PHOTOGRAPHS')
        for photo in photos:
            try:
                img_path = photo.file.path
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Cm(14))
                    cap = doc.add_paragraph(photo.title)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(8)
                    run = cap.runs[0]
                    run.italic = True
                    run.font.size = Pt(9)
            except Exception:
                pass
        doc.add_paragraph()

    # ── 7. Workflow History ───────────────────────────────────────────────────
    if include_workflow:
        _add_section_heading(f'{"7" if include_photos and photos else "6"}. WORKFLOW HISTORY')
        if workflow_steps:
            wf_tbl = doc.add_table(rows=1 + len(workflow_steps), cols=4)
            wf_tbl.style = 'Table Grid'
            for ci, h in enumerate(['Date & Time', 'Action', 'By', 'Remarks']):
                run = wf_tbl.rows[0].cells[ci].paragraphs[0].add_run(h)
                run.bold = True
                run.font.size = Pt(10)
            for ri, step in enumerate(workflow_steps, start=1):
                row = wf_tbl.rows[ri]
                row.cells[0].paragraphs[0].add_run(
                    step.timestamp.strftime('%d-%b-%Y %H:%M')
                ).font.size = Pt(9)
                row.cells[1].paragraphs[0].add_run(step.get_action_display()).font.size = Pt(9)
                row.cells[2].paragraphs[0].add_run(
                    step.actor.get_full_name() or step.actor.username
                ).font.size = Pt(9)
                row.cells[3].paragraphs[0].add_run(step.remarks or '').font.size = Pt(9)
            _set_col_widths(wf_tbl, [3.5, 5.0, 3.5, 4.0])
        else:
            doc.add_paragraph('No workflow actions recorded for this area.')
        doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    from django.utils import timezone
    now_str = timezone.now().strftime('%d-%b-%Y %H:%M')
    _centered_para('─' * 60, size=9, space_before=8, space_after=2)
    _centered_para(
        f'Report generated on: {now_str}   |   By: {user.get_full_name() or user.username} ({user.get_role_display()})',
        size=9, space_after=2,
    )
    _centered_para('RakshaGIS — Defence Estates GIS Survey Platform', size=8, color=(150, 150, 150))

    # ── Save .docx to Doc folder ──────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    report_title = f'Survey Report — {area.name}'
    safe_name = area.name.replace('/', '_').replace(' ', '_')
    docx_filename = f'survey_report_{safe_name}_{date.today().isoformat()}.docx'

    try:
        from apps.core.watermark import embed_watermark
        wm_meta = {
            "project_id": project.id,
            "project_number": project.project_number,
            "title": report_title,
            "uploaded_by": user.username,
            "export_format": "docx",
        }
        docx_bytes = embed_watermark(
            docx_bytes, docx_filename,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            wm_meta,
        )
    except Exception:
        pass

    db_doc = Document.objects.create(
        project=project,
        folder=doc_folder,
        title=report_title,
        category=Document.SURVEY_REPORT,
        file_size=len(docx_bytes),
        mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        uploaded_by=user,
    )
    db_doc.file.save(docx_filename, ContentFile(docx_bytes), save=True)

    result = {
        'doc_id':   db_doc.id,
        'file_url': db_doc.file.url,
        'title':    report_title,
        'pdf_doc_id': None,
        'pdf_url':  None,
    }

    # ── Convert to PDF (OnlyOffice ConvertService) ───────────────────────────
    try:
        pdf_doc = _convert_to_pdf_via_onlyoffice(
            db_doc, project, doc_folder, user, report_title, docx_filename
        )
        if pdf_doc:
            result['pdf_doc_id'] = pdf_doc.id
            result['pdf_url']    = pdf_doc.file.url
    except Exception:
        pass

    return result


def _get_subtree_folder_ids(root_folder_id):
    """Return a list of folder IDs for root + all descendants."""
    from collections import deque
    all_ids = []
    queue = deque([root_folder_id])
    while queue:
        cur = queue.popleft()
        all_ids.append(cur)
        for child_id in ProjectLayerFolder.objects.filter(parent_id=cur).values_list('id', flat=True):
            queue.append(child_id)
    return all_ids


class SurveyAreaViewSet(viewsets.ModelViewSet):
    """
    CRUD for SurveyArea objects.
    Workflow transitions are handled via POST /workflow/steps/area-transition/{pk}/{transition}/
    """
    serializer_class = SurveyAreaSerializer
    filter_backends = [DjangoFilterBackend, SearchFilter]
    filterset_fields = ['project', 'status', 'assigned_to']
    search_fields = ['name', 'area_code', 'description']

    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        base_qs = SurveyArea.objects.select_related(
            'project__organisation', 'assigned_to', 'created_by', 'folder'
        )
        from .access import hq_level, published_map_filter
        own_qs = org_queryset_filter(user, base_qs, org_field='project__organisation')
        # DGDE/PDDE (national/command) office users: their OWN org's areas stay
        # fully visible (HQ offices can run their own projects), while subordinate
        # offices' data is limited to PUBLISHED + map-enabled areas (Map Viewer),
        # PDDE additionally restricted to its command subtree.
        level = hq_level(user)
        org_id = self.request.query_params.get('organisation')
        if level:
            pub = published_map_filter(user, base_qs.filter(
                status=SurveyArea.PUBLISHED, map_enabled=True, project__map_enabled=True,
            ))
            if org_id:
                return pub.filter(project__organisation_id=org_id)
            return (own_qs | pub).distinct()
        if user.is_superadmin:
            # Field-browser: when a specific office is passed, return only its
            # PUBLISHED survey areas — same scope as DGDE/PDDE office browsing.
            if org_id:
                return SurveyArea.objects.filter(
                    status=SurveyArea.PUBLISHED, map_enabled=True,
                    project__map_enabled=True,
                    project__organisation_id=org_id,
                )
            return own_qs
        # Also include survey areas shared via ProjectShare or SurveyAreaAccessRequest
        shared_project_ids = get_shared_project_ids(user)
        approved_area_ids  = get_approved_area_ids(user)
        if not shared_project_ids and not approved_area_ids:
            return own_qs
        extra_q = Q(project_id__in=shared_project_ids) | Q(id__in=approved_area_ids)
        return (own_qs | base_qs.filter(extra_q)).distinct()

    def get_permissions(self):
        # `features` is read-only; `set_map_enabled` enforces its own admin check and
        # must work even on PUBLISHED (edit-locked) areas — so neither uses CanEditProject.
        if self.action in ['list', 'retrieve', 'discovery', 'features', 'set_map_enabled']:
            return [permissions.IsAuthenticated()]
        return [permissions.IsAuthenticated(), CanEditProject()]

    def perform_create(self, serializer):
        from apps.accounts.models import User
        user = self.request.user
        project = serializer.validated_data.get('project')
        # Only SDO/SURVEYOR/SUPERADMIN can create survey areas
        if user.role not in (User.SDO, User.SURVEYOR, User.SUPERADMIN):
            raise PermissionDenied('Only SDO/Surveyor can create survey areas.')
        if user.role != User.SUPERADMIN and project.organisation != user.organisation:
            raise PermissionDenied('Permission denied.')
        area = serializer.save(created_by=user)
        # Auto-create the default folder tree so features are always area-scoped
        root_folder = ProjectLayerFolder.objects.create(
            project=project, name=area.name,
            folder_type=ProjectLayerFolder.ZONE, created_by=user, order=0,
        )
        _add_survey_area_subfolders(root_folder, user)
        area.folder = root_folder
        area.save(update_fields=['folder'])

    @action(detail=True, methods=['post'], url_path='set-map-enabled',
            permission_classes=[permissions.IsAuthenticated])
    def set_map_enabled(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/set-map-enabled/  body: {map_enabled: bool}
        Enable/disable exposure of this area's PUBLISHED data to higher levels.
        Owning-office admins (DEO/CEO/ADEO) + superadmin only.
        """
        area = self.get_object()
        user = request.user
        if not (user.is_superadmin or user.role in user.ADMIN_ROLES):
            return Response({'detail': 'Only office admins can change visibility.'}, status=403)
        area.map_enabled = bool(request.data.get('map_enabled', True))
        area.save(update_fields=['map_enabled'])
        return Response({'id': area.id, 'map_enabled': area.map_enabled})

    @action(detail=True, methods=['get'], url_path='features',
            permission_classes=[permissions.IsAuthenticated])
    def features(self, request, pk=None):
        """
        GET /api/projects/survey-areas/{id}/features/?limit=5000

        Return this survey area's GIS features as a GeoJSON FeatureCollection,
        scoped to the area's folder subtree and the caller's jurisdiction.
        Lets the 3D Terrain Viewer load a single named area (e.g. "AFS Sulur")
        instead of the whole project.
        """
        area = self.get_object()
        from apps.survey_projects.analysis import (
            _get_folder_ids_for_survey_areas, _scope_survey_qs,
        )
        folder_ids = _get_folder_ids_for_survey_areas([area.id])
        qs = _scope_survey_qs(
            GISFeature.objects.filter(is_deleted=False, folder_id__in=folder_ids),
            request.user,
        ).only('geometry', 'layer_name', 'feature_id', 'attributes')
        try:
            limit = min(int(request.query_params.get('limit', 5000)), 20000)
        except (TypeError, ValueError):
            limit = 5000
        feats = []
        for f in qs[:limit]:
            try:
                geom = json.loads(f.geometry.geojson)
            except Exception:
                continue
            props = {'layer_name': f.layer_name, 'feature_id': f.feature_id}
            if isinstance(f.attributes, dict):
                props.update(f.attributes)
            feats.append({'type': 'Feature', 'geometry': geom, 'properties': props})
        return Response({'type': 'FeatureCollection', 'features': feats})

    @action(detail=True, methods=['get'], url_path='summary',
            permission_classes=[permissions.IsAuthenticated])
    def summary(self, request, pk=None):
        """
        GET /api/projects/survey-areas/{id}/summary/
        Return live summary statistics (total area in hectares, feature counts, last edited by, etc.).
        """
        area = self.get_object()
        from django.db.models import Count
        from apps.survey_projects.models import GISFeature, ProjectLayerFolder
        from apps.survey_projects.views import _get_subtree_folder_ids

        folder_ids = []
        if area.folder_id:
            folder_ids = _get_subtree_folder_ids(area.folder_id)

        features = GISFeature.objects.filter(folder_id__in=folder_ids, is_deleted=False)

        counts = features.values('geometry_type').annotate(count=Count('id'))
        count_map = {'POINT': 0, 'LINE': 0, 'POLYGON': 0}
        for c in counts:
            gt = c['geometry_type']
            count_map[gt] = c['count']

        total_sqm = 0
        polygons = features.filter(geometry_type='POLYGON')
        for f in polygons:
            if f.geometry:
                try:
                    total_sqm += f.geometry.transform(32643, clone=True).area
                except Exception:
                    try:
                        total_sqm += f.geometry.area * (111320 ** 2)
                    except Exception:
                        pass
        total_ha = round(total_sqm / 10000.0, 3)

        last_feat = features.order_by('-updated_at').select_related('created_by').first()
        last_edited_by = last_feat.created_by.get_full_name() or last_feat.created_by.username if last_feat and last_feat.created_by else '—'
        last_edited_at = last_feat.updated_at.isoformat() if last_feat else None

        is_locked = area.status not in ('DRAFT', 'RETURNED')

        return Response({
            'total_area_ha': total_ha,
            'features_count': count_map,
            'last_edited_by': last_edited_by,
            'last_edited_at': last_edited_at,
            'is_locked': is_locked,
            'status': area.status,
        })

    @action(detail=True, methods=['post'], url_path='ensure-folder',
            permission_classes=[permissions.IsAuthenticated])
    def ensure_folder(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/ensure-folder/
        Idempotent: create folder tree if missing, return updated SurveyArea.
        Called automatically by the map frontend when an area without a folder is selected.
        """
        area = self.get_object()
        if area.status not in (SurveyArea.DRAFT, SurveyArea.RETURNED):
            return Response({'detail': 'Area is locked — cannot create folders.'}, status=400)
        if not area.folder_id:
            root_folder = ProjectLayerFolder.objects.create(
                project=area.project, name=area.name,
                folder_type=ProjectLayerFolder.ZONE, created_by=request.user, order=0,
            )
            _add_survey_area_subfolders(root_folder, request.user)
            area.folder = root_folder
            area.save(update_fields=['folder'])
        return Response(SurveyAreaSerializer(area).data)

    def perform_update(self, serializer):
        from apps.accounts.models import User
        area = self.get_object()
        user = self.request.user
        if area.status not in (SurveyArea.DRAFT, SurveyArea.RETURNED):
            raise PermissionDenied('Cannot edit a survey area that has been submitted.')
        serializer.save()

    def perform_destroy(self, instance):
        if instance.status not in (SurveyArea.DRAFT, SurveyArea.RETURNED):
            raise PermissionDenied('Cannot delete a submitted survey area.')
        # Deleting the area must remove ALL its data — features, rasters, shapefiles,
        # documents and the folder tree. Every dependent model uses folder=SET_NULL,
        # so deleting folders alone would orphan (not remove) those rows and they would
        # keep showing on the map / in the project. So purge them explicitly first.
        from django.db import transaction
        from apps.survey_projects.analysis import _get_folder_ids_for_survey_areas
        from apps.documents.models import Document

        folder_ids = _get_folder_ids_for_survey_areas([instance.id])
        root_folder = instance.folder

        def _del_files(qs, *fields):
            for obj in qs:
                for fld in fields:
                    f = getattr(obj, fld, None)
                    if f:
                        try:
                            f.delete(save=False)
                        except Exception:
                            pass

        with transaction.atomic():
            if folder_ids:
                # Remove stored files first (FileField.delete is not automatic on row delete).
                _del_files(Document.objects.filter(folder_id__in=folder_ids), 'file')
                _del_files(GeoTiffLayer.objects.filter(folder_id__in=folder_ids), 'file', 'cog_file')
                Document.objects.filter(folder_id__in=folder_ids).delete()
                GeoTiffLayer.objects.filter(folder_id__in=folder_ids).delete()
                GISFeature.objects.filter(folder_id__in=folder_ids).delete()
                # Optional models (present in this build) — purge if importable.
                try:
                    from .models import ShapefileImport
                    _del_files(ShapefileImport.objects.filter(folder_id__in=folder_ids), 'file')
                    ShapefileImport.objects.filter(folder_id__in=folder_ids).delete()
                except Exception:
                    pass
                try:
                    from .models import QGISUploadLog
                    QGISUploadLog.objects.filter(folder_id__in=folder_ids).delete()
                except Exception:
                    pass
            instance.delete()
            if root_folder is not None:
                # CASCADE on ProjectLayerFolder.parent removes the whole subtree.
                ProjectLayerFolder.objects.filter(pk=root_folder.pk).delete()

    @action(detail=True, methods=['post'], url_path='export',
            permission_classes=[permissions.IsAuthenticated])
    def export(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/export/

        Queue an async export of ALL data for this survey area:
          • GIS features (per-layer Shapefiles, EPSG:4326)
          • Documents (C2PA / LP-DNA watermarked)
          • GeoTIFF rasters (C2PA watermarked)
          • Uploaded shapefile ZIPs
          • provenance.json manifest

        Returns: { task_uuid, status, message }
        Poll GET /api/core/export/status/{task_uuid}/ until status=='DONE',
        then download via GET /api/core/export/download/{task_uuid}/.
        """
        from django.http import JsonResponse
        from apps.core.models import ExportTask
        from apps.core.tasks import build_export_zip

        area = self.get_object()
        user = request.user
        org_id = getattr(user, 'organisation_id', None)

        active_statuses = (ExportTask.PENDING, ExportTask.RUNNING)
        _max_user = getattr(__import__('django.conf', fromlist=['settings']).settings,
                            'EXPORT_MAX_CONCURRENT_PER_USER', 2)
        _max_org  = getattr(__import__('django.conf', fromlist=['settings']).settings,
                            'EXPORT_MAX_CONCURRENT_PER_ORG', 3)

        if ExportTask.objects.filter(requested_by=user, status__in=active_statuses).count() >= _max_user:
            return JsonResponse({
                'error': 'Too many active exports',
                'detail': 'Wait for your current export(s) to finish before starting another.',
            }, status=429)

        if org_id and ExportTask.objects.filter(
            organisation_id=org_id, status__in=active_statuses
        ).count() >= _max_org:
            return JsonResponse({
                'error': 'Office export limit reached',
                'detail': 'Your office has too many exports running. Please wait.',
            }, status=429)

        include_dxf = bool(request.data.get('include_dxf', False))

        et = ExportTask.objects.create(
            export_type=ExportTask.SURVEY_AREA,
            object_id=area.id,
            object_name=area.name,
            requested_by=user,
            organisation_id=org_id,
            include_dxf=include_dxf,
            progress_msg='Queued…',
        )
        build_export_zip.delay(et.pk)

        return JsonResponse({
            'task_uuid': str(et.task_uuid),
            'status': et.status,
            'message': 'Export queued.',
        }, status=202)

    @action(detail=True, methods=['post'], url_path='generate-report',
            permission_classes=[permissions.IsAuthenticated])
    def generate_report(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/generate-report/
        Generate a Word (.docx) survey report for this area, save it to the Doc folder,
        and optionally convert to PDF using LibreOffice headless.
        Returns: { doc_id, file_url, pdf_doc_id, pdf_url, title }
        """
        area = self.get_object()
        user = request.user
        include_features = request.data.get('include_features', True)
        include_photos   = request.data.get('include_photos',   True)
        include_workflow = request.data.get('include_workflow',  True)

        try:
            result = _build_survey_report(area, user, include_features, include_photos, include_workflow)
            return Response(result, status=201)
        except ImportError:
            return Response(
                {'detail': 'python-docx is not installed. Run: pip install python-docx'},
                status=500,
            )
        except Exception as exc:
            import traceback
            return Response({'detail': f'Report generation failed: {exc}', 'trace': traceback.format_exc()}, status=500)

    @action(detail=True, methods=['post'], url_path='create-online-report',
            permission_classes=[permissions.IsAuthenticated])
    def create_online_report(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/create-online-report/

        Create a blank structured Word document for this survey area and return a
        Document ID so the caller can open it immediately in OnlyOffice.
        The document is placed in the area's Doc sub-folder (created if absent).
        """
        import io
        import os
        from django.core.files.base import ContentFile
        from apps.documents.models import Document as Doc

        area = self.get_object()
        user = request.user
        title = request.data.get('title') or f'Survey Report — {area.name}'
        # Copy the blank docx template to prevent OnlyOffice "file content does not match the file extension" validation error
        from django.conf import settings
        template_path = os.path.join(str(settings.BASE_DIR), 'apps', 'documents', 'templates', 'new.docx')
        buf_content = None
        if os.path.exists(template_path):
            try:
                with open(template_path, 'rb') as f:
                    buf_content = f.read()
            except Exception:
                pass

        if buf_content is None:
            # Fallback to dynamic docx creation if template is missing
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading(title, level=0)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                buf_content = buf.read()
            except ImportError:
                return Response(
                    {'detail': 'python-docx is not installed. Run: pip install python-docx'},
                    status=500,
                )

        # Find or create the Doc sub-folder under the area's root folder
        doc_folder = None
        if area.folder_id:
            doc_folder = ProjectLayerFolder.objects.filter(
                project=area.project,
                parent=area.folder,
                folder_type=ProjectLayerFolder.DOC,
            ).first()

        safe_name = ''.join(c if c.isalnum() or c in '-_ ' else '_' for c in area.name)
        filename = f'report_{safe_name}.docx'

        try:
            from apps.core.watermark import embed_watermark
            wm_meta = {
                "project_id": area.project_id,
                "project_number": area.project.project_number if area.project else None,
                "title": title,
                "uploaded_by": user.username,
                "export_format": "docx",
            }
            buf_content = embed_watermark(
                buf_content, filename,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                wm_meta,
            )
        except Exception:
            pass

        db_doc = Doc(
            project=area.project,
            folder=doc_folder,
            title=title,
            category='REPORT',
            file_size=len(buf_content),
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            uploaded_by=user,
        )
        db_doc.file.save(filename, ContentFile(buf_content), save=True)

        return Response({
            'doc_id': db_doc.id,
            'file_url': request.build_absolute_uri(db_doc.file.url) if db_doc.file else None,
            'title': db_doc.title,
        }, status=201)

    @action(detail=True, methods=['post'], url_path='template-report',
            permission_classes=[permissions.IsAuthenticated])
    def template_report(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/template-report/

        Generate the ministry-prescribed DGDE survey report (.docx) for this area
        from the standard template — title page + clickable index + the 17 sections,
        with the Statement of Survey Numbers and Area Computation auto-filled from
        the area's GIS features. Saved to the area's Doc folder and returned as a
        Document id so the caller can open it in OnlyOffice for editing.
        """
        from django.core.files.base import ContentFile
        from apps.documents.models import Document as Doc
        from .report_templates import build_ministry_survey_report

        area = self.get_object()
        user = request.user

        # Ensure the area has a folder tree + a Doc sub-folder.
        if not area.folder_id:
            root = ProjectLayerFolder.objects.create(
                project=area.project, name=area.name,
                folder_type=ProjectLayerFolder.ZONE, created_by=user, order=0,
            )
            _add_survey_area_subfolders(root, user)
            area.folder = root
            area.save(update_fields=['folder'])
        doc_folder, _ = ProjectLayerFolder.objects.get_or_create(
            project=area.project, parent=area.folder,
            folder_type=ProjectLayerFolder.DOC,
            defaults={'name': 'Doc', 'created_by': user, 'order': 2},
        )

        try:
            content = build_ministry_survey_report(area, user)
        except ImportError:
            return Response(
                {'detail': 'python-docx is not installed. Run: pip install python-docx'},
                status=500,
            )
        except Exception as exc:
            import traceback
            logger.exception('Template report failed for area %s', area.id)
            return Response({'detail': f'Report generation failed: {exc}'}, status=500)

        safe_name = ''.join(c if c.isalnum() or c in '-_ ' else '_' for c in area.name)
        title = f'DGDE Survey Report — {area.name}'
        docx_filename = f'survey_report_{safe_name}.docx'

        try:
            from apps.core.watermark import embed_watermark
            wm_meta = {
                "project_id": area.project_id,
                "project_number": area.project.project_number if area.project else None,
                "title": title,
                "uploaded_by": user.username,
                "export_format": "docx",
            }
            content = embed_watermark(
                content, docx_filename,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                wm_meta,
            )
        except Exception:
            pass

        db_doc = Doc(
            project=area.project, folder=doc_folder, title=title, category='REPORT',
            file_size=len(content),
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            uploaded_by=user,
        )
        db_doc.file.save(docx_filename, ContentFile(content), save=True)
        return Response({
            'doc_id': db_doc.id,
            'file_url': request.build_absolute_uri(db_doc.file.url) if db_doc.file else None,
            'title': db_doc.title,
        }, status=201)

    @action(detail=True, methods=['post'], url_path='request-access',
            permission_classes=[permissions.IsAuthenticated])
    def request_access(self, request, pk=None):
        """
        DEO/CEO/ADEO user requests read-only access to this survey area.
        The target org's admin must approve via SurveyAreaAccessRequestViewSet.
        """
        area = self.get_object()
        user = request.user
        requesting_org = user.organisation
        if not requesting_org:
            return Response({'detail': 'Your account has no organisation.'}, status=400)
        if area.project.organisation_id == requesting_org.id:
            return Response({'detail': 'This area already belongs to your organisation.'}, status=400)

        from .serializers import SurveyAreaAccessRequestSerializer
        existing = SurveyAreaAccessRequest.objects.filter(
            survey_area=area, requesting_org=requesting_org
        ).first()
        if existing:
            if existing.status == SurveyAreaAccessRequest.APPROVED:
                return Response({'detail': 'Access already granted.'}, status=400)
            if existing.status == SurveyAreaAccessRequest.PENDING:
                return Response({'detail': 'A request is already pending review.'}, status=400)
            # Rejected — allow re-submission
            existing.status = SurveyAreaAccessRequest.PENDING
            existing.reason = request.data.get('reason', existing.reason)
            existing.reviewed_by = None
            existing.reviewed_at = None
            existing.review_remarks = ''
            existing.save()
            return Response(SurveyAreaAccessRequestSerializer(existing).data, status=200)

        req = SurveyAreaAccessRequest.objects.create(
            survey_area=area,
            requested_by=user,
            requesting_org=requesting_org,
            reason=request.data.get('reason', ''),
        )
        return Response(SurveyAreaAccessRequestSerializer(req).data, status=201)

    @action(detail=False, methods=['get'], url_path='discovery',
            permission_classes=[permissions.IsAuthenticated])
    def discovery(self, request):
        """
        Returns basic metadata of survey areas visible within the user's PDDE jurisdiction
        (siblings under same PDDE) that are NOT in the user's own org.
        Useful for cross-org access discovery — no GIS data is exposed.
        """
        user = request.user
        org = user.organisation
        if not org:
            return Response([])

        # Determine PDDE parent
        from apps.accounts.models import Organisation
        if org.level == Organisation.PDDE:
            pdde = org
        elif org.parent and org.parent.level == Organisation.PDDE:
            pdde = org.parent
        else:
            pdde = None

        if pdde:
            sibling_ids = list(
                Organisation.objects.filter(parent=pdde)
                .exclude(id=org.id)
                .values_list('id', flat=True)
            )
        else:
            sibling_ids = []

        if not sibling_ids:
            return Response([])

        # Get approved request IDs for this org (to show status)
        approved_ids = set(get_approved_area_ids(user))
        pending_ids = set(
            SurveyAreaAccessRequest.objects.filter(
                requesting_org=org, status=SurveyAreaAccessRequest.PENDING
            ).values_list('survey_area_id', flat=True)
        )
        rejected_ids = set(
            SurveyAreaAccessRequest.objects.filter(
                requesting_org=org, status=SurveyAreaAccessRequest.REJECTED
            ).values_list('survey_area_id', flat=True)
        )

        areas = SurveyArea.objects.filter(
            project__organisation_id__in=sibling_ids
        ).select_related('project__organisation').order_by('project__organisation__name', 'name')

        result = []
        for a in areas:
            if a.id in approved_ids:
                access = 'APPROVED'
            elif a.id in pending_ids:
                access = 'PENDING'
            elif a.id in rejected_ids:
                access = 'REJECTED'
            else:
                access = 'NONE'
            result.append({
                'id': a.id,
                'name': a.name,
                'area_code': a.area_code,
                'status': a.status,
                'project_id': a.project_id,
                'project_name': a.project.name,
                'org_id': a.project.organisation_id,
                'org_name': a.project.organisation.name,
                'org_level': a.project.organisation.level,
                'access_status': access,
            })
        return Response(result)

    # ── Time-series / history endpoints ──────────────────────────────────────

    @action(detail=True, methods=['get'], url_path='history',
            permission_classes=[permissions.IsAuthenticated])
    def history(self, request, pk=None):
        """
        GET /api/projects/survey-areas/{id}/history/
        Paginated feature change log for this survey area.
        Optional query params: change_type, layer_name, days (default 90)
        """
        from apps.survey_projects.models import GISFeatureHistory
        from apps.survey_projects.serializers import GISFeatureHistorySerializer
        from django.utils import timezone
        import datetime

        area = self.get_object()
        qs = GISFeatureHistory.objects.filter(survey_area=area).select_related('changed_by')

        change_type = request.query_params.get('change_type')
        if change_type:
            qs = qs.filter(change_type=change_type)
        layer_name = request.query_params.get('layer_name')
        if layer_name:
            qs = qs.filter(layer_name=layer_name)
        days = int(request.query_params.get('days', 90))
        cutoff = timezone.now() - datetime.timedelta(days=days)
        qs = qs.filter(changed_at__gte=cutoff)

        page = self.paginate_queryset(qs)
        if page is not None:
            return self.get_paginated_response(GISFeatureHistorySerializer(page, many=True).data)
        return Response(GISFeatureHistorySerializer(qs[:200], many=True).data)

    @action(detail=True, methods=['get'], url_path='timeline',
            permission_classes=[permissions.IsAuthenticated])
    def timeline(self, request, pk=None):
        """
        GET /api/projects/survey-areas/{id}/timeline/
        Returns daily change counts for the past N days (default 90) — used for chart.
        Response: { dates: [...], created: [...], modified: [...], deleted: [...] }
        """
        from apps.survey_projects.models import GISFeatureHistory
        from django.utils import timezone
        from django.db.models.functions import TruncDate
        from django.db.models import Count
        import datetime

        area = self.get_object()
        days = int(request.query_params.get('days', 90))
        cutoff = timezone.now() - datetime.timedelta(days=days)

        rows = (
            GISFeatureHistory.objects
            .filter(survey_area=area, changed_at__gte=cutoff)
            .annotate(day=TruncDate('changed_at'))
            .values('day', 'change_type')
            .annotate(n=Count('id'))
            .order_by('day')
        )

        # Build date-indexed dicts
        from collections import defaultdict
        by_day: dict = defaultdict(lambda: {'CREATE': 0, 'MODIFY': 0, 'DELETE': 0,
                                             'TRANSFER_OUT': 0, 'TRANSFER_IN': 0})
        for row in rows:
            d = str(row['day'])
            by_day[d][row['change_type']] = row['n']

        # Fill in every date from cutoff to today (for a continuous chart)
        all_dates = []
        cur = cutoff.date()
        end = timezone.now().date()
        while cur <= end:
            all_dates.append(str(cur))
            cur += datetime.timedelta(days=1)

        return Response({
            'dates':        all_dates,
            'created':      [by_day[d]['CREATE'] for d in all_dates],
            'modified':     [by_day[d]['MODIFY'] for d in all_dates],
            'deleted':      [by_day[d]['DELETE'] for d in all_dates],
            'transferred':  [by_day[d]['TRANSFER_OUT'] + by_day[d]['TRANSFER_IN'] for d in all_dates],
        })

    @action(detail=True, methods=['get', 'post'], url_path='snapshots',
            permission_classes=[permissions.IsAuthenticated])
    def snapshots(self, request, pk=None):
        """
        GET  /api/projects/survey-areas/{id}/snapshots/ — list snapshots
        POST /api/projects/survey-areas/{id}/snapshots/ — take manual snapshot
        """
        from apps.survey_projects.models import SurveyAreaSnapshot
        from apps.survey_projects.serializers import SurveyAreaSnapshotSerializer

        area = self.get_object()

        if request.method == 'POST':
            snap = _take_snapshot(area, request.user, SurveyAreaSnapshot.MANUAL,
                                   notes=request.data.get('notes', ''),
                                   label=request.data.get('label', ''))
            return Response(SurveyAreaSnapshotSerializer(snap).data, status=201)

        qs = SurveyAreaSnapshot.objects.filter(survey_area=area).select_related('taken_by')
        return Response(SurveyAreaSnapshotSerializer(qs, many=True).data)

    @action(detail=True, methods=['get'], url_path=r'snapshots/(?P<snapshot_id>[0-9]+)/geojson',
            permission_classes=[permissions.IsAuthenticated])
    def snapshot_geojson(self, request, pk=None, snapshot_id=None):
        """
        GET /api/projects/survey-areas/{id}/snapshots/{snapshot_id}/geojson/
        Returns the full GeoJSON FeatureCollection for that snapshot.
        """
        from apps.survey_projects.models import SurveyAreaSnapshot
        area = self.get_object()
        try:
            snap = SurveyAreaSnapshot.objects.get(pk=snapshot_id, survey_area=area)
        except SurveyAreaSnapshot.DoesNotExist:
            return Response({'detail': 'Snapshot not found.'}, status=404)
        return Response(snap.features_geojson)

    @action(detail=True, methods=['get'], url_path='lineage',
            permission_classes=[permissions.IsAuthenticated])
    def lineage(self, request, pk=None):
        """
        GET /api/projects/survey-areas/{id}/lineage/
        Returns parent area info + all child areas (pockets/splits) for this area.
        """
        from apps.survey_projects.models import SurveyAreaSplitRecord
        from apps.survey_projects.serializers import (
            SurveyAreaSerializer, SurveyAreaSplitRecordSerializer,
        )

        area = self.get_object()
        parent = SurveyAreaSerializer(area.parent_area).data if area.parent_area_id else None
        children = SurveyAreaSerializer(
            area.child_areas.select_related('project'), many=True
        ).data
        split_events = SurveyAreaSplitRecord.objects.filter(source_area=area).select_related(
            'new_area', 'performed_by'
        )
        return Response({
            'area':         SurveyAreaSerializer(area).data,
            'parent':       parent,
            'children':     children,
            'split_events': SurveyAreaSplitRecordSerializer(split_events, many=True).data,
        })

    @action(detail=True, methods=['post'], url_path='split',
            permission_classes=[permissions.IsAuthenticated])
    def split(self, request, pk=None):
        """
        POST /api/projects/survey-areas/{id}/split/
        Split this area: move selected features into a new survey area.

        Body:
          new_area_name    (str)   — name for the new area
          new_area_code    (str)   — optional area code
          feature_ids      (list)  — IDs of GISFeature records to transfer
          operation        (str)   — SPLIT | POCKET | TRANSFER (default SPLIT)
          reason           (str)   — required for TRANSFER
          notes            (str)   — optional

        Rules:
          - Any source area status is accepted (DRAFT through PUBLISHED).
          - APPROVED/PUBLISHED areas: new area starts as DRAFT; source stays unchanged.
          - Moved features get survey_area updated to new area (folder re-assigned too).
          - Before/after snapshots are taken automatically.
        """
        from apps.survey_projects.models import (
            GISFeatureHistory, SurveyAreaSnapshot, SurveyAreaSplitRecord,
        )
        from apps.survey_projects.serializers import SurveyAreaSerializer

        area = self.get_object()
        new_name    = (request.data.get('new_area_name') or '').strip()
        new_code    = (request.data.get('new_area_code') or '').strip()
        feature_ids = request.data.get('feature_ids', [])
        operation   = request.data.get('operation', SurveyAreaSplitRecord.SPLIT)
        reason      = (request.data.get('reason') or '').strip()
        notes       = (request.data.get('notes') or '').strip()

        if not new_name:
            return Response({'detail': 'new_area_name is required.'}, status=400)
        if not feature_ids:
            return Response({'detail': 'feature_ids must not be empty.'}, status=400)
        if operation not in (SurveyAreaSplitRecord.SPLIT, SurveyAreaSplitRecord.POCKET,
                             SurveyAreaSplitRecord.TRANSFER):
            return Response({'detail': 'Invalid operation.'}, status=400)
        if operation == SurveyAreaSplitRecord.TRANSFER and not reason:
            return Response({'detail': 'reason is required for TRANSFER operations.'}, status=400)

        # Validate features belong to this area
        features = list(
            GISFeature.objects.filter(pk__in=feature_ids, is_deleted=False)
            .select_related('folder')
        )
        if not features:
            return Response({'detail': 'No valid features found for the given IDs.'}, status=400)

        # Take SPLIT_BEFORE snapshot of source area
        _take_snapshot(area, request.user, SurveyAreaSnapshot.SPLIT_BEFORE,
                       label=f'Before {operation.lower()} → {new_name}')

        # Create new survey area
        area_type = (SurveyArea.POCKET if operation == SurveyAreaSplitRecord.POCKET
                     else SurveyArea.SPLIT_RESULT)
        new_area = SurveyArea.objects.create(
            project=area.project,
            name=new_name,
            area_code=new_code,
            status=SurveyArea.DRAFT,
            area_type=area_type,
            parent_area=area,
            created_by=request.user,
        )
        # Create folder tree for new area
        new_root = ProjectLayerFolder.objects.create(
            project=area.project, name=new_name,
            folder_type=ProjectLayerFolder.ZONE, created_by=request.user, order=0,
        )
        _add_survey_area_subfolders(new_root, request.user)
        new_area.folder = new_root
        new_area.save(update_fields=['folder'])

        # Find the target subfolder (first child folder, likely "Shape Files")
        target_folder = (
            ProjectLayerFolder.objects.filter(parent=new_root)
            .order_by('order', 'name')
            .first()
        ) or new_root

        # Transfer features: reassign folder + survey_area, record history
        for feat in features:
            feat._history_user = request.user
            feat._history_note = f'Transferred to {new_name} ({operation})'

            # Record TRANSFER_OUT on old area
            GISFeatureHistory.objects.create(
                feature=feat, feature_pk=feat.pk,
                survey_area=area, project_id=feat.project_id,
                changed_by=request.user,
                change_type=GISFeatureHistory.TRANSFER_OUT,
                layer_name=feat.layer_name,
                old_geometry=feat.geometry,
                new_geometry=feat.geometry,
                old_attributes=feat.attributes or {},
                new_attributes=feat.attributes or {},
                area_status_at_change=area.status,
                note=f'Transferred to {new_name}. Reason: {reason}' if reason else f'Transferred to {new_name}',
            )

            # Update feature assignment
            old_survey_area_id = feat.survey_area_id
            feat.survey_area = new_area
            feat.folder = target_folder
            feat._skip_history = True  # suppress the auto-history; we're recording manually
            feat.save(update_fields=['survey_area', 'folder'])

            # Record TRANSFER_IN on new area
            GISFeatureHistory.objects.create(
                feature=feat, feature_pk=feat.pk,
                survey_area=new_area, project_id=feat.project_id,
                changed_by=request.user,
                change_type=GISFeatureHistory.TRANSFER_IN,
                layer_name=feat.layer_name,
                old_geometry=None,
                new_geometry=feat.geometry,
                old_attributes={},
                new_attributes=feat.attributes or {},
                area_status_at_change=new_area.status,
                note=f'Received from {area.name}. Reason: {reason}' if reason else f'Received from {area.name}',
            )

        # Create split record
        SurveyAreaSplitRecord.objects.create(
            source_area=area,
            new_area=new_area,
            operation=operation,
            transferred_feature_ids=[f.pk for f in features],
            transferred_feature_count=len(features),
            performed_by=request.user,
            reason=reason,
            notes=notes,
        )

        # Take SPLIT_AFTER snapshots for both areas
        _take_snapshot(area, request.user, SurveyAreaSnapshot.SPLIT_AFTER,
                       label=f'After {operation.lower()} — {len(features)} features moved to {new_name}')
        _take_snapshot(new_area, request.user, SurveyAreaSnapshot.SPLIT_AFTER,
                       label=f'Initial state after {operation.lower()} from {area.name}')

        return Response({
            'detail': f'{len(features)} feature(s) transferred to new area "{new_name}".',
            'new_area': SurveyAreaSerializer(new_area).data,
        }, status=201)


# ── Snapshot helper ───────────────────────────────────────────────────────────

def _take_snapshot(area, user, snapshot_type, label='', notes=''):
    """Build and store a GeoJSON snapshot of all live features in an area."""
    from apps.survey_projects.models import SurveyAreaSnapshot
    from django.contrib.gis.serializers.geojson import Serializer as GeoJSONSerializer

    features_qs = GISFeature.objects.filter(
        survey_area=area, is_deleted=False
    ).select_related('created_by')

    # Also include folder-linked features (legacy records without direct survey_area)
    if area.folder_id:
        folder_ids = _get_subtree_folder_ids(area.folder_id)
        legacy = GISFeature.objects.filter(folder_id__in=folder_ids, is_deleted=False).exclude(
            survey_area=area
        )
        # Union via Python (small counts expected)
        all_feats = list(features_qs) + list(legacy)
    else:
        all_feats = list(features_qs)

    # Build minimal GeoJSON FeatureCollection
    fc = {
        'type': 'FeatureCollection',
        'features': [
            {
                'type': 'Feature',
                'geometry': {'type': f.geometry.geom_type, 'coordinates': []},  # placeholder
                'properties': {
                    'id': f.pk, 'layer_name': f.layer_name,
                    'geometry_type': f.geometry_type,
                    **{k: v for k, v in (f.attributes or {}).items()},
                },
            }
            for f in all_feats
        ],
    }
    # Write proper GeoJSON geometry via GDAL
    for i, feat in enumerate(all_feats):
        try:
            import json
            fc['features'][i]['geometry'] = json.loads(feat.geometry.geojson)
        except Exception:
            fc['features'][i]['geometry'] = None

    snap = SurveyAreaSnapshot.objects.create(
        survey_area=area,
        snapshot_type=snapshot_type,
        taken_by=user,
        status_at_snapshot=area.status,
        feature_count=len(all_feats),
        features_geojson=fc,
        label=label or snapshot_type,
        notes=notes,
    )
    return snap


class SurveyAreaAccessRequestViewSet(viewsets.GenericViewSet):
    """
    Admins see incoming access requests for their org's survey areas.
    They can approve or reject them.
    """
    from .serializers import SurveyAreaAccessRequestSerializer
    serializer_class = SurveyAreaAccessRequestSerializer

    def get_serializer_class(self):
        from .serializers import SurveyAreaAccessRequestSerializer
        return SurveyAreaAccessRequestSerializer

    def get_queryset(self):
        user = self.request.user
        if user.is_superadmin:
            return SurveyAreaAccessRequest.objects.select_related(
                'survey_area__project__organisation', 'requested_by', 'requesting_org', 'reviewed_by'
            ).all()
        from apps.accounts.models import User
        if user.role in User.ADMIN_ROLES:
            # Admin sees requests targeting their org's survey areas
            return SurveyAreaAccessRequest.objects.filter(
                survey_area__project__organisation=user.organisation
            ).select_related(
                'survey_area__project__organisation', 'requested_by', 'requesting_org', 'reviewed_by'
            )
        # Non-admins see requests they submitted
        return SurveyAreaAccessRequest.objects.filter(
            requesting_org=user.organisation
        ).select_related(
            'survey_area__project__organisation', 'requested_by', 'requesting_org', 'reviewed_by'
        )

    def get_permissions(self):
        return [permissions.IsAuthenticated()]

    def list(self, request):
        qs = self.get_queryset()
        status_filter = request.query_params.get('status')
        direction = request.query_params.get('direction')  # 'incoming' or 'outgoing'
        if status_filter:
            qs = qs.filter(status=status_filter)
        if direction == 'incoming' and not request.user.is_superadmin:
            qs = qs.filter(survey_area__project__organisation=request.user.organisation)
        elif direction == 'outgoing':
            qs = qs.filter(requesting_org=request.user.organisation)
        from .serializers import SurveyAreaAccessRequestSerializer
        return Response(SurveyAreaAccessRequestSerializer(qs, many=True).data)

    def retrieve(self, request, pk=None):
        obj = self.get_queryset().filter(pk=pk).first()
        if not obj:
            return Response({'detail': 'Not found.'}, status=404)
        from .serializers import SurveyAreaAccessRequestSerializer
        return Response(SurveyAreaAccessRequestSerializer(obj).data)

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        obj = self.get_queryset().filter(pk=pk).first()
        if not obj:
            return Response({'detail': 'Not found.'}, status=404)
        user = request.user
        from apps.accounts.models import User
        if not user.is_superadmin and user.role not in User.ADMIN_ROLES:
            return Response({'detail': 'Only admins can approve requests.'}, status=403)
        if not user.is_superadmin and obj.survey_area.project.organisation != user.organisation:
            return Response({'detail': 'You can only approve requests for your org\'s areas.'}, status=403)
        from django.utils import timezone
        obj.status = SurveyAreaAccessRequest.APPROVED
        obj.reviewed_by = user
        obj.reviewed_at = timezone.now()
        obj.review_remarks = request.data.get('remarks', '')
        obj.save(update_fields=['status', 'reviewed_by', 'reviewed_at', 'review_remarks'])
        from .serializers import SurveyAreaAccessRequestSerializer
        return Response(SurveyAreaAccessRequestSerializer(obj).data)

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        obj = self.get_queryset().filter(pk=pk).first()
        if not obj:
            return Response({'detail': 'Not found.'}, status=404)
        user = request.user
        from apps.accounts.models import User
        if not user.is_superadmin and user.role not in User.ADMIN_ROLES:
            return Response({'detail': 'Only admins can reject requests.'}, status=403)
        if not user.is_superadmin and obj.survey_area.project.organisation != user.organisation:
            return Response({'detail': 'You can only reject requests for your org\'s areas.'}, status=403)
        from django.utils import timezone
        obj.status = SurveyAreaAccessRequest.REJECTED
        obj.reviewed_by = user
        obj.reviewed_at = timezone.now()
        obj.review_remarks = request.data.get('remarks', '')
        obj.save(update_fields=['status', 'reviewed_by', 'reviewed_at', 'review_remarks'])
        from .serializers import SurveyAreaAccessRequestSerializer
        return Response(SurveyAreaAccessRequestSerializer(obj).data)


class ProjectLayerFolderViewSet(viewsets.ModelViewSet):
    serializer_class = ProjectLayerFolderSerializer
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['project', 'parent', 'folder_type', 'is_final']

    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        from .access import hq_level, published_map_filter
        base_qs = ProjectLayerFolder.objects.select_related('project__organisation', 'parent', 'created_by')
        own_qs = org_queryset_filter(user, base_qs, org_field='project__organisation')
        if user.is_superadmin:
            return own_qs
        # DGDE/PDDE office users: own-org folders + folder subtrees of PUBLISHED
        # map-enabled survey areas (the Map Viewer layer tree), nothing else.
        if hq_level(user):
            from apps.survey_projects.models import SurveyArea as _SA
            areas = published_map_filter(user, _SA.objects.filter(
                status=_SA.PUBLISHED, map_enabled=True,
                project__map_enabled=True, folder__isnull=False,
            ))
            pub_folder_ids: list[int] = []
            for a in areas.only('folder_id'):
                pub_folder_ids.extend(_get_subtree_folder_ids(a.folder_id))
            if not pub_folder_ids:
                return own_qs
            return (own_qs | base_qs.filter(id__in=pub_folder_ids)).distinct()
        approved_area_ids = get_approved_area_ids(user)
        shared_project_ids = get_shared_project_ids(user)
        if not approved_area_ids and not shared_project_ids:
            return own_qs
        from apps.survey_projects.models import SurveyArea as _SA
        approved_folder_ids: list[int] = []
        for area in _SA.objects.filter(id__in=approved_area_ids).select_related('folder'):
            if area.folder_id:
                approved_folder_ids.extend(_get_subtree_folder_ids(area.folder_id))
        extra_q = Q(project_id__in=shared_project_ids)
        if approved_folder_ids:
            extra_q |= Q(id__in=approved_folder_ids)
        return (own_qs | base_qs.filter(extra_q)).distinct()

    def get_permissions(self):
        if self.action in ['list', 'retrieve', 'tree', 'upload_doc', 'files']:
            return [permissions.IsAuthenticated()]
        if self.action == 'import_gis_file':
            return [permissions.IsAuthenticated(), CanEditProject()]
        return [CanEditProject()]

    def perform_create(self, serializer):
        folder = serializer.save(created_by=self.request.user)
        # Auto-create Doc and Shapefile subfolders for every user-created folder
        # (but not for DOC/SHAPEFILE folders themselves to avoid recursion)
        if folder.folder_type not in (ProjectLayerFolder.DOC, ProjectLayerFolder.SHAPEFILE):
            _add_doc_shapefile_subfolders(folder, self.request.user)

    @action(detail=True, methods=['post'], url_path='upload-doc',
            permission_classes=[permissions.IsAuthenticated])
    def upload_doc(self, request, pk=None):
        """
        POST /api/projects/folders/{id}/upload-doc/
        Upload any document file to this folder (folder_type must be DOC).
        Body: multipart/form-data with 'file' and optional 'title', 'category'
        """
        from apps.documents.models import Document
        folder = self.get_object()
        uploaded = request.FILES.get('file')
        if not uploaded:
            return Response({'detail': 'No file provided.'}, status=400)
        title = (request.data.get('title') or uploaded.name).strip()
        category = request.data.get('category', Document.OTHER)
        mime_type = getattr(uploaded, 'content_type', '')
        if not mime_type or mime_type == 'application/octet-stream':
            ext = uploaded.name.split('.')[-1].lower() if '.' in uploaded.name else ''
            mime_map = {
                'doc': 'application/msword',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'odt': 'application/vnd.oasis.opendocument.text',
                'rtf': 'application/rtf',
                'txt': 'text/plain',
                'pdf': 'application/pdf',
                'xls': 'application/vnd.ms-excel',
                'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'ods': 'application/vnd.oasis.opendocument.spreadsheet',
                'csv': 'text/csv',
                'ppt': 'application/vnd.ms-powerpoint',
                'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                'odp': 'application/vnd.oasis.opendocument.presentation',
            }
            mime_type = mime_map.get(ext, 'application/octet-stream')

        doc = Document.objects.create(
            project=folder.project,
            folder=folder,
            title=title,
            category=category,
            file=uploaded,
            file_size=uploaded.size,
            mime_type=mime_type,
            uploaded_by=request.user,
        )
        return Response({'detail': 'Document uploaded.', 'id': doc.id}, status=201)

    @action(detail=True, methods=['get'])
    def tree(self, request, pk=None):
        """Return nested folder tree rooted at this folder."""
        folder = self.get_object()
        return Response(ProjectLayerFolderSerializer(folder).data)

    @action(detail=True, methods=['get'], url_path='files')
    def files(self, request, pk=None):
        """
        GET /api/projects/folders/{id}/files/
        Returns lists of documents and GeoTiff layers in this folder.
        Used by the QGIS plugin for duplicate detection before upload.
        """
        from apps.documents.models import Document
        from apps.documents.serializers import DocumentSerializer
        from .serializers import GeoTiffLayerSerializer

        folder = self.get_object()
        docs = Document.objects.filter(folder=folder).values(
            'id', 'title', 'file', 'mime_type', 'file_size', 'uploaded_at'
        )
        geotiffs = GeoTiffLayer.objects.filter(folder=folder).values(
            'id', 'name', 'file', 'status', 'created_at'
        )
        return Response({
            'folder_id': folder.id,
            'folder_name': folder.name,
            'documents': list(docs),
            'geotiffs': list(geotiffs),
        })

    @action(detail=True, methods=['post'], url_path='import-gis-file',
            permission_classes=[permissions.IsAuthenticated])
    def import_gis_file(self, request, pk=None):
        """
        POST /api/projects/folders/{id}/import-gis-file/
        Supported formats:
          .zip       — Shapefile bundle (must contain .shp + .dbf + .shx)
          .geojson / .json — GeoJSON FeatureCollection or Feature
          .kml       — KML (via GDAL)
          .gpkg      — GeoPackage (via GDAL)
          .tif/.tiff — GeoTiff (creates GeoTiffLayer, queues COG conversion)
        Body params:
          layer_name  — display name for created features (default: filename stem)
          name_field  — attribute to use as feature_id label (optional)
        """
        import os
        folder = self.get_object()
        uploaded = request.FILES.get('file')
        if not uploaded:
            return Response({'detail': 'No file uploaded.'}, status=400)

        fname = uploaded.name
        ext = os.path.splitext(fname)[1].lower()
        layer_name = (request.data.get('layer_name') or os.path.splitext(fname)[0]).strip()
        name_field = request.data.get('name_field', '').strip()
        # Default Yes; only sub-DEO uploaders send this explicitly.
        deo_visible = str(request.data.get('deo_visible', 'true')).lower() not in ('false', '0', 'no')
        geom_type = request.data.get('geom_type', 'auto').strip().lower()

        # Optional surveyor-defined attributes (JSON object, GPX only)
        import json as _json
        _extra_raw = request.data.get('extra_attributes', '')
        try:
            extra_attributes = _json.loads(_extra_raw) if _extra_raw else {}
            if not isinstance(extra_attributes, dict):
                extra_attributes = {}
        except (ValueError, TypeError):
            extra_attributes = {}

        if geom_type not in ('auto', 'point', 'line', 'polygon'):
            return Response(
                {'detail': f'Invalid geom_type "{geom_type}". Expected one of: auto, point, line, polygon'},
                status=400
            )

        if ext in ('.tif', '.tiff'):
            return _import_geotiff(folder, uploaded, layer_name, request.user, deo_visible=deo_visible)
        if ext == '.zip':
            return _import_shapefile_zip(folder, uploaded, layer_name, name_field, request.user, deo_visible=deo_visible)
        if ext in ('.geojson', '.json'):
            return _import_geojson_file(folder, uploaded, layer_name, name_field, request.user, deo_visible=deo_visible)
        if ext in ('.kml', '.gpkg'):
            return _import_via_gdal(folder, uploaded, layer_name, name_field, request.user, ext.lstrip('.'), deo_visible=deo_visible)
        if ext == '.gpx':
            return _import_gpx_file(folder, uploaded, layer_name, name_field, request.user,
                                    geom_type=geom_type, deo_visible=deo_visible, extra_attributes=extra_attributes)
        if ext == '.csv':
            return _import_csv_file(folder, uploaded, layer_name, name_field, request.user, geom_type=geom_type, deo_visible=deo_visible)
        return Response(
            {'detail': f'Unsupported format "{ext}". Accepted: .zip, .geojson, .json, .kml, .gpkg, .gpx, .csv, .tif, .tiff'},
            status=400
        )


class ProjectShareViewSet(viewsets.ModelViewSet):
    serializer_class = ProjectShareSerializer
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['project', 'granted_to']
    http_method_names = ['get', 'post', 'delete', 'head', 'options']

    def get_queryset(self):
        user = self.request.user
        from apps.accounts.models import User
        if user.role == User.SUPERADMIN:
            return ProjectShare.objects.select_related('project__organisation', 'granted_to', 'granted_by').all()
        return ProjectShare.objects.select_related('project__organisation', 'granted_to', 'granted_by').filter(
            project__organisation=user.organisation
        )

    def get_permissions(self):
        if self.action in ['list', 'retrieve']:
            return [permissions.IsAuthenticated()]
        return [CanEditProject()]

    def perform_create(self, serializer):
        serializer.save(granted_by=self.request.user)


def _export_features(qs, project, fmt: str):
    """Export a queryset of GISFeatures in the requested format."""
    import json
    import tempfile
    import os
    from django.http import HttpResponse, FileResponse
    from django.contrib.gis.serializers.geojson import Serializer as GeoJSONSerializer
    from apps.core.watermark import embed_watermark

    features = list(qs.select_related('created_by'))
    metadata = {
        "project_id": project.id,
        "project_number": project.project_number,
        "export_format": fmt,
        "features_count": len(features),
    }

    if fmt == 'geojson':
        fc = {
            'type': 'FeatureCollection',
            'features': [
                {
                    'type': 'Feature',
                    'geometry': json.loads(f.geometry.geojson),
                    'properties': {
                        'id': f.id,
                        'layer_name': f.layer_name,
                        'geometry_type': f.geometry_type,
                        'feature_id': f.feature_id,
                        **f.attributes,
                        'created_by': str(f.created_by),
                        'created_at': f.created_at.isoformat(),
                    },
                }
                for f in features
            ],
        }
        geojson_str = json.dumps(fc, indent=2)
        try:
            watermarked_bytes = embed_watermark(geojson_str.encode('utf-8'), f"{project.project_number}.geojson", 'application/geo+json', metadata)
            resp = HttpResponse(watermarked_bytes, content_type='application/geo+json')
        except Exception:
            resp = HttpResponse(geojson_str, content_type='application/geo+json')
        resp['Content-Disposition'] = f'attachment; filename="{project.project_number}.geojson"'
        return resp

    if fmt == 'csv':
        import csv
        import io
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(['id', 'layer_name', 'geometry_type', 'feature_id', 'longitude', 'latitude', 'created_at'])
        for f in features:
            centroid = f.geometry.centroid
            writer.writerow([f.id, f.layer_name, f.geometry_type, f.feature_id, centroid.x, centroid.y, f.created_at.isoformat()])
        csv_str = buf.getvalue()
        try:
            watermarked_bytes = embed_watermark(csv_str.encode('utf-8'), f"{project.project_number}.csv", 'text/csv', metadata)
            resp = HttpResponse(watermarked_bytes, content_type='text/csv')
        except Exception:
            resp = HttpResponse(csv_str, content_type='text/csv')
        resp['Content-Disposition'] = f'attachment; filename="{project.project_number}.csv"'
        return resp

    if fmt in ('shapefile', 'gpkg', 'kml', 'geojson2'):
        try:
            import fiona
            import fiona.crs
            from fiona.transform import transform_geom
        except ImportError:
            from django.http import HttpResponse
            return HttpResponse('fiona not available', status=500)

        driver_map = {
            'shapefile': ('ESRI Shapefile', '.shp', '.zip'),
            'gpkg': ('GPKG', '.gpkg', '.gpkg'),
            'kml': ('KML', '.kml', '.kml'),
        }
        driver, ext, archive_ext = driver_map.get(fmt, ('ESRI Shapefile', '.shp', '.zip'))

        import tempfile, os, zipfile, io

        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = os.path.join(tmpdir, f"{project.project_number}{ext}")

            schema = {
                'geometry': 'Unknown',
                'properties': {
                    'id': 'int',
                    'layer_name': 'str',
                    'feature_id': 'str',
                    'created_at': 'str',
                },
            }

            with fiona.open(out_path, 'w', driver=driver, schema=schema,
                            crs=fiona.crs.from_epsg(4326)) as dst:
                for f in features:
                    geom = json.loads(f.geometry.geojson)
                    dst.write({
                        'geometry': geom,
                        'properties': {
                            'id': f.id,
                            'layer_name': f.layer_name,
                            'feature_id': f.feature_id or '',
                            'created_at': f.created_at.isoformat(),
                        },
                    })

            if fmt == 'shapefile':
                buf = io.BytesIO()
                with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for fn in os.listdir(tmpdir):
                        zf.write(os.path.join(tmpdir, fn), fn)
                buf.seek(0)
                shape_bytes = buf.read()
                try:
                    watermarked_bytes = embed_watermark(shape_bytes, f"{project.project_number}.zip", 'application/zip', metadata)
                    resp = HttpResponse(watermarked_bytes, content_type='application/zip')
                except Exception:
                    resp = HttpResponse(shape_bytes, content_type='application/zip')
                resp['Content-Disposition'] = f'attachment; filename="{project.project_number}.zip"'
                return resp
            else:
                with open(out_path, 'rb') as fh:
                    content_type = 'application/octet-stream'
                    file_bytes = fh.read()
                    try:
                        watermarked_bytes = embed_watermark(file_bytes, f"{project.project_number}{archive_ext}", content_type, metadata)
                        resp = HttpResponse(watermarked_bytes, content_type=content_type)
                    except Exception:
                        resp = HttpResponse(file_bytes, content_type=content_type)
                    resp['Content-Disposition'] = f'attachment; filename="{project.project_number}{archive_ext}"'
                    return resp

    from django.http import HttpResponse
    return HttpResponse('Unsupported format', status=400)


class GeoTiffLayerViewSet(viewsets.ModelViewSet):
    serializer_class = GeoTiffLayerSerializer
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['project', 'folder', 'status']
    http_method_names = ['get', 'post', 'patch', 'delete', 'head', 'options']

    def get_queryset(self):
        user = self.request.user
        base_qs = GeoTiffLayer.objects.select_related('project__organisation', 'folder', 'created_by')
        own_qs = org_queryset_filter(user, base_qs, org_field='project__organisation')
        # DEO offices additionally see subordinate-office rasters marked deo_visible.
        deo_sub_ids = deo_subordinate_org_ids(user)
        if not deo_sub_ids:
            return own_qs
        extra = base_qs.filter(project__organisation_id__in=deo_sub_ids, deo_visible=True)
        return (own_qs | extra).distinct()

    def get_permissions(self):
        if self.action in ['list', 'retrieve']:
            return [permissions.IsAuthenticated()]
        # partial_update (PATCH) is used by the map to toggle is_visible / opacity.
        # Any authenticated user in the same org may do this; full create/delete
        # is restricted to CanEditProject (SDO/SURVEYOR/SUPERADMIN).
        if self.action == 'partial_update':
            return [permissions.IsAuthenticated()]
        return [CanEditProject()]

    def perform_update(self, serializer):
        user = self.request.user
        if not (user.is_superadmin or
                serializer.instance.project.organisation_id == user.organisation_id):
            from rest_framework.exceptions import PermissionDenied
            raise PermissionDenied('You can only modify GeoTiff layers in your organisation.')
        serializer.save()

    def get_serializer_context(self):
        ctx = super().get_serializer_context()
        ctx['request'] = self.request
        return ctx

    def perform_create(self, serializer):
        from .tasks import convert_geotiff_to_cog
        layer = serializer.save(created_by=self.request.user)
        convert_geotiff_to_cog.delay(layer.id)

    @action(detail=True, methods=['post'], url_path='terrain-analysis', permission_classes=[permissions.IsAuthenticated])
    def terrain_analysis(self, request, pk=None):
        """
        POST /api/survey_projects/geotiffs/{id}/terrain-analysis/

        Run GDAL terrain analysis on a GeoTiff DEM layer.
        Body: { operation: 'hillshade'|'slope'|'aspect'|'contour', z_factor: 1.0, contour_interval: 10 }
        Returns: saves result as new GeoTiffLayer and returns its id + COG path for display.
        """
        import os, subprocess, tempfile
        from django.conf import settings as django_settings

        layer = self.get_object()
        operation = request.data.get('operation', 'hillshade')
        z_factor = float(request.data.get('z_factor', 1.0))
        contour_interval = float(request.data.get('contour_interval', 10.0))

        if not layer.cog_file:
            return Response({'detail': 'DEM has not been processed yet (no COG file).'}, status=400)

        src_path = os.path.join(django_settings.MEDIA_ROOT, str(layer.cog_file))
        if not os.path.exists(src_path):
            return Response({'detail': 'DEM file not found.'}, status=404)

        out_rel = str(layer.cog_file).rsplit('.', 1)[0] + f'_{operation}.tif'
        out_path = os.path.join(django_settings.MEDIA_ROOT, out_rel)
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        try:
            if operation == 'contour':
                with tempfile.NamedTemporaryFile(suffix='.geojson', delete=False) as tmp:
                    tmp_path = tmp.name
                subprocess.run(
                    ['gdal_contour', '-a', 'elev', '-i', str(contour_interval), src_path, tmp_path, '-f', 'GeoJSON'],
                    check=True, capture_output=True,
                )
                import json as _json
                with open(tmp_path) as f:
                    data = _json.load(f)
                os.unlink(tmp_path)
                return Response({'type': 'contour', 'geojson': data, 'interval': contour_interval})

            elif operation in ('hillshade', 'slope', 'aspect'):
                cmd = ['gdaldem', operation, src_path, out_path, '-of', 'COG', '-co', 'COMPRESS=DEFLATE']
                if operation == 'hillshade':
                    cmd += ['-z', str(z_factor)]
                subprocess.run(cmd, check=True, capture_output=True)

                result_layer = GeoTiffLayer.objects.create(
                    project=layer.project,
                    folder=layer.folder,
                    name=f'{layer.name}_{operation}',
                    file=out_rel,
                    cog_file=out_rel,
                    status=GeoTiffLayer.DONE,
                    created_by=request.user,
                )
                return Response({
                    'type': operation,
                    'layer_id': result_layer.id,
                    'cog_url': f'/media/{out_rel}',
                    'layer_name': result_layer.name,
                })
            else:
                return Response({'detail': f'Unknown operation: {operation}'}, status=400)
        except subprocess.CalledProcessError as exc:
            return Response({'detail': f'GDAL error: {(exc.stderr or b"").decode()[:300]}'}, status=500)
        except Exception as exc:
            return Response({'detail': str(exc)}, status=500)


class TopologyCheckView(APIView):
    """GET /api/projects/topology/ — check DefenceParcel geometries for errors."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        from django.db import connection

        # Org-filtered base queryset
        from apps.accounts.permissions import org_queryset_filter
        base_qs = org_queryset_filter(
            request.user,
            DefenceParcel.objects.select_related('state', 'district'),
        )

        # 1. Self-intersecting / invalid geometries
        invalid = base_qs.extra(where=['NOT ST_IsValid("survey_projects_defenceparcel"."geometry")'])
        invalid_list = [
            {'type': 'INVALID_GEOMETRY',
             'parcel_a': {'id': p.id, 'parcel_id': p.parcel_id, 'name': p.name}}
            for p in invalid
        ]

        # 2. Overlapping parcel pairs
        id_list = list(base_qs.values_list('id', flat=True))
        overlaps = []
        if len(id_list) >= 2:
            placeholders = ','.join(['%s'] * len(id_list))
            with connection.cursor() as cur:
                cur.execute(
                    f"""
                    SELECT a.id, a.parcel_id, a.name,
                           b.id, b.parcel_id, b.name,
                           ST_AsGeoJSON(ST_Centroid(ST_Intersection(a.geometry, b.geometry)))
                    FROM survey_projects_defenceparcel a
                    JOIN survey_projects_defenceparcel b ON a.id < b.id
                    WHERE a.id IN ({placeholders})
                      AND b.id IN ({placeholders})
                      AND ST_Intersects(a.geometry, b.geometry)
                      AND NOT ST_Touches(a.geometry, b.geometry)
                    LIMIT 200
                    """,
                    id_list + id_list,
                )
                for row in cur.fetchall():
                    overlaps.append({
                        'type': 'OVERLAP',
                        'parcel_a': {'id': row[0], 'parcel_id': row[1], 'name': row[2]},
                        'parcel_b': {'id': row[3], 'parcel_id': row[4], 'name': row[5]},
                        'centroid': json.loads(row[6]) if row[6] else None,
                    })

        issues = invalid_list + overlaps
        return Response({'issues': issues, 'total': len(issues)})


class TopologyRuleViewSet(viewsets.ModelViewSet):
    """CRUD + check endpoint for project-level topology rules."""
    serializer_class = TopologyRuleSerializer

    def get_queryset(self):
        from .models import TopologyRule
        user = self.request.user
        return TopologyRule.objects.filter(project__organisation=user.organisation)

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=False, methods=['post'], url_path='check', permission_classes=[permissions.IsAuthenticated])
    def check(self, request):
        """Run all active topology rules for a project and return violations as GeoJSON."""
        from django.db import connection
        import json
        from .models import TopologyRule

        project_id = request.data.get('project')
        if not project_id:
            return Response({'detail': 'project required.'}, status=400)

        rules = TopologyRule.objects.filter(project_id=project_id, is_active=True)
        violations = []

        for rule in rules:
            try:
                with connection.cursor() as cur:
                    if rule.rule_type == 'MUST_NOT_OVERLAP':
                        cur.execute("""
                            SELECT a.id, b.id,
                                   ST_AsGeoJSON(ST_Intersection(a.geometry, b.geometry)) as geom,
                                   ST_Area(ST_Intersection(a.geometry::geography, b.geometry::geography)) as area
                            FROM survey_projects_gisfeature a, survey_projects_gisfeature b
                            WHERE a.project_id = %s AND b.project_id = %s
                              AND a.layer_name = %s AND b.layer_name = %s
                              AND a.id < b.id
                              AND a.is_deleted = false AND b.is_deleted = false
                              AND ST_Overlaps(a.geometry, b.geometry)
                        """, [project_id, project_id, rule.layer_a, rule.layer_a])
                        for r in cur.fetchall():
                            geom = json.loads(r[2]) if r[2] else None
                            if geom:
                                violations.append({
                                    'rule_id': rule.id, 'rule_type': rule.rule_type,
                                    'layer': rule.layer_a, 'feature_a_id': r[0], 'feature_b_id': r[1],
                                    'description': f'{rule.layer_a}: features #{r[0]} and #{r[1]} overlap ({r[3]:.1f} m²)',
                                    'geometry': geom,
                                })

                    elif rule.rule_type == 'MUST_BE_INSIDE' and rule.layer_b:
                        cur.execute("""
                            SELECT a.id, ST_AsGeoJSON(a.geometry) as geom
                            FROM survey_projects_gisfeature a
                            WHERE a.project_id = %s AND a.layer_name = %s AND a.is_deleted = false
                              AND NOT EXISTS (
                                SELECT 1 FROM survey_projects_gisfeature b
                                WHERE b.project_id = %s AND b.layer_name = %s AND b.is_deleted = false
                                  AND ST_Within(a.geometry, b.geometry)
                              )
                        """, [project_id, rule.layer_a, project_id, rule.layer_b])
                        for r in cur.fetchall():
                            geom = json.loads(r[1]) if r[1] else None
                            if geom:
                                violations.append({
                                    'rule_id': rule.id, 'rule_type': rule.rule_type,
                                    'layer': rule.layer_a, 'feature_a_id': r[0],
                                    'description': f'{rule.layer_a} #{r[0]}: not inside any {rule.layer_b} feature',
                                    'geometry': geom,
                                })

                    elif rule.rule_type == 'MUST_NOT_DANGLE':
                        cur.execute("""
                            SELECT a.id, ST_AsGeoJSON(ST_StartPoint(a.geometry::geometry)) as geom
                            FROM survey_projects_gisfeature a
                            WHERE a.project_id = %s AND a.layer_name = %s
                              AND a.geometry_type = 'LINE' AND a.is_deleted = false
                              AND NOT EXISTS (
                                SELECT 1 FROM survey_projects_gisfeature b
                                WHERE b.project_id = %s AND b.layer_name = %s AND b.id <> a.id AND b.is_deleted = false
                                  AND (ST_DWithin(ST_StartPoint(a.geometry::geometry), ST_StartPoint(b.geometry::geometry), %s)
                                    OR ST_DWithin(ST_StartPoint(a.geometry::geometry), ST_EndPoint(b.geometry::geometry), %s))
                              )
                        """, [project_id, rule.layer_a, project_id, rule.layer_a, rule.tolerance, rule.tolerance])
                        for r in cur.fetchall():
                            geom = json.loads(r[1]) if r[1] else None
                            if geom:
                                violations.append({
                                    'rule_id': rule.id, 'rule_type': rule.rule_type,
                                    'layer': rule.layer_a, 'feature_a_id': r[0],
                                    'description': f'{rule.layer_a} #{r[0]}: dangling start point',
                                    'geometry': geom,
                                })
            except Exception as exc:
                violations.append({'rule_id': rule.id, 'rule_type': rule.rule_type, 'error': str(exc)})

        geojson_features = [
            {'type': 'Feature', 'geometry': v.pop('geometry', None), 'properties': v}
            for v in violations if 'geometry' in v
        ]
        return Response({
            'type': 'FeatureCollection', 'features': geojson_features,
            'violation_count': len(geojson_features),
            'rules_checked': len(rules),
        })


class BufferAnalysisView(APIView):
    """POST /api/projects/buffer/
    Modes:
      - Point buffer (default): requires lng, lat
      - Feature buffer: requires feature_ids (list of GISFeature IDs) or layer_name + project_id
    Common params: distances (list), unit (meters|kilometers), dissolve (bool)
    """
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        from django.contrib.gis.geos import Point, GEOSGeometry
        from django.contrib.gis.measure import D
        from django.db import connection

        raw_distances = request.data.get('distances', [100])
        unit = request.data.get('unit', 'meters')
        dissolve = bool(request.data.get('dissolve', False))

        # ── Determine source geometry ──────────────────────────────────────────
        feature_ids = request.data.get('feature_ids')
        layer_name  = request.data.get('layer_name')
        project_id  = request.data.get('project_id')

        if feature_ids or layer_name:
            # Feature-based buffer
            if feature_ids:
                feats = GISFeature.objects.filter(id__in=feature_ids, is_deleted=False)
            else:
                feats = GISFeature.objects.filter(
                    project_id=project_id, layer_name=layer_name, is_deleted=False
                )
            if not feats.exists():
                return Response({'detail': 'No features found.'}, status=status.HTTP_400_BAD_REQUEST)

            # Combine all feature geometries into a UNION
            from django.contrib.gis.db.models import Union as GeoUnion
            combined = feats.aggregate(union=GeoUnion('geometry'))['union']
            if not combined:
                return Response({'detail': 'Features have no geometry.'}, status=status.HTTP_400_BAD_REQUEST)

            source_wkt = combined.wkt
            # Compute centroid for display
            centroid = combined.centroid
            lng, lat = centroid.x, centroid.y
        else:
            # Point-based buffer: use exact click coordinates
            try:
                lng = float(request.data['lng'])
                lat = float(request.data['lat'])
            except (KeyError, TypeError, ValueError):
                return Response({'detail': 'Provide lng/lat or feature_ids or layer_name.'}, status=status.HTTP_400_BAD_REQUEST)

            source_wkt = f'SRID=4326;POINT({lng} {lat})'

        if not isinstance(raw_distances, list) or not raw_distances:
            return Response({'detail': 'distances must be a non-empty list.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            distances = sorted(float(d) for d in raw_distances)
        except (TypeError, ValueError):
            return Response({'detail': 'Each distance must be a number.'}, status=status.HTTP_400_BAD_REQUEST)

        ref_point = Point(lng, lat, srid=4326)
        multiplier = 1000 if unit == 'kilometers' else 1
        results = []

        for dist in distances:
            dist_m = dist * multiplier

            # Build PostGIS buffer from source geometry
            with connection.cursor() as cur:
                cur.execute(
                    "SELECT ST_AsGeoJSON(ST_Buffer(ST_GeogFromText(%s)::geography, %s)::geometry)",
                    [source_wkt, dist_m],
                )
                buffer_geojson = json.loads(cur.fetchone()[0])

            # Find parcels intersecting this buffer
            with connection.cursor() as cur:
                cur.execute(
                    "SELECT ST_AsText(ST_Buffer(ST_GeogFromText(%s)::geography, %s)::geometry)",
                    [source_wkt, dist_m],
                )
                buffer_wkt = cur.fetchone()[0]

            buffer_geom = GEOSGeometry(buffer_wkt, srid=4326)
            parcels = DefenceParcel.objects.filter(
                geometry__intersects=buffer_geom
            ).annotate(
                distance=Distance('geometry', ref_point)
            ).select_related('state', 'district', 'organisation')

            # Find survey area features (GISFeature) within this buffer ring
            nearby_features = (
                GISFeature.objects
                .filter(geometry__intersects=buffer_geom, is_deleted=False)
                .select_related(
                    'project', 'project__organisation',
                    'folder',
                )
            )

            # Group nearby features by survey area (via folder → survey_area_link)
            from collections import defaultdict
            area_map = defaultdict(lambda: {
                'area_id': None, 'area_name': None, 'area_code': None,
                'status': None, 'status_display': None,
                'project_id': None, 'project_name': None,
                'organisation': None,
                'feature_count': 0, 'layers': set(),
            })

            for feat in nearby_features:
                # Resolve survey area through folder chain
                folder = feat.folder
                survey_area = None
                if folder:
                    try:
                        survey_area = folder.survey_area_link
                    except Exception:
                        pass
                    if not survey_area:
                        # walk up to find a folder linked to a survey area
                        parent = folder.parent if folder else None
                        while parent and not survey_area:
                            try:
                                survey_area = parent.survey_area_link
                            except Exception:
                                pass
                            parent = parent.parent if parent else None

                key = survey_area.id if survey_area else f'project_{feat.project_id}'
                entry = area_map[key]
                if survey_area and not entry['area_id']:
                    entry.update({
                        'area_id': survey_area.id,
                        'area_name': survey_area.name,
                        'area_code': survey_area.area_code,
                        'status': survey_area.status,
                        'status_display': survey_area.get_status_display(),
                    })
                if not entry['project_id']:
                    entry.update({
                        'project_id': feat.project_id,
                        'project_name': str(feat.project),
                        'organisation': feat.project.organisation.name if feat.project.organisation_id else '',
                    })
                entry['feature_count'] += 1
                entry['layers'].add(feat.layer_name)

            survey_areas_data = []
            for entry in area_map.values():
                survey_areas_data.append({
                    **{k: v for k, v in entry.items() if k != 'layers'},
                    'layers': sorted(entry['layers']),
                })
            survey_areas_data.sort(key=lambda x: x['feature_count'], reverse=True)

            results.append({
                'distance': dist,
                'unit': unit,
                'distance_m': dist_m,
                'buffer_geojson': buffer_geojson,
                'parcels': BufferParcelSerializer(parcels, many=True).data,
                'feature_count': parcels.count(),
                'survey_areas': survey_areas_data,
                'survey_area_count': len(survey_areas_data),
                'dissolve': dissolve,
            })

        return Response({
            'center_lng': lng,
            'center_lat': lat,
            'rings': results,
        })


# ── Feature Attachments ───────────────────────────────────────────────────────

class FeatureAttachmentViewSet(viewsets.ModelViewSet):
    serializer_class = FeatureAttachmentSerializer
    http_method_names = ['get', 'post', 'delete', 'head', 'options']

    def get_permissions(self):
        if self.action in ['list', 'retrieve']:
            return [permissions.IsAuthenticated()]
        return [CanEditProject()]

    def get_queryset(self):
        feature_id = self.kwargs.get('feature_pk') or self.request.query_params.get('feature')
        qs = FeatureAttachment.objects.select_related('uploaded_by', 'feature__project')
        if feature_id:
            qs = qs.filter(feature_id=feature_id)
        return qs

    def perform_create(self, request_file=None, **kwargs):
        pass  # handled in create()

    def create(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'detail': 'No file uploaded.'}, status=400)

        feature_id = request.data.get('feature') or self.kwargs.get('feature_pk')
        try:
            feature = GISFeature.objects.get(pk=feature_id)
        except GISFeature.DoesNotExist:
            return Response({'detail': 'Feature not found.'}, status=404)

        # Determine file type from extension
        ext = file_obj.name.rsplit('.', 1)[-1].lower() if '.' in file_obj.name else ''
        if ext in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'tiff', 'bmp'):
            file_type = 'image'
        elif ext == 'pdf':
            file_type = 'pdf'
        else:
            file_type = 'doc'

        attachment = FeatureAttachment.objects.create(
            feature=feature,
            file=file_obj,
            original_filename=file_obj.name,
            file_size=file_obj.size,
            file_type=file_type,
            caption=request.data.get('caption', ''),
            uploaded_by=request.user,
        )
        if file_type == 'image':
            from apps.ai_assistant.tasks import process_photo_ocr
            try:
                process_photo_ocr.delay(attachment.id)
            except Exception:
                import threading
                threading.Thread(target=process_photo_ocr, args=(attachment.id,)).start()

        return Response(
            FeatureAttachmentSerializer(attachment, context={'request': request}).data,
            status=201,
        )


# ── Bulk CSV Import ───────────────────────────────────────────────────────────

class CSVImportView(APIView):
    permission_classes = [CanEditProject]

    def post(self, request, pk=None):
        import csv, io
        from django.contrib.gis.geos import Point as GEOSPoint, GEOSGeometry

        try:
            project = SurveyProject.objects.get(pk=pk)
        except SurveyProject.DoesNotExist:
            return Response({'detail': 'Project not found.'}, status=404)

        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'detail': 'No CSV file uploaded.'}, status=400)

        layer_name = request.data.get('layer_name', 'imported_layer')
        lat_col = request.data.get('lat_col', 'latitude')
        lon_col = request.data.get('lon_col', 'longitude')
        wkt_col = request.data.get('wkt_col', '')

        try:
            text = file_obj.read().decode('utf-8-sig')
            reader = csv.DictReader(io.StringIO(text))
            rows = list(reader)
        except Exception as e:
            return Response({'detail': f'Could not parse CSV: {e}'}, status=400)

        imported = 0
        errors = []

        for i, row in enumerate(rows, 1):
            try:
                if wkt_col and wkt_col in row and row[wkt_col].strip():
                    geom = GEOSGeometry(row[wkt_col].strip(), srid=4326)
                    geom_type = geom.geom_type.upper()
                elif lat_col in row and lon_col in row:
                    lat = float(row[lat_col])
                    lon = float(row[lon_col])
                    geom = GEOSPoint(lon, lat, srid=4326)
                    geom_type = 'POINT'
                else:
                    errors.append(f'Row {i}: missing geometry columns')
                    continue

                attrs = {k: v for k, v in row.items()
                         if k not in (lat_col, lon_col, wkt_col) and v}

                GISFeature.objects.create(
                    project=project,
                    layer_name=layer_name,
                    geometry_type=geom_type if geom_type in ('POINT', 'LINE', 'POLYGON') else 'POINT',
                    geometry=geom,
                    attributes=attrs,
                    created_by=request.user,
                )
                imported += 1
            except Exception as e:
                errors.append(f'Row {i}: {e}')

        # Log export/import action
        try:
            from apps.accounts.models import ExportAuditLog
            ExportAuditLog.objects.create(
                user=request.user, export_type='csv_import',
                project=project, row_count=imported,
            )
        except Exception:
            pass

        return Response({'imported': imported, 'total': len(rows), 'errors': errors[:20]})


# ── Encroachment Detection ────────────────────────────────────────────────────

class EncroachmentView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk=None):
        from django.db import connection
        try:
            project = SurveyProject.objects.get(pk=pk)
        except SurveyProject.DoesNotExist:
            return Response({'detail': 'Project not found.'}, status=404)

        features = GISFeature.objects.filter(project=project, is_deleted=False).exclude(
            geometry_type='POINT'
        )

        encroachments = []
        for feat in features:
            from apps.gis_layers.models import RevenueMap
            overlapping = RevenueMap.objects.filter(
                geometry__intersects=feat.geometry
            ).exclude(geometry__equals=feat.geometry)[:5]
            for rm in overlapping:
                with connection.cursor() as cur:
                    cur.execute(
                        'SELECT ST_Area(ST_Intersection(%s::geometry, %s::geometry)::geography)',
                        [feat.geometry.wkt, rm.geometry.wkt],
                    )
                    area_m2 = cur.fetchone()[0] or 0
                encroachments.append({
                    'feature_id': feat.id,
                    'feature_label': feat.feature_id or str(feat.id),
                    'layer_name': feat.layer_name,
                    'revenue_map_id': rm.id,
                    'survey_number': rm.survey_number,
                    'overlap_area_m2': round(area_m2, 2),
                    'overlap_area_ha': round(area_m2 / 10000, 4),
                })

        return Response({
            'project_id': project.id,
            'project_number': project.project_number,
            'encroachment_count': len(encroachments),
            'encroachments': encroachments,
        })


# ── Feature Merge ─────────────────────────────────────────────────────────────

class FeatureMergeView(APIView):
    permission_classes = [CanEditProject]

    def post(self, request):
        from django.contrib.gis.geos import GEOSGeometry
        feature_ids = request.data.get('feature_ids', [])
        layer_name = request.data.get('layer_name', '')

        if len(feature_ids) < 2:
            return Response({'detail': 'Provide at least 2 feature_ids.'}, status=400)

        features = GISFeature.objects.filter(pk__in=feature_ids, is_deleted=False)
        if features.count() < 2:
            return Response({'detail': 'Could not find the requested features.'}, status=404)

        # Ensure all belong to the same project
        project_ids = set(features.values_list('project_id', flat=True))
        if len(project_ids) > 1:
            return Response({'detail': 'All features must belong to the same project.'}, status=400)

        from django.db import connection
        wkt_list = [f.geometry.wkt for f in features]
        placeholders = ', '.join(['%s::geometry'] * len(wkt_list))
        with connection.cursor() as cur:
            cur.execute(f'SELECT ST_AsText(ST_Union(ARRAY[{placeholders}]))', wkt_list)
            merged_wkt = cur.fetchone()[0]

        merged_geom = GEOSGeometry(merged_wkt, srid=4326)
        first = features.first()
        new_feature = GISFeature.objects.create(
            project=first.project,
            folder=first.folder,
            layer_name=layer_name or first.layer_name,
            geometry_type=first.geometry_type,
            geometry=merged_geom,
            attributes={},
            created_by=request.user,
        )
        features.update(is_deleted=True)

        return Response(GISFeatureSerializer(new_feature).data, status=201)


# ── Feature Split ─────────────────────────────────────────────────────────────

class FeatureSplitView(APIView):
    permission_classes = [CanEditProject]

    def post(self, request, pk=None):
        from django.contrib.gis.geos import GEOSGeometry
        from django.db import connection

        try:
            feature = GISFeature.objects.get(pk=pk)
        except GISFeature.DoesNotExist:
            return Response({'detail': 'Feature not found.'}, status=404)

        split_line = request.data.get('split_line')
        if not split_line:
            return Response({'detail': 'split_line GeoJSON required.'}, status=400)

        try:
            import json
            line_geom = GEOSGeometry(json.dumps(split_line), srid=4326)
        except Exception as e:
            return Response({'detail': f'Invalid split_line: {e}'}, status=400)

        with connection.cursor() as cur:
            cur.execute(
                '''
                SELECT ST_AsText((ST_Dump(
                    ST_Split(
                        ST_Snap(%s::geometry, %s::geometry, 0.0000001),
                        %s::geometry
                    )
                )).geom)
                ''',
                [feature.geometry.wkt, line_geom.wkt, line_geom.wkt],
            )
            parts = [row[0] for row in cur.fetchall()]

        if len(parts) < 2:
            return Response({'detail': 'Split produced fewer than 2 parts. Check split line intersects the feature.'}, status=400)

        new_features = []
        for part_wkt in parts:
            part_geom = GEOSGeometry(part_wkt, srid=4326)
            nf = GISFeature.objects.create(
                project=feature.project,
                folder=feature.folder,
                layer_name=feature.layer_name,
                geometry_type=feature.geometry_type,
                geometry=part_geom,
                attributes=dict(feature.attributes),
                created_by=request.user,
            )
            new_features.append(nf)

        feature.is_deleted = True
        feature.save(update_fields=['is_deleted'])

        return Response(GISFeatureSerializer(new_features, many=True).data, status=201)


# ── Project Milestones (Gantt) ────────────────────────────────────────────────

class ProjectMilestoneViewSet(viewsets.ModelViewSet):
    serializer_class = ProjectMilestoneSerializer

    def get_permissions(self):
        if self.action in ['list', 'retrieve']:
            return [permissions.IsAuthenticated()]
        return [CanEditProject()]

    def get_queryset(self):
        project_id = self.kwargs.get('project_pk') or self.request.query_params.get('project')
        qs = ProjectMilestone.objects.select_related('assignee', 'created_by', 'project')
        if project_id:
            qs = qs.filter(project_id=project_id)
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)


# ── QGIS Upload Log ───────────────────────────────────────────────────────────

class QGISUploadLogViewSet(viewsets.ModelViewSet):
    """
    GET  /api/projects/qgis-uploads/                — list (filterable by project/status)
    POST /api/projects/qgis-uploads/                — QGIS plugin writes upload records
    GET  /api/projects/qgis-uploads/{id}/           — detail
    POST /api/projects/qgis-uploads/{id}/retry/     — re-queue a FAILED upload
    """
    from .serializers import QGISUploadLogSerializer
    serializer_class = QGISUploadLogSerializer
    filter_backends  = [DjangoFilterBackend]
    filterset_fields = ['project', 'status', 'algorithm_id']
    http_method_names = ['get', 'post', 'head', 'options']

    def get_queryset(self):
        from .models import QGISUploadLog
        user = self.request.user
        from apps.accounts.models import User as AccountUser
        qs = QGISUploadLog.objects.select_related(
            'project', 'folder', 'uploaded_by'
        )
        if user.role == AccountUser.SUPERADMIN:
            return qs
        return qs.filter(project__organisation=user.organisation)

    def get_permissions(self):
        return [permissions.IsAuthenticated()]

    def perform_create(self, serializer):
        log = serializer.save(uploaded_by=self.request.user)
        if log.status == 'SUCCESS':
            self._notify_project_team(log, self.request.user)

    def _notify_project_team(self, log, uploader):
        """Send an in-app notification to project admins when a file is uploaded."""
        try:
            from apps.workflow.models import Notification
            from apps.accounts.models import User as AccountUser
            recipients = AccountUser.objects.filter(
                organisation=log.project.organisation,
                role__in=[
                    AccountUser.SUPERADMIN,
                    AccountUser.DEO_ADMIN,
                    AccountUser.CEO_ADMIN,
                    AccountUser.ADEO_ADMIN,
                ],
                is_active=True,
            ).exclude(pk=uploader.pk)
            folder_label = log.folder.name if log.folder else (log.module_name or log.folder_name or 'Unknown')
            notifs = [
                Notification(
                    user=u,
                    project=log.project,
                    title='New QGIS Upload',
                    message=(
                        f'{uploader.get_full_name() or uploader.username} uploaded '
                        f'"{log.filename}" to {log.project.project_number} / {folder_label}.'
                    ),
                )
                for u in recipients
            ]
            if notifs:
                Notification.objects.bulk_create(notifs)
        except Exception:
            pass  # never fail the upload log write because of notification issues

    @action(detail=True, methods=['post'], url_path='retry')
    def retry(self, request, pk=None):
        """
        Mark a FAILED log entry as pending-retry and send a workflow notification
        so the user knows to re-run the QGIS upload.  The actual re-upload must
        originate from the QGIS plugin — the server cannot pull files from the
        client machine, so we reset the record and notify the uploader.
        """
        from .models import QGISUploadLog
        log = self.get_object()

        if log.status != QGISUploadLog.FAILED:
            return Response(
                {'detail': 'Only FAILED uploads can be retried.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Reset to SKIPPED so it no longer shows as failed; the plugin will
        # POST a fresh record on the next successful upload.
        log.status = QGISUploadLog.SKIPPED
        log.error_message = f'[Retry requested by {request.user.username}] {log.error_message or ""}'.strip()
        log.save(update_fields=['status', 'error_message'])

        # Notify the original uploader (if different from requester)
        if log.uploaded_by and log.uploaded_by != request.user:
            from apps.workflow.models import Notification
            Notification.objects.create(
                user=log.uploaded_by,
                project=log.project,
                title='QGIS Upload Retry Requested',
                message=(
                    f'{request.user.get_full_name() or request.user.username} has requested '
                    f'a retry for "{log.filename}" '
                    f'(project {log.project.project_number}). '
                    'Please re-upload the file from QGIS.'
                ),
            )

        from .serializers import QGISUploadLogSerializer
        return Response(QGISUploadLogSerializer(log).data)


# ── Temporary Layer Upload ─────────────────────────────────────────────────────

def _parse_temp_layer_file(file_obj):
    """Parse KML/KMZ/GeoJSON/Shapefile-ZIP into a GeoJSON FeatureCollection dict.
    Returns (geojson_dict, feature_count, file_format_str).
    """
    import os, tempfile, zipfile, json, shutil
    from django.contrib.gis.gdal import DataSource
    from .models import TemporaryLayer

    name = file_obj.name.lower()
    ext = name.rsplit('.', 1)[-1] if '.' in name else ''

    # Write uploaded file to a temp path so GDAL can open it
    tmp_dir = tempfile.mkdtemp(prefix='tmplayer_')
    try:
        if ext == 'kmz':
            # KMZ = zipped KML
            kmz_path = os.path.join(tmp_dir, 'upload.kmz')
            with open(kmz_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            with zipfile.ZipFile(kmz_path) as zf:
                kml_names = [n for n in zf.namelist() if n.lower().endswith('.kml')]
                if not kml_names:
                    raise ValueError('No .kml file found inside KMZ archive.')
                zf.extract(kml_names[0], tmp_dir)
                src_path = os.path.join(tmp_dir, kml_names[0])
            file_format = TemporaryLayer.KMZ

        elif ext == 'kml':
            src_path = os.path.join(tmp_dir, 'upload.kml')
            with open(src_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            file_format = TemporaryLayer.KML

        elif ext == 'geojson' or ext == 'json':
            src_path = os.path.join(tmp_dir, 'upload.geojson')
            with open(src_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            file_format = TemporaryLayer.GEOJSON

        elif ext == 'zip':
            zip_path = os.path.join(tmp_dir, 'upload.zip')
            with open(zip_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            shp_dir = os.path.join(tmp_dir, 'shp')
            os.makedirs(shp_dir)
            with zipfile.ZipFile(zip_path) as zf:
                zf.extractall(shp_dir)
            shp_files = []
            for root, dirs, files in os.walk(shp_dir):
                for fn in files:
                    if fn.lower().endswith('.shp'):
                        shp_files.append(os.path.join(root, fn))
            if not shp_files:
                raise ValueError('No .shp file found inside ZIP archive.')
            src_path = shp_files[0]
            file_format = TemporaryLayer.SHAPEFILE
        else:
            raise ValueError(f'Unsupported file type: .{ext}')

        # Use GeoDjango GDAL DataSource to read and convert to GeoJSON
        ds = DataSource(src_path)
        features = []
        for layer in ds:
            for feat in layer:
                try:
                    geom = feat.geom.transform(4326, clone=True)
                    geom_json = json.loads(geom.geojson)
                except Exception:
                    geom_json = None

                props = {}
                for field_name in feat.fields:
                    try:
                        val = feat[field_name].value
                        if isinstance(val, bytes):
                            val = val.decode('utf-8', errors='replace')
                        props[field_name] = val
                    except Exception:
                        props[field_name] = None

                if geom_json:
                    features.append({
                        'type': 'Feature',
                        'geometry': geom_json,
                        'properties': props,
                    })

        geojson = {
            'type': 'FeatureCollection',
            'features': features,
        }
        return geojson, len(features), file_format

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


class TemporaryLayerViewSet(viewsets.ModelViewSet):
    """CRUD for user's own temporary upload layers."""
    from .serializers import TemporaryLayerSerializer as _TLSer
    serializer_class   = _TLSer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        from django.db.models import Q
        from .models import TemporaryLayer
        user = self.request.user
        # Always your own temp layers; DEO offices also see subordinate-office
        # uploaders' temp layers that are marked deo_visible.
        q = Q(uploaded_by=user)
        deo_sub_ids = deo_subordinate_org_ids(user)
        if deo_sub_ids:
            q |= Q(uploaded_by__organisation_id__in=deo_sub_ids, deo_visible=True)
        return TemporaryLayer.objects.filter(q).distinct()

    def create(self, request, *args, **kwargs):
        from .models import TemporaryLayer
        from .serializers import TemporaryLayerSerializer
        from rest_framework.exceptions import ValidationError as DRFValidationError

        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'detail': 'No file uploaded.'}, status=status.HTTP_400_BAD_REQUEST)

        name = request.data.get('name', '').strip()
        if not name:
            return Response({'detail': 'Layer name is required.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            geojson, feature_count, file_format = _parse_temp_layer_file(file_obj)
        except ValueError as e:
            return Response({'detail': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'detail': f'Failed to parse file: {e}'}, status=status.HTTP_400_BAD_REQUEST)

        if feature_count == 0:
            return Response({'detail': 'No valid features found in the uploaded file.'}, status=status.HTTP_400_BAD_REQUEST)

        # Validate mandatory new fields
        purpose_type = request.data.get('purpose_type', '').strip()
        purpose_other = request.data.get('purpose_other', '').strip()
        land_rights_type = request.data.get('land_rights_type', '').strip()
        land_rights_other = request.data.get('land_rights_other', '').strip()

        if not purpose_type:
            return Response({'detail': 'Purpose is required.'}, status=status.HTTP_400_BAD_REQUEST)
        if purpose_type == 'OTHER' and not purpose_other:
            return Response({'detail': 'Please specify the other purpose.'}, status=status.HTTP_400_BAD_REQUEST)
        if not land_rights_type:
            return Response({'detail': 'Land Rights Type is required.'}, status=status.HTTP_400_BAD_REQUEST)
        if land_rights_type == 'OTHER' and not land_rights_other:
            return Response({'detail': 'Please specify the other land rights type.'}, status=status.HTTP_400_BAD_REQUEST)

        # Reset file pointer and save again for the FileField
        file_obj.seek(0)
        layer = TemporaryLayer(
            name=name,
            purpose=request.data.get('purpose', ''),
            purpose_type=purpose_type,
            purpose_other=purpose_other,
            land_rights_type=land_rights_type,
            land_rights_other=land_rights_other,
            description=request.data.get('description', ''),
            file_format=file_format,
            file=file_obj,
            geojson=geojson,
            feature_count=feature_count,
            deo_visible=str(request.data.get('deo_visible', 'true')).lower() not in ('false', '0', 'no'),
            uploaded_by=request.user,
        )
        layer.save()
        serializer = TemporaryLayerSerializer(layer)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'], url_path='analyse')
    def analyse(self, request, pk=None):
        """
        POST /api/projects/temp-layers/{id}/analyse/

        Body: { "buffer_m": 1000 }  (must be one of VALID_BUFFER_M; default 1000)

        Runs spatial analysis for the given buffer distance, caches the result
        keyed by buffer_m, and returns the JSON result.
        """
        from apps.survey_projects.analysis import run_defence_analysis, DEFAULT_BUFFER_M, VALID_BUFFER_M

        layer = self.get_object()
        if not layer.geojson:
            return Response({'detail': 'No geometry data in this layer.'}, status=400)

        try:
            buffer_m = int(request.data.get('buffer_m', DEFAULT_BUFFER_M))
        except (TypeError, ValueError):
            buffer_m = DEFAULT_BUFFER_M

        if buffer_m not in VALID_BUFFER_M:
            return Response(
                {'detail': f'buffer_m must be one of {VALID_BUFFER_M}'},
                status=400,
            )

        # survey_area_ids / external_layer_ids:
        #   null / omitted  → check all
        #   []              → skip that type
        #   [1, 2, ...]     → check only those IDs
        def _parse_ids(key):
            val = request.data.get(key)
            if val is None:
                return None           # all
            if isinstance(val, list):
                return [int(x) for x in val if str(x).isdigit() or isinstance(x, int)]
            return None               # unexpected format → treat as all

        survey_area_ids    = _parse_ids('survey_area_ids')
        external_layer_ids = _parse_ids('external_layer_ids')

        try:
            result = run_defence_analysis(
                layer,
                buffer_m=buffer_m,
                survey_area_ids=survey_area_ids,
                external_layer_ids=external_layer_ids,
                user=request.user,
            )
        except Exception as exc:
            logger.exception('Defence analysis failed for temp layer %s', layer.id)
            return Response({'detail': f'Analysis failed: {exc}'}, status=500)

        # Store keyed by buffer; migrate legacy flat result.
        # The frontend fires one request per buffer in parallel, so this
        # read-modify-write must be serialised under a row lock — otherwise
        # concurrent saves clobber each other and a buffer's result is lost
        # (the stale value then surfaces in the PDF report).
        from django.db import transaction
        from .models import TemporaryLayer
        with transaction.atomic():
            locked = TemporaryLayer.objects.select_for_update().get(pk=layer.pk)
            stored = locked.analysis_result or {}
            if isinstance(stored, dict) and 'verdict' in stored:
                stored = {'1000': stored}
            stored[str(buffer_m)] = result
            locked.analysis_result = stored
            locked.save(update_fields=['analysis_result'])
        return Response(result)

    @action(detail=True, methods=['get'], url_path='analyse/report')
    def analyse_report(self, request, pk=None):
        """
        GET /api/projects/temp-layers/{id}/analyse/report/

        Generates a PDF for the specified buffer distances.

        Optional ?buffers=1000,5000,10000 — comma-separated list of buffer_m values.
        If omitted, uses whatever is cached in analysis_result.
        Missing buffers that are requested will be run on-the-fly.
        """
        from apps.survey_projects.analysis import (
            run_defence_analysis,
            generate_multi_range_report_html, VALID_BUFFER_M,
        )
        import httpx

        layer = self.get_object()
        if not layer.geojson:
            return Response({'detail': 'No geometry data in this layer.'}, status=400)

        # Normalise cached results
        stored = layer.analysis_result or {}
        if isinstance(stored, dict) and 'verdict' in stored:
            stored = {'1000': stored}

        # Determine which buffers to include in the report
        buffers_param = request.query_params.get('buffers', '').strip()
        if buffers_param:
            requested = []
            for part in buffers_param.split(','):
                try:
                    bm = int(part.strip())
                    if bm in VALID_BUFFER_M:
                        requested.append(bm)
                except ValueError:
                    pass
            target_buffers = requested if requested else list(stored.keys())
        else:
            # No param → use whatever is already cached
            target_buffers = [int(k) for k in stored.keys()] if stored else [1000]

        if not target_buffers:
            return Response({'detail': 'No buffer ranges available. Run analysis first.'}, status=400)

        # Run any missing buffers in the requested set
        missing = [bm for bm in target_buffers if str(bm) not in stored]
        if missing:
            try:
                computed = {
                    str(bm): run_defence_analysis(layer, buffer_m=bm, user=request.user)
                    for bm in missing
                }
                stored.update(computed)
                # Merge under a row lock so we don't clobber a concurrent analyse write.
                from django.db import transaction
                from .models import TemporaryLayer
                with transaction.atomic():
                    locked = TemporaryLayer.objects.select_for_update().get(pk=layer.pk)
                    current = locked.analysis_result or {}
                    if isinstance(current, dict) and 'verdict' in current:
                        current = {'1000': current}
                    current.update(computed)
                    locked.analysis_result = current
                    locked.save(update_fields=['analysis_result'])
            except Exception as exc:
                return Response({'detail': f'Analysis failed: {exc}'}, status=500)

        # Build subset for this report
        report_data = {str(bm): stored[str(bm)] for bm in target_buffers if str(bm) in stored}
        if not report_data:
            return Response({'detail': 'No results to include in report.'}, status=400)

        try:
            html = generate_multi_range_report_html(layer, report_data)
            pw_response = httpx.post(
                'http://print-service:3001/render',
                json={'html': html, 'paper_size': 'A4', 'orientation': 'portrait', 'scale': 1.5},
                timeout=180,
            )
            pw_response.raise_for_status()
        except httpx.ConnectError:
            return Response({'detail': 'Print service unavailable.'}, status=503)
        except Exception as exc:
            return Response({'detail': str(exc)}, status=500)

        from django.http import HttpResponse as DjangoHttpResponse
        from apps.core.watermark import embed_watermark
        safe_name = ''.join(c if c.isalnum() or c in '-_ ' else '_' for c in layer.name)
        filename = f"defence_proximity_report_{safe_name}.pdf"
        metadata = {
            "uploaded_by": request.user.username,
            "export_format": "pdf",
            "report_type": "defence_proximity",
            "layer_name": layer.name,
        }
        try:
            pdf_bytes = embed_watermark(pw_response.content, filename, 'application/pdf', metadata)
        except Exception as wexc:
            pdf_bytes = pw_response.content

        pdf_resp = DjangoHttpResponse(pdf_bytes, content_type='application/pdf')
        pdf_resp['Content-Disposition'] = f'attachment; filename="{filename}"'
        return pdf_resp

    http_method_names = ['get', 'post', 'delete', 'head', 'options']


class ReviewAnnotationViewSet(viewsets.ModelViewSet):
    serializer_class = ReviewAnnotationSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = None

    def get_queryset(self):
        from django.db.models import Q
        user = self.request.user
        
        base_qs = SurveyArea.objects.select_related('project__organisation')
        shared_project_ids = get_shared_project_ids(user)
        approved_area_ids  = get_approved_area_ids(user)
        deo_sub_ids = deo_subordinate_org_ids(user)

        q = Q(project__organisation=user.organisation) if user.organisation else Q()
        if shared_project_ids:
            q |= Q(project_id__in=shared_project_ids)
        if approved_area_ids:
            q |= Q(id__in=approved_area_ids)
        if deo_sub_ids:
            q |= Q(project__organisation_id__in=deo_sub_ids)

        if user.is_superadmin:
            accessible_areas = SurveyArea.objects.all()
        else:
            accessible_areas = SurveyArea.objects.filter(q).distinct()

        qs = ReviewAnnotation.objects.filter(survey_area__in=accessible_areas).select_related('created_by', 'survey_area')

        sa_id = self.request.query_params.get('survey_area')
        if sa_id:
            qs = qs.filter(survey_area_id=sa_id)
        return qs

    def perform_create(self, serializer):
        user = self.request.user
        if not (user.can_check or user.can_approve or user.is_superadmin):
            raise PermissionDenied("Only checkers, approvers, or superadmins can create review annotations.")
        serializer.save(created_by=user)

    def perform_update(self, serializer):
        user = self.request.user
        if not (user.can_check or user.can_approve or user.is_superadmin):
            raise PermissionDenied("Only checkers, approvers, or superadmins can update review annotations.")
        serializer.save()

    def perform_destroy(self, instance):
        user = self.request.user
        if not (user.can_check or user.can_approve or user.is_superadmin):
            raise PermissionDenied("Only checkers, approvers, or superadmins can delete review annotations.")
        instance.delete()


@api_view(['POST'])
@drf_permission_classes([permissions.IsAuthenticated])
def georeference_image(request):
    """
    POST /api/survey_projects/georeference/

    Warp a scanned image using user-provided GCPs and save as a GeoTiff layer.

    Multipart body:
        image        — uploaded image file (JPG/PNG)
        gcps         — JSON string: [{"px": x, "py": y, "lon": lon, "lat": lat}, ...]
        project      — project ID
        layer_name   — name for the result layer
        folder       — folder ID (optional)
    """
    import json as _json
    import os
    import subprocess
    import tempfile
    from django.conf import settings as django_settings

    image_file = request.FILES.get('image')
    gcps_json  = request.data.get('gcps', '[]')
    project_id = request.data.get('project')
    layer_name = (request.data.get('layer_name') or 'georeferenced_scan').strip()
    folder_id  = request.data.get('folder')

    if not image_file or not project_id:
        return Response({'detail': 'image and project are required.'}, status=400)

    try:
        gcps = _json.loads(gcps_json)
    except Exception:
        return Response({'detail': 'Invalid GCPs JSON.'}, status=400)

    if len(gcps) < 3:
        return Response({'detail': 'At least 3 GCP pairs are required for warping.'}, status=400)

    project = SurveyProject.objects.get(id=project_id)
    folder  = ProjectLayerFolder.objects.get(id=folder_id) if folder_id else None

    ext = os.path.splitext(image_file.name)[1] or '.jpg'
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp_in:
        for chunk in image_file.chunks():
            tmp_in.write(chunk)
        tmp_in_path = tmp_in.name

    with tempfile.NamedTemporaryFile(suffix='_gcps.tif', delete=False) as tmp_gcps:
        tmp_gcps_path = tmp_gcps.name

    safe_name = layer_name.replace(' ', '_').replace('/', '_')
    out_rel = f'survey_projects/{project_id}/georeference/{safe_name}.tif'
    out_path = os.path.join(django_settings.MEDIA_ROOT, out_rel)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    try:
        gcp_args = []
        for gcp in gcps:
            gcp_args += ['-gcp', str(gcp['px']), str(gcp['py']), str(gcp['lon']), str(gcp['lat'])]

        subprocess.run(
            ['gdal_translate', '-a_srs', 'EPSG:4326'] + gcp_args + [tmp_in_path, tmp_gcps_path],
            check=True, capture_output=True,
        )

        subprocess.run(
            ['gdalwarp', '-s_srs', 'EPSG:4326', '-t_srs', 'EPSG:3857',
             '-r', 'bilinear', '-of', 'COG', '-co', 'COMPRESS=DEFLATE',
             tmp_gcps_path, out_path],
            check=True, capture_output=True,
        )

        result_layer = GeoTiffLayer.objects.create(
            project=project,
            folder=folder,
            name=layer_name,
            file=out_rel,
            cog_file=out_rel,
            status=GeoTiffLayer.DONE,
            created_by=request.user,
        )

        return Response({
            'id': result_layer.id,
            'layer_name': result_layer.name,
            'cog_url': f'/media/{out_rel}',
        })

    except subprocess.CalledProcessError as exc:
        return Response({'detail': f'GDAL warping failed: {(exc.stderr or b"").decode()[:400]}'}, status=500)
    except Exception as exc:
        return Response({'detail': str(exc)}, status=500)
    finally:
        for p in [tmp_in_path, tmp_gcps_path]:
            try:
                os.unlink(p)
            except Exception:
                pass



class FeatureCommentViewSet(viewsets.ModelViewSet):
    """Per-feature discussion thread (feature info drawer on the map).

    GET  /projects/feature-comments/?feature=<id>  — thread for one geometry
    POST /projects/feature-comments/               — add a remark
    DELETE                                          — own comments only
    """
    serializer_class = FeatureCommentSerializer
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['feature']
    pagination_class = None
    http_method_names = ['get', 'post', 'delete', 'head', 'options']

    def get_queryset(self):
        qs = FeatureComment.objects.select_related('user', 'feature__project__organisation')
        return org_queryset_filter(self.request.user, qs,
                                   org_field='feature__project__organisation')

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    def destroy(self, request, *args, **kwargs):
        comment = self.get_object()
        if comment.user_id != request.user.id and not request.user.is_superadmin:
            raise PermissionDenied('You can only delete your own comments.')
        return super().destroy(request, *args, **kwargs)

"""
Generate RakshaGIS Project Write-Up as a Word (.docx) document.
Run: python3 generate_writeup.py
Output: RakshaGIS_Project_WriteUp.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette (hex strings) ─────────────────────────────────────────────
NAVY      = '0A1E3C'
OLIVE     = '4A6A00'
ACCENT    = '1F5EA8'
LIGHT_BG  = 'F0F4F8'
WHITE     = 'FFFFFF'
MID_GRAY  = '555555'
DARK_TEXT = '1A1A2E'


def rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_h_rule(doc, color=NAVY, thick='6'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), thick)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = rgb(NAVY)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), NAVY)
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = rgb(ACCENT)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = rgb(OLIVE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = rgb(DARK_TEXT)
    return p


def add_bullet(doc, label, detail, label_color=ACCENT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(1)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = rgb(label_color)
    r2 = p.add_run(detail)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = rgb(DARK_TEXT)
    return p


def add_plain_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb(DARK_TEXT)
    return p


def add_table(doc, rows, header=None, col_widths=None):
    cols = len(rows[0]) if rows else 2
    table = doc.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        hrow = table.add_row()
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            set_cell_bg(cell, NAVY)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = rgb(WHITE)

    for ridx, row_data in enumerate(rows):
        r = table.add_row()
        bg = LIGHT_BG if ridx % 2 == 0 else WHITE
        for cidx, val in enumerate(row_data):
            cell = r.cells[cidx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = rgb(DARK_TEXT)
            if cidx == 0:
                run.bold = True

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────────────────────
#  BUILD THE DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ── Cover block ───────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(4)
title_p.paragraph_format.space_after  = Pt(0)
r1 = title_p.add_run('RAKSHA')
r1.bold = True; r1.font.size = Pt(36); r1.font.color.rgb = rgb(NAVY)
r2 = title_p.add_run('GIS')
r2.bold = True; r2.font.size = Pt(36); r2.font.color.rgb = rgb(ACCENT)

sub = doc.add_paragraph('Defence Estate GIS Survey Platform')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
for run in sub.runs:
    run.font.size = Pt(14); run.font.color.rgb = rgb(OLIVE); run.bold = True

org_p = doc.add_paragraph('Directorate General of Defence Estates (DGDE)  |  Government of India')
org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
org_p.paragraph_format.space_after = Pt(2)
for run in org_p.runs:
    run.font.size = Pt(11); run.font.color.rgb = rgb(MID_GRAY)

date_p = doc.add_paragraph(datetime.date.today().strftime('%B %Y'))
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(12)
for run in date_p.runs:
    run.font.size = Pt(10); run.font.color.rgb = rgb(MID_GRAY)

add_h_rule(doc, NAVY, '12')

# ── 1. Project Title ──────────────────────────────────────────────────────────
add_heading(doc, '1. Project Title', level=1)

t = doc.add_table(rows=1, cols=1)
t.style = 'Table Grid'
cell = t.rows[0].cells[0]
set_cell_bg(cell, NAVY)
cell.width = Inches(6.2)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RakshaGIS — Defence Estate Geographic Information System & Survey Platform')
run.bold = True; run.font.size = Pt(12); run.font.color.rgb = rgb(WHITE)
doc.add_paragraph()

add_body(doc,
    'RakshaGIS is a full-stack, self-hosted enterprise GIS platform developed for the Directorate '
    'General of Defence Estates (DGDE) to digitise, manage, review, and publish defence land surveys '
    'across India. The system consolidates field surveying, internal approval workflows, in-browser '
    'document editing, AI-assisted document processing, and map publishing into a single role-gated '
    'application — hosted entirely on-premise with zero dependency on commercial cloud services or '
    'internet access after initial installation.')

# ── 2. Objective ──────────────────────────────────────────────────────────────
add_heading(doc, '2. Objective of the Project', level=1)

add_body(doc,
    'Defence estate land management has historically relied on paper maps, disconnected spreadsheets, '
    'and manual review processes — creating delays, data inconsistencies, and security risks. '
    'RakshaGIS was commissioned to address these challenges with the following objectives:')

for label, detail in [
    ('Digitise land surveys',
     'Replace paper-based field surveys with an interactive web GIS that lets authorised surveyors draw, '
     'annotate, and submit georeferenced features directly from their workstations.'),
    ('Enforce a formal, survey-area-wise approval chain',
     'Implement a role-based state machine (Draft → Submitted → Under Review → Approved → Published) '
     'operating at the survey-area (pocket) level, so different pockets within the same project can '
     'progress independently through Checker and Approver review.'),
    ('Maintain data sovereignty',
     'Keep all spatial data, documents, and AI inference within DGDE infrastructure — '
     'no data ever leaves the on-premise deployment. The system operates fully offline after '
     'initial installation on a dedicated secure line.'),
    ('Organisation-wise data isolation',
     'Enforce strict row-level data isolation so each office (DEO/CEO/ADEO) sees only its own data, '
     'while PDDE offices see their command jurisdiction and DGDE has all-India visibility. '
     'Controlled cross-org access is available via a formal request-and-approval mechanism.'),
    ('Provide spatial analytics',
     'Enable buffer analysis, topology checking, version comparison, and COG raster overlay in the browser, '
     'eliminating the need for separate GIS software licences such as ArcGIS.'),
    ('AI-assisted reporting',
     'Automatically summarise uploaded survey documents and generate structured reports '
     'using a local large language model (Ollama / LocalAI) — without internet access or API subscriptions. '
     'Generated reports open directly in an integrated online document editor.'),
    ('Scalable multi-organisation support',
     'Support a five-level organisational hierarchy (DGDE → PDDE → DEO → CEO → ADEO) '
     'with per-organisation row-level data isolation and delegated administration.'),
]:
    add_bullet(doc, label, detail, ACCENT)

# ── 3. Key Features / Work Done ───────────────────────────────────────────────
add_heading(doc, '3. Key Features / Work Done', level=1)

# 3.1
add_heading(doc, '3.1  Interactive Map & Spatial Analysis', level=2)
for text in [
    'Multi-basemap support — OSM, XYZ tile layers, WMS/WMTS, and Bhuvan (ISRO) satellite imagery.',
    'Local offline raster tile server — India OSM tiles served by overv/openstreetmap-tile-server after '
    'one-time import (./build.sh --import-osm, ~800 MB India PBF from Geofabrik). '
    'No internet required after import; tiles cached with 7-day immutable headers.',
    'Feature drawing with snapping: Point, Line, and Polygon tools with snap-to-vertex and live area/perimeter feedback while drawing.',
    'Edit existing features — move and reshape vertices; changes are automatically persisted to the database.',
    'Box Select tool — drag a rectangle to highlight features; selected count shown in an overlay badge.',
    'Identify tool — click any feature to inspect all attributes in a side drawer.',
    'Buffer Analysis — N configurable distance rings (metres/kilometres); PostGIS geographically-accurate buffer intersected against all defence parcels; results in a colour-coded tabbed modal with mini-map; downloadable as Excel (one sheet per ring) and PDF.',
    'Topology Checker — detects invalid geometries and overlapping parcels via PostGIS ST_IsValid and ST_Intersects; issues listed in a modal with parcel identifiers.',
    'Measure tool — real-time line length display as the user draws.',
    'Cloud-Optimized GeoTIFF (COG) overlay — upload satellite imagery; Celery converts to COG; per-layer opacity and visibility controls.',
    'Shapefile import — upload a zipped .shp archive; background Celery task parses features and bulk-creates them into the selected project.',
    'Attribute Table Panel — bottom-docked spreadsheet of all project features; inline attribute editing; Field Calculator with ${field} expression support; CSV export; double-click row zooms the map to that feature.',
    'Print to PDF — jsPDF A4/A3/Letter layout with captured map canvas, north arrow, scale bar, legend, optional coordinate grid, and DGDE-branded header/footer.',
    'Map Bookmarks — save and recall named map extents stored in localStorage.',
    'Go-to Coordinate — jump directly to any Lat/Lon position.',
    'Per-layer symbology — colour picker per layer name; label toggle renders feature IDs as map text.',
    'Admin boundary tile overlay from pg_tileserv (MapVector Tile) showing District boundaries.',
    'Undo last drawn feature — removes from map source and calls DELETE on the backend.',
    'Auto-load: the map remembers the last active project across browser sessions so layers are visible immediately on open without manual project selection.',
]:
    add_plain_bullet(doc, text)

# 3.2
add_heading(doc, '3.2  Survey Workflow — Survey-Area-Wise Approval', level=2)
add_body(doc,
    'A key architectural principle of RakshaGIS is that workflow operates at the survey-area (pocket) level, '
    'not at the project level. This means different sections of the same project can progress through '
    'Checker and Approver review independently.')
for text in [
    'Survey Areas: SDOs/Surveyors create named survey areas within a project and link each area to a '
    'project folder; the workflow tracks that pocket individually.',
    'State machine per area: DRAFT → SUBMITTED → UNDER_REVIEW → APPROVED → PUBLISHED, with RETURNED return paths at each stage.',
    'Role-based transitions: SDO/Surveyor submits to Checker; Checker sends to Approver or returns to SDO with mandatory remarks; Approver approves or returns from review; DEO Admin publishes.',
    'Mandatory remarks on return: Return transitions require a remarks field so the submitter understands what to correct.',
    'Folder lock on submission: when a survey area is submitted, its linked folder and all descendant sub-folders (Doc, Shapefile, Raster) are immediately read-only — no draw, edit, delete, upload, or import operations are allowed until the area is returned for revision.',
    'Visual lock indicators: locked folders display a gold lock icon in the folder tree; the map toolbar hides all write tools and shows a banner.',
    'Per-area audit log: every transition, with actor, remarks, and timestamp, is displayed inline on each survey area card in the project detail page.',
    'In-app notifications: all affected users are notified on every state transition.',
    'Versioned layer folders: Phase → Zone → Year → Ver-I / Ver-II / … / Final, auto-created on first use. Active version is detected or created automatically; Final folder is auto-created and features are copied when a survey area is approved.',
    'Project sharing: grant read access to specific users outside the creating organisation.',
]:
    add_plain_bullet(doc, text)

# 3.3
add_heading(doc, '3.3  Version Comparison — Split-Screen Map', level=2)
for text in [
    'Split-screen map view comparing any two VERSION folders within a project side-by-side.',
    'Shared OpenLayers View instance — panning or zooming one panel automatically mirrors the other.',
    'Version A features rendered in blue; Version B in orange for instant visual diff.',
    'Version selectors: dropdowns listing only VERSION-type folders in the project.',
]:
    add_plain_bullet(doc, text)

# 3.4
add_heading(doc, '3.4  Organisation-Wise Data Isolation & Cross-Org Access', level=2)
add_body(doc,
    'All data in RakshaGIS is scoped strictly to the user\'s organisation. '
    'A controlled request-and-approval mechanism allows DEO offices to request read access '
    'to specific survey areas from neighbouring CEO/ADEO offices under the same PDDE command.')
for text in [
    'Strict data isolation: each organisation (DEO/CEO/ADEO) can only see its own projects, survey areas, features, folders, and documents.',
    'DGDE (SUPERADMIN/VIEWER): all-India access — reads all organisations\' data.',
    'PDDE (PDDE_VIEWER): command-level access — reads all data within own jurisdiction subtree.',
    'DEO / CEO / ADEO: own-organisation data only by default.',
    'Cross-org access requests: DEO browses a discovery list of sibling-org survey areas (metadata only — name, organisation, status; no GIS data) and submits a formal access request with reason.',
    'Approval flow: request goes to the target org\'s admin; once approved, requesting org can view features, folders, and documents for that survey area.',
    'Access Requests page: three-tab UI — Discover Areas (sibling orgs under same PDDE), My Requests (outgoing with status), Incoming Requests (admin-only with Approve/Reject and mandatory remarks).',
    'Pending badge: sidebar "Data Access" item shows a live count of pending incoming requests for admins, refreshed every 60 seconds.',
    'ProjectShare: projects can also be shared directly to specific users for read access.',
]:
    add_plain_bullet(doc, text)

# 3.5
add_heading(doc, '3.5  Online Document Editing (OnlyOffice)', level=2)
add_body(doc,
    'RakshaGIS integrates OnlyOffice Community Document Server as a self-hosted, fully offline '
    'document editor. Users can view and edit Word, Excel, PowerPoint, and PDF files directly '
    'in the browser without any external service.')
for text in [
    'OnlyOffice Community Document Server 8.2.2 runs as a Docker service within the internal network — no licence, no registration, no internet connection.',
    'Supports .docx, .odt, .xls, .xlsx, .pptx, .ppt, .pdf — full editing capability for all office formats.',
    'JWT-signed configuration: Django generates a PyJWT-signed editor config for each document open; OnlyOffice validates the token before serving the editor.',
    'Edit callback: when a user saves a document in OnlyOffice, the updated file is automatically downloaded by Django and saved back (version incremented, file overwritten).',
    'AI-generated survey reports open directly in OnlyOffice for review and annotation without downloading.',
    'Any document in a project\'s folder tree can be opened with a single click on the OnlyOffice icon next to supported file types.',
    'Full-screen embedded editor (95vw × 90vh modal) with fallback download button if the editor cannot load.',
    'OnlyOffice spell-check and telemetry are disabled for fully offline operation.',
]:
    add_plain_bullet(doc, text)

# 3.6
add_heading(doc, '3.6  User & Organisation Management', level=2)
for text in [
    'Five-level organisation hierarchy: DGDE → PDDE → DEO → CEO → ADEO.',
    'Ten distinct roles: SUPERADMIN, PDDE_VIEWER, VIEWER, DEO_ADMIN, CEO_ADMIN, ADEO_ADMIN, SDO, SURVEYOR, CHECKER, APPROVER — with fine-grained permission classes on every API endpoint.',
    'Organisation records extended with Office ID (5-character alphanumeric code), Officer Name, mobile, landline, email, address, pincode, State and District linkage with cascade-validated dropdowns.',
    'Admin-role protection: DEO/CEO/ADEO/SUPERADMIN accounts cannot be deleted or deactivated — prevents accidental system lockout.',
    'Force-logout action sets is_active=False; JWT tokens become invalid immediately on the next request without any token blacklist database overhead.',
    'Password management: users change own password with current-password verification; admins can set any non-admin user\'s password; separate "Change My Password" endpoint in the header dropdown.',
]:
    add_plain_bullet(doc, text)

# 3.7
add_heading(doc, '3.7  Master Data Management (SuperAdmin)', level=2)
for text in [
    'Full CRUD for State, District, Taluk, and Village with cascading dropdowns (selecting a State filters Districts, selecting a District filters Taluks, and so on).',
    'PostGIS geometry fields on each hierarchy level — boundaries can be stored and visualised on the map.',
    'SUPERADMIN-only access enforced at API and UI level.',
    'Dedicated master data pages: /master/states, /master/districts, /master/taluks, /master/villages.',
]:
    add_plain_bullet(doc, text)

# 3.8
add_heading(doc, '3.8  Documents & AI Assistant', level=2)
for text in [
    'Per-project document upload (PDF, images) with MIME-type validation using python-magic.',
    'Background AI processing via Ollama — text extracted by pdfplumber, summarised by the local LLM.',
    'Auto-generated structured survey report per project; report opens directly in the OnlyOffice editor for review and annotation.',
    'Interactive chat interface: users ask free-text questions; the local LLM answers with project document context.',
    'Multiple AI backend options selectable via Docker Compose profiles: Ollama, LocalAI (AIO CPU/GPU), llama.cpp (CPU/GPU), AnythingLLM — all on-premise, all offline.',
    'AI model downloads are a one-time SuperAdmin task after installation; no internet access required thereafter.',
    'All AI inference runs on-premise — no data leaves the deployment host.',
]:
    add_plain_bullet(doc, text)

# 3.9
add_heading(doc, '3.9  UI Themes', level=2)
for text in [
    'Six built-in themes: Dark, Light, Navy, Forest, Midnight, Saffron.',
    'All UI surfaces — header, sidebar, drawers, map panels — use CSS custom properties (--bg-base, --bg-card, --accent, etc.) ensuring complete theme coverage.',
    'Theme selection persisted in localStorage across sessions.',
]:
    add_plain_bullet(doc, text)

# 3.10
add_heading(doc, '3.10  Infrastructure, DevOps & Offline Operation', level=2)
add_body(doc,
    'RakshaGIS is designed for fully offline operation on a dedicated secure line after initial installation. '
    'No service requires internet access during normal operation.')
for text in [
    'Fully containerised via Docker Compose: PostgreSQL/PostGIS 16, Redis 7, Celery workers, Gunicorn, Nginx 1.27, pg_tileserv, OnlyOffice 8.2.2.',
    'Two-network Docker design: raksha-net (internal: true) blocks outbound internet from all backend containers; raksha-edge (external) is attached to nginx only, enabling port 80/443 publishing to clients.',
    'All Docker image versions pinned — build.sh --save-images bundles all images into a single archive for air-gapped deployment; --load-images restores them on the target machine without any internet access.',
    'Local offline raster tile server (overv/openstreetmap-tile-server:2.3.0) in compose profile "osm" — started automatically after build.sh --import-osm; serves India OSM tiles at /osm-tiles/{z}/{x}/{y}.png.',
    'India PBF import (./build.sh --import-osm): downloads india-latest.osm.pbf from Geofabrik (~800 MB, one-time), runs osm2pgsql import inside the tile server container (2–4 hours), then starts the osm-tiles service.',
    'Certbot (HTTPS) in profile "https" — not started on dedicated/intranet deployments; suitable for public-domain SSL activation when needed.',
    'Prometheus + Grafana in profile "monitoring" — not started by default; enable for Phase 5 dashboard setup.',
    'build.sh — smart Docker image build that skips rebuilding when Dockerfile and requirements are unchanged (SHA-256 hash comparison).',
    'RakshaGIS.sh — one-command service manager: start, stop, status, logs, backup, restore, Django management commands.',
    'Layer exports: GeoJSON, Shapefile, CSV, KML, GeoPackage via Fiona/Shapely.',
    'django-prometheus metrics exposed at /metrics/ for Prometheus scraping.',
    'extra_hosts: host.docker.internal:host-gateway — allows containers to reach Ollama running on the host machine even with internal networking.',
]:
    add_plain_bullet(doc, text)

# ── 4. Technologies / Tools Used ─────────────────────────────────────────────
add_heading(doc, '4. Technologies / Tools Used', level=1)

tech_rows = [
    ('Backend framework',     'Python 3.11 + Django 4.2',              'REST API, ORM, authentication'),
    ('REST API layer',        'Django REST Framework 3.15',             'Serialisers, viewsets, permissions'),
    ('Spatial database',      'PostgreSQL 16 + PostGIS 3.4',            'Geometry storage, spatial queries, topology'),
    ('Async task queue',      'Celery 5.3 + Redis 7',                   'COG conversion, shapefile import, AI tasks'),
    ('Vector tile server',    'pg_tileserv',                            'MVT admin boundary tiles from PostGIS'),
    ('Raster tile server',    'overv/openstreetmap-tile-server 2.3.0',  'Offline India OSM raster tiles'),
    ('Document editor',       'OnlyOffice Community Server 8.2.2',      'In-browser Word/Excel/PPT/PDF editing'),
    ('AI / LLM inference',    'Ollama, LocalAI, llama.cpp, AnythingLLM','Local document summarisation and chat (profile-selectable)'),
    ('PDF text extraction',   'pdfplumber 0.11',                        'Extracts text from uploaded survey PDFs'),
    ('GIS import/export',     'Fiona 1.9 + Shapely 2.0',               'Shapefile I/O, GeoJSON, KML, GeoPackage'),
    ('HTTP client',           'httpx 0.27',                             'Ollama / LocalAI API calls from Django'),
    ('Frontend framework',    'React 18 + TypeScript + Vite',          'SPA, type safety, fast HMR builds'),
    ('Map library',           'OpenLayers 10.2',                        'Interactive map, vector/raster/MVT layers'),
    ('UI component library',  'Ant Design 5.20',                        'Tables, forms, modals, drawers, pickers'),
    ('State management',      'Zustand 4.5 + TanStack Query 5',        'Global UI state, server-state caching'),
    ('PDF generation',        'jsPDF 4.2 + jspdf-autotable 5',         'Print layout, buffer analysis PDF export'),
    ('Excel export',          'SheetJS (xlsx) 0.18',                   'Buffer analysis Excel workbook'),
    ('Map canvas capture',    'html2canvas 1.4',                        'OL canvas to image for PDF embedding'),
    ('COG display',           'geotiff.js 3.0.5',                       'WebGLTileLayer Cloud-Optimized GeoTIFF'),
    ('JWT authentication',    'SimpleJWT 5.3',                          'Short-lived access + rotating refresh tokens'),
    ('JWT document signing',  'PyJWT 2.x',                             'OnlyOffice editor config + callback signing'),
    ('Monitoring',            'Prometheus v2.55.1 + Grafana 11.4.2',   'Metrics collection, dashboards, alerting'),
    ('Web server',            'Nginx 1.27 + Gunicorn (4 workers)',      'Reverse proxy, static file serving'),
    ('Containerisation',      'Docker 24 + Docker Compose v2',         'Isolated, reproducible on-premise deployment'),
    ('API documentation',     'drf-spectacular 0.27',                   'OpenAPI 3 schema + Swagger UI'),
    ('File validation',       'python-magic 0.4',                       'MIME-type check on all uploaded files'),
]

add_table(doc, tech_rows,
          header=('Category', 'Technology / Version', 'Purpose'),
          col_widths=[1.6, 2.4, 2.6])

# ── 5. Expected Outcomes / Benefits ──────────────────────────────────────────
add_heading(doc, '5. Expected Outcomes / Benefits', level=1)

for label, detail in [
    ('End-to-end digital survey workflow',
     'Eliminates paper-based submission and manual data entry. Surveys are drawn, reviewed, and published '
     'entirely within the platform — reducing processing time from weeks to days.'),
    ('Granular, pocket-level approval',
     'The survey-area-wise workflow means different sections of the same project can progress independently '
     'through Checker and Approver sign-off, avoiding bottlenecks where a single large project blocks all work.'),
    ('Accurate, versioned spatial data',
     'Every edit is tied to a named version (Ver-I, Ver-II, Final); no data is overwritten, '
     'enabling full roll-back and side-by-side version comparison.'),
    ('Enforced approval chain',
     'The state machine prevents publication without Checker and Approver sign-off, '
     'reducing the risk of unauthorised or incomplete surveys entering the public record.'),
    ('Data integrity through folder locking',
     'Once a survey area is submitted, its entire folder tree is read-only — preventing concurrent edits '
     'that could corrupt the submitted dataset during review.'),
    ('Strict data sovereignty',
     'All data, including AI inference and document editing, remains within DGDE infrastructure. '
     'Row-level data isolation prevents cross-organisation data leakage. '
     'The system runs fully offline after installation on a dedicated secure line.'),
    ('Controlled cross-org collaboration',
     'The access request mechanism allows DEO offices to formally request and receive read access to '
     'specific survey areas from neighbouring CEO/ADEO offices — enabling data sharing without '
     'compromising isolation for other areas.'),
    ('Seamless document workflows',
     'OnlyOffice integration eliminates the download-edit-re-upload cycle. AI-generated reports '
     'open directly in the editor for review, annotation, and sign-off within the same browser tab.'),
    ('Reduced GIS software dependency',
     'Buffer analysis, topology checking, feature editing, print-to-PDF, and attribute management — '
     'capabilities previously requiring ArcGIS licences — are now available in any modern browser.'),
    ('AI-assisted productivity',
     'Automatic document summarisation and report generation reduce the time an Approver needs '
     'to review a project from hours to minutes.'),
    ('Low operational overhead',
     'Docker Compose deployment, smart build script, air-gapped image archive, and single-command service '
     'manager allow a small IT team to maintain the platform without specialist DevOps knowledge.'),
]:
    add_bullet(doc, label, detail, ACCENT)

# ── 6. Roadmap & Future Scope ─────────────────────────────────────────────────
add_heading(doc, '6. Roadmap & Future Scope', level=1)

add_body(doc,
    'Several items originally identified as future scope have since been delivered. '
    'The platform status is summarised below, followed by the enhancements that remain '
    'planned for future iterations.')

add_heading(doc, '6.1  Delivered Since Initial Scope', level=2)
for label, detail in [
    ('HTTPS with internal CA',
     'Nginx TLS is wired up: the docker-compose "https" profile and certbot service are in place; '
     'activate by supplying the DGDE internal Certificate Authority certificate.'),
    ('Prometheus / Grafana monitoring',
     'django-prometheus exposes metrics at /metrics/; Prometheus v2.55 and Grafana 11.4 ship in the '
     '"monitoring" profile. Import of DGDE-specific dashboards is the remaining step.'),
    ('Automated boundary dispute detection',
     'Implemented as a pre-submission check: a PostGIS overlap query flags newly submitted features '
     'that intersect PUBLISHED features from other organisations and records a DisputeReport.'),
    ('Advanced AI integration',
     'Parcel boundaries are auto-extracted from GeoTIFFs and scanned maps (classical GIS pipeline + '
     'AI vision pipeline) with best-effort Survey-Number OCR; local-LLM training-export and '
     'DGDE-expert model creation tooling is included for fine-tuning.'),
    ('Admin boundary data load',
     'State, district, taluk and village master tables with shapefile import enable spatial filtering '
     'and boundary overlay on the map.'),
    ('3D terrain and elevation overlay',
     'A Cesium 3D viewer provides elevation query, profile and slope analysis over SRTM/Cartosat DEM, '
     'with survey features and external layers draped on the terrain.'),
    ('Real-time collaborative editing',
     'WebSocket-based concurrent editing (Django Channels / Daphne) with live presence lets multiple '
     'surveyors work on the same project simultaneously.'),
    ('Multi-language UI',
     'Localised into English, Hindi, Tamil, Telugu, Bengali, Kannada and Marathi.'),
    ('Automated disaster-recovery backups',
     'Scheduled, encrypted PostgreSQL dumps with automated rotation are in place (off-site replication '
     'to NIC Meghraj remains optional/pending).'),
    ('Regulatory reporting',
     'Survey-area report generation (.docx, editable in OnlyOffice) and proximity / encroachment '
     'analysis reports are delivered; ministry-prescribed template formats are in progress.'),
]:
    add_bullet(doc, label, detail, OLIVE)

add_heading(doc, '6.2  Planned for Future Iterations', level=2)
for label, detail in [
    ('Mobile field surveying app',
     'A Progressive Web App (PWA) or React Native client for GPS-assisted feature capture '
     'in the field with offline sync on reconnection.'),
    ('National land registry integration',
     'REST API connectors to DILRMP (Digital India Land Records Modernisation Programme) '
     'for cross-referencing defence parcels against state revenue records.'),
    ('PKI / smart-card authentication',
     'Integration with DGDE\'s internal PKI for mutual TLS client-certificate login, '
     'replacing password-based authentication for high-security operations.'),
    ('Off-site backup replication',
     'Automated push of encrypted backups to NIC Meghraj cloud object storage for off-site recovery, '
     'extending the existing local encrypted-backup rotation.'),
    ('Ministry-format report templates',
     'Pre-built templates that compile survey statistics, ownership summaries and encroachment '
     'analysis into ministry-prescribed formats, building on the current reporting engine.'),
]:
    add_bullet(doc, label, detail, OLIVE)

# ── Footer rule ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top_el = OxmlElement('w:top')
top_el.set(qn('w:val'), 'single')
top_el.set(qn('w:sz'), '6')
top_el.set(qn('w:space'), '1')
top_el.set(qn('w:color'), NAVY)
pBdr.append(top_el)
pPr.append(pBdr)

footer_p = doc.add_paragraph(
    f'RakshaGIS — Defence Estate GIS Survey Platform  |  '
    f'Prepared: {datetime.date.today().strftime("%d %B %Y")}  |  '
    'DGDE, Government of India  |  CONFIDENTIAL'
)
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(4)
for run in footer_p.runs:
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MID_GRAY)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = 'RakshaGIS_Project_WriteUp.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
